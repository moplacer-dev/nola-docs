from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify
from flask_wtf import FlaskForm
from flask_wtf.csrf import CSRFProtect
from wtforms import StringField, TextAreaField, SelectField, SelectMultipleField, SubmitField, FieldList, FormField, IntegerField, HiddenField, BooleanField, RadioField
from wtforms.validators import DataRequired, Length, Optional, ValidationError
from docxtpl import DocxTemplate, RichText, InlineImage
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from werkzeug.utils import secure_filename
from flask_wtf.file import FileField, FileAllowed
import os
import shutil
import tempfile
from datetime import datetime
import PyPDF2
import re

# Add database and authentication imports
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_login import LoginManager, login_required, current_user, login_user, logout_user
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Update app configuration
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-key-change-this')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'temp_uploads'

# Database configuration with Render compatibility
database_url = os.environ.get('DATABASE_URL', 'sqlite:///nola_docs.db')
# Fix for Render PostgreSQL URL format
if database_url.startswith('postgres://'):
    database_url = database_url.replace('postgres://', 'postgresql://', 1)

app.config['SQLALCHEMY_DATABASE_URI'] = database_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Environment detection
app.config['IS_PRODUCTION'] = os.environ.get('FLASK_ENV') == 'production' or 'render.com' in os.environ.get('RENDER_EXTERNAL_URL', '')

# Create upload folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize database and migrations
from models import db, User, FormDraft, GeneratedDocument, TemplateFile, ActivityLog, State, Standard, Module, ModuleStandardMapping, IplModule, IplEntry
db.init_app(app)
migrate = Migrate(app, db)

# CLI commands for development/maintenance have been archived
# These scripts are now available in archive/dev-scripts/ if needed

# Helper functions for standards/modules queries
def get_framework_for(state: str, subject: str) -> str:
    """Get the framework for a given state and subject"""
    return 'CCSS-M' if subject.upper() == 'MATH' else 'NGSS'

def get_all_standards(state: str, grade: int, subject: str) -> list:
    """Get all standard codes for a given state, grade, and subject"""
    fw = get_framework_for(state, subject)
    q = Standard.query.filter_by(framework=fw, subject=subject.upper())
    
    if subject.upper() == 'MATH':
        q = q.filter_by(grade_level=int(grade))
    else:
        # For Science (NGSS), standards are now separated by grade level
        q = q.filter_by(grade_level=int(grade))
    
    return [s.code for s in q.order_by(Standard.code).all()]

def get_module_to_standards(subject: str, grade: int) -> dict:
    """Get mapping of module titles to their covered standard codes"""
    q = (db.session.query(Module.title, Standard.code)
         .join(ModuleStandardMapping, Module.id==ModuleStandardMapping.module_id)
         .join(Standard, Standard.id==ModuleStandardMapping.standard_id)
         .filter(Module.subject==subject.upper()))
    
    if subject.upper() == 'MATH':
        q = q.filter(Module.grade_level==int(grade), Standard.framework=='CCSS-M')
    else:
        # For Science (NGSS), standards are now separated by grade level
        q = q.filter(Standard.framework=='NGSS', Standard.grade_level==int(grade))
    
    out = {}
    for title, code in q:
        out.setdefault(title, set()).add(code)
    return out

# Flask-Login setup
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'auth.login'
login_manager.login_message = 'Please log in to access this page.'
login_manager.login_message_category = 'info'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# CSRF Protection setup
csrf = CSRFProtect(app)

# Register blueprints
from auth import auth
app.register_blueprint(auth)

# Helper function to escape XML special characters for DOCX
def escape_xml(text):
    """Escape special XML characters for safe inclusion in DOCX templates"""
    if not text:
        return text
    
    # XML special characters that need escaping
    replacements = {
        '&': '&amp;',
        '<': '&lt;',
        '>': '&gt;',
        '"': '&quot;',
        "'": '&apos;'
    }
    
    # Apply replacements
    for char, escaped in replacements.items():
        text = text.replace(char, escaped)
    
    return text

def load_moduleanswerkey_draft_into_form(form, user_id):
    """Load the most recent moduleanswerkey draft data into the provided form"""
    try:
        # Get the most recent draft for this user
        latest_draft = FormDraft.query.filter_by(
            user_id=user_id, 
            form_type='moduleanswerkey'
        ).order_by(FormDraft.updated_at.desc()).first()
        
        if not latest_draft:
            print("🔍 No moduleanswerkey draft found for user")
            return False
            
        print(f"🔍 Loading draft data from draft ID {latest_draft.id}")
        form_data = latest_draft.form_data
        
        # Populate basic fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        
        # Populate pre-test questions
        pretest_data = form_data.get('pretest_questions', [])
        for i, question_data in enumerate(pretest_data):
            if i < len(form.pretest_questions):
                form.pretest_questions[i].question_text.data = question_data.get('question_text', '')
                form.pretest_questions[i].choice_a.data = question_data.get('choice_a', '')
                form.pretest_questions[i].choice_b.data = question_data.get('choice_b', '')
                form.pretest_questions[i].choice_c.data = question_data.get('choice_c', '')
                form.pretest_questions[i].choice_d.data = question_data.get('choice_d', '')
                # Clean correct_answer to ensure it's valid
                correct_answer = question_data.get('correct_answer', '').strip().upper()
                form.pretest_questions[i].correct_answer.data = correct_answer if correct_answer in ['A', 'B', 'C', 'D'] else ''
        
        # Populate RCA sessions
        rca_data = form_data.get('rca_sessions', [])
        print(f"🔍 Loading RCA sessions data: {len(rca_data)} sessions found")
        for i, session_data in enumerate(rca_data):
            if i < len(form.rca_sessions):
                questions = session_data.get('questions', [])
                for j, question_data in enumerate(questions):
                    if j < len(form.rca_sessions[i].questions):
                        form.rca_sessions[i].questions[j].question_text.data = question_data.get('question_text', '')
                        form.rca_sessions[i].questions[j].choice_a.data = question_data.get('choice_a', '')
                        form.rca_sessions[i].questions[j].choice_b.data = question_data.get('choice_b', '')
                        form.rca_sessions[i].questions[j].choice_c.data = question_data.get('choice_c', '')
                        form.rca_sessions[i].questions[j].choice_d.data = question_data.get('choice_d', '')
                        # Clean correct_answer to ensure it's valid
                        correct_answer = question_data.get('correct_answer', '').strip().upper()
                        form.rca_sessions[i].questions[j].correct_answer.data = correct_answer if correct_answer in ['A', 'B', 'C', 'D'] else ''
        
        # Populate post-test questions
        posttest_data = form_data.get('posttest_questions', [])
        for i, question_data in enumerate(posttest_data):
            if i < len(form.posttest_questions):
                form.posttest_questions[i].question_text.data = question_data.get('question_text', '')
                form.posttest_questions[i].choice_a.data = question_data.get('choice_a', '')
                form.posttest_questions[i].choice_b.data = question_data.get('choice_b', '')
                form.posttest_questions[i].choice_c.data = question_data.get('choice_c', '')
                form.posttest_questions[i].choice_d.data = question_data.get('choice_d', '')
                # Clean correct_answer to ensure it's valid
                correct_answer = question_data.get('correct_answer', '').strip().upper()
                form.posttest_questions[i].correct_answer.data = correct_answer if correct_answer in ['A', 'B', 'C', 'D'] else ''
        
        # Populate PBA sessions
        pba_data = form_data.get('pba_sessions', [])
        print(f"🔍 Loading PBA sessions data: {len(pba_data)} sessions found")
        for i, session_data in enumerate(pba_data):
            if i < len(form.pba_sessions):
                form.pba_sessions[i].session_number.data = session_data.get('session_number', '')
                form.pba_sessions[i].activity_name.data = session_data.get('activity_name', '')
                
                questions = session_data.get('assessment_questions', [])
                for j, question_data in enumerate(questions):
                    if j < len(form.pba_sessions[i].assessment_questions):
                        form.pba_sessions[i].assessment_questions[j].question.data = question_data.get('question', '')
                        form.pba_sessions[i].assessment_questions[j].correct_answer.data = question_data.get('correct_answer', '')
        
        # Populate vocabulary
        vocab_data = form_data.get('vocabulary', [])
        for i, term_data in enumerate(vocab_data):
            if i < len(form.vocabulary):
                form.vocabulary[i].term.data = term_data.get('term', '')
                form.vocabulary[i].definition.data = term_data.get('definition', '')
        
        # Populate portfolio checklist
        portfolio_data = form_data.get('portfolio_checklist', [])
        for i, item_data in enumerate(portfolio_data):
            if i < len(form.portfolio_checklist):
                form.portfolio_checklist[i].product.data = item_data.get('product', '')
                form.portfolio_checklist[i].session_number.data = item_data.get('session_number', '')
        
        print(f"🔍 Successfully loaded simplified module answer key draft data")
        return True
        
    except Exception as e:
        print(f"🔍 Error loading draft data into form: {e}")
        import traceback
        traceback.print_exc()
        return False

def load_module_answer_key2_draft_into_form(form, user_id):
    """Load the most recent module_answer_key2 draft data into the provided form"""
    try:
        # Get the most recent draft for this user
        latest_draft = FormDraft.query.filter_by(
            user_id=user_id, 
            form_type='module_answer_key2'
        ).order_by(FormDraft.updated_at.desc()).first()
        
        if not latest_draft:
            print("🔍 No module_answer_key2 draft found for user")
            return False
            
        print(f"🔍 Loading draft data from Module Answer Key 2.0 draft ID {latest_draft.id}")
        form_data = latest_draft.form_data
        
        # Populate basic fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        
        # Populate pre-test questions
        pretest_data = form_data.get('pretest_questions', [])
        for i, question_data in enumerate(pretest_data):
            if i < len(form.pretest_questions):
                form.pretest_questions[i].question_text.data = question_data.get('question_text', '')
                form.pretest_questions[i].choice_a.data = question_data.get('choice_a', '')
                form.pretest_questions[i].choice_b.data = question_data.get('choice_b', '')
                form.pretest_questions[i].choice_c.data = question_data.get('choice_c', '')
                form.pretest_questions[i].choice_d.data = question_data.get('choice_d', '')
                # Clean correct_answer to ensure it's valid
                correct_answer = question_data.get('correct_answer', '').strip().upper()
                form.pretest_questions[i].correct_answer.data = correct_answer if correct_answer in ['A', 'B', 'C', 'D'] else ''
        
        # Populate RCA sessions
        rca_data = form_data.get('rca_sessions', [])
        print(f"🔍 Loading RCA sessions data: {len(rca_data)} sessions found")
        for i, session_data in enumerate(rca_data):
            if i < len(form.rca_sessions):
                questions = session_data.get('questions', [])
                for j, question_data in enumerate(questions):
                    if j < len(form.rca_sessions[i].questions):
                        form.rca_sessions[i].questions[j].question_text.data = question_data.get('question_text', '')
                        form.rca_sessions[i].questions[j].choice_a.data = question_data.get('choice_a', '')
                        form.rca_sessions[i].questions[j].choice_b.data = question_data.get('choice_b', '')
                        form.rca_sessions[i].questions[j].choice_c.data = question_data.get('choice_c', '')
                        form.rca_sessions[i].questions[j].choice_d.data = question_data.get('choice_d', '')
                        # Clean correct_answer to ensure it's valid
                        correct_answer = question_data.get('correct_answer', '').strip().upper()
                        form.rca_sessions[i].questions[j].correct_answer.data = correct_answer if correct_answer in ['A', 'B', 'C', 'D'] else ''
        
        # Populate post-test questions
        posttest_data = form_data.get('posttest_questions', [])
        for i, question_data in enumerate(posttest_data):
            if i < len(form.posttest_questions):
                form.posttest_questions[i].question_text.data = question_data.get('question_text', '')
                form.posttest_questions[i].choice_a.data = question_data.get('choice_a', '')
                form.posttest_questions[i].choice_b.data = question_data.get('choice_b', '')
                form.posttest_questions[i].choice_c.data = question_data.get('choice_c', '')
                form.posttest_questions[i].choice_d.data = question_data.get('choice_d', '')
                # Clean correct_answer to ensure it's valid
                correct_answer = question_data.get('correct_answer', '').strip().upper()
                form.posttest_questions[i].correct_answer.data = correct_answer if correct_answer in ['A', 'B', 'C', 'D'] else ''
        
        # Populate PBA sessions
        pba_data = form_data.get('pba_sessions', [])
        print(f"🔍 Loading PBA sessions data: {len(pba_data)} sessions found")
        for i, session_data in enumerate(pba_data):
            if i < len(form.pba_sessions):
                form.pba_sessions[i].session_number.data = session_data.get('session_number', '')
                form.pba_sessions[i].activity_name.data = session_data.get('activity_name', '')
                
                questions = session_data.get('assessment_questions', [])
                for j, question_data in enumerate(questions):
                    if j < len(form.pba_sessions[i].assessment_questions):
                        form.pba_sessions[i].assessment_questions[j].question.data = question_data.get('question', '')
                        form.pba_sessions[i].assessment_questions[j].correct_answer.data = question_data.get('correct_answer', '')
        
        # Populate vocabulary
        vocab_data = form_data.get('vocabulary', [])
        for i, term_data in enumerate(vocab_data):
            if i < len(form.vocabulary):
                form.vocabulary[i].term.data = term_data.get('term', '')
                form.vocabulary[i].definition.data = term_data.get('definition', '')
        
        # Populate portfolio checklist
        portfolio_data = form_data.get('portfolio_checklist', [])
        for i, item_data in enumerate(portfolio_data):
            if i < len(form.portfolio_checklist):
                form.portfolio_checklist[i].product.data = item_data.get('product', '')
                form.portfolio_checklist[i].session_number.data = item_data.get('session_number', '')
        
        print(f"🔍 Successfully loaded Module Answer Key 2.0 draft data")
        return True
        
    except Exception as e:
        print(f"🔍 Error loading Module Answer Key 2.0 draft data into form: {e}")
        import traceback
        traceback.print_exc()
        return False

def create_correlation_table(selected_modules, grade_level, subject):
    """Generate a formatted correlation table subdoc with specific formatting requirements"""
    from docx import Document
    from docx.shared import Inches, Pt
    
    # Create a new document to hold just the table
    subdoc = Document()
    
    # Get all standards for the selected grade and subject
    standards = Standard.query.filter_by(
        grade_level=grade_level,
        subject=subject
    ).order_by(Standard.code).limit(10).all()  # Limit to 10 for testing
    
    if not standards:
        # Add a simple paragraph if no standards found
        p = subdoc.add_paragraph("No standards found for this grade level and subject.")
        return subdoc
    
    # Get module objects
    modules = Module.query.filter(
        Module.id.in_(selected_modules),
        Module.subject == subject
    ).all()
    
    if not modules:
        # Add a simple paragraph if no modules found
        p = subdoc.add_paragraph("No modules found.")
        return subdoc
    
    # Create a simple table: 1 column for standards + 1 column per module
    num_cols = 1 + len(modules)
    num_rows = 1 + len(standards)  # Header + standards rows
    
    table = subdoc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # Configure header row
    header_cells = table.rows[0].cells
    
    # Header: "Standards" column
    header_cells[0].text = 'Standards'
    
    # Header: Module columns
    for i, module in enumerate(modules):
        header_cells[i + 1].text = module.title[:20]  # Truncate long names
    
    # Fill data rows - simplified version
    for row_idx, standard in enumerate(standards, 1):
        row = table.rows[row_idx]
        
        # Standards column
        standards_cell = row.cells[0]
        standards_cell.text = standard.code
        
        # Module columns - simplified without complex XML formatting
        for col_idx, module in enumerate(modules, 1):
            module_cell = row.cells[col_idx]
            
            # Check if this module covers this standard
            mapping = ModuleStandardMapping.query.filter_by(
                module_id=module.id,
                standard_id=standard.id,
                grade_level=grade_level
            ).first()
            
            if mapping:
                module_cell.text = 'X'
            else:
                module_cell.text = ''
    
    return subdoc


def create_hlp_table_subdoc(doc, selected_modules):
    """Generate a formatted horizontal lesson plan table subdoc using DocxTemplate subdoc approach."""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docxtpl import Subdoc  # Add missing import like correlation reports
    from models import LessonPlanModule, LessonPlanSession, LessonPlanEnrichment
    
    def format_text_consistently(text):
        """Apply consistent formatting to text data for HLP display"""
        if not text or text == 'N/A':
            return text
        
        # Split by common separators to handle lists better
        parts = []
        for part in text.split(';'):
            # Strip whitespace and capitalize first letter of each part
            part = part.strip()
            if part:
                # Capitalize first letter while preserving other capitalization
                part = part[0].upper() + part[1:] if len(part) > 1 else part.upper()
                parts.append(part)
        
        return '; '.join(parts)
    
    print(f"DEBUG HLP: Starting subdoc creation for {len(selected_modules)} module IDs: {selected_modules}")
    
    # Create subdoc from the DocxTemplate instance
    subdoc = doc.new_subdoc()
    print(f"DEBUG HLP: Created subdoc: {type(subdoc)}")
    
    # Get the selected modules
    modules = LessonPlanModule.query.filter(LessonPlanModule.id.in_(selected_modules)).all()
    print(f"DEBUG HLP: Found {len(modules)} modules from database")
    
    if not modules:
        print("DEBUG HLP: No modules found, adding error paragraph")
        p = subdoc.add_paragraph("No modules found.")
        return subdoc
    
    # Create table structure: 1 module column + 1 column per module
    num_cols = 1 + len(modules)
    
    # Find the maximum number of sessions across all modules
    max_sessions = 0
    for module in modules:
        session_count = module.sessions.count()
        max_sessions = max(max_sessions, session_count)
    
    print(f"DEBUG HLP: Found max {max_sessions} sessions across all modules")
    
    # Row structure: HORIZONTAL LESSON PLAN header + Module/Section header + Focus + Goals + Material List + Teacher Prep + PBA rows for each session + Module/Section for Enrichments + Activities (enrichments data)
    row_labels = ['Module:\nSection', 'Focus', 'Goals', 'Material\nList', 'Teacher\nPrep', 'PBA']
    num_rows = 1 + (len(row_labels) * max_sessions) + 2  # +1 for HORIZONTAL LESSON PLAN header, +2 for final Module:Section + Activities rows
    
    print(f"DEBUG HLP: Creating HLP table with {num_rows} rows and {num_cols} cols")
    table = subdoc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Helper function to shade a cell light grey
    def shade_cell_grey(cell):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')  # Light grey
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Helper function to shade a cell with custom color
    def shade_cell_custom(cell, color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Helper function to set cell margins
    def set_cell_margins(cell, top=0.02, bottom=0.02, left=0.05, right=0.05):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        tc_mar = tcPr.find(qn('w:tcMar'))
        if tc_mar is None:
            tc_mar = OxmlElement('w:tcMar')
            tcPr.append(tc_mar)
        
        # Convert inches to twips (1 inch = 1440 twips)
        def inches_to_twips(inches):
            return str(int(inches * 1440))
        
        # Set margins
        for margin_name, inches in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            margin_elem = tc_mar.find(qn(f'w:{margin_name}'))
            if margin_elem is None:
                margin_elem = OxmlElement(f'w:{margin_name}')
                tc_mar.append(margin_elem)
            margin_elem.set(qn('w:w'), inches_to_twips(inches))
            margin_elem.set(qn('w:type'), 'dxa')
    
    # Helper function to set row height
    def set_row_height(row, height_inches):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is None:
            trHeight = OxmlElement('w:trHeight')
            trPr.append(trHeight)
        trHeight.set(qn('w:val'), str(int(height_inches * 1440)))  # Convert to twips
        trHeight.set(qn('w:hRule'), 'exact')
    
    # Helper function to set cell text with specific formatting
    def set_cell_text(cell, text, font_name, font_size, bold=False, center=True, color=None):
        cell.text = text
        for paragraph in cell.paragraphs:
            if center:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.bold = bold
                if color:
                    run.font.color.rgb = RGBColor.from_string(color)
        
        # Set vertical alignment to center
        from docx.enum.table import WD_ALIGN_VERTICAL
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Set cell margins for all cells
        set_cell_margins(cell)
    
    # Set up column widths
    table.columns[0].width = Inches(1.2)  # Label column
    for i in range(1, num_cols):
        table.columns[i].width = Inches(2.0)  # Module columns
    
    # Add HORIZONTAL LESSON PLAN header row
    header_row = table.rows[0]
    set_row_height(header_row, 0.3)  # 0.3" height
    
    # Merge all cells in the header row
    merged_cell = header_row.cells[0]
    for i in range(1, num_cols):
        merged_cell.merge(header_row.cells[i])
    
    # Set header formatting
    set_cell_text(merged_cell, "HORIZONTAL LESSON PLAN", 'Rockwell', 11, bold=True, center=True, color='FFFFFF')
    shade_cell_custom(merged_cell, '13205A')  # Dark blue background
    
    # Fill all session blocks dynamically (starting from row 1 now)
    for session_num in range(1, max_sessions + 1):
        session_start_row = 1 + ((session_num - 1) * len(row_labels))  # +1 for header row offset
        
        print(f"DEBUG HLP: Creating Session {session_num} block starting at row {session_start_row}")
        
        # Session header row
        session_header_row = table.rows[session_start_row]
        set_cell_text(session_header_row.cells[0], row_labels[0], 'Rockwell', 8, bold=True)
        shade_cell_grey(session_header_row.cells[0])
        
        for i, module in enumerate(modules, 1):
            module_header = f"{module.name}:\nSession {session_num}"
            set_cell_text(session_header_row.cells[i], module_header, 'Rockwell', 9, bold=True)
            shade_cell_grey(session_header_row.cells[i])
        
        # Session data rows (Focus, Goals, Materials, Teacher Prep, PBA)
        for row_idx, label in enumerate(row_labels[1:], 1):
            current_row = table.rows[session_start_row + row_idx]
            
            # Label column
            set_cell_text(current_row.cells[0], label, 'Rockwell', 8, bold=True)
            shade_cell_grey(current_row.cells[0])
            
            # Data columns for each module
            for module_idx, module in enumerate(modules, 1):
                # Get session data for this module and session number
                session = module.sessions.filter_by(session_number=session_num).first()
                
                if session:
                    if label == 'Focus':
                        data_text = format_text_consistently(session.focus or 'N/A')
                    elif label == 'Goals':
                        data_text = format_text_consistently(session.objectives or 'N/A')
                    elif label == 'Material\nList':
                        data_text = session.materials or 'N/A'
                    elif label == 'Teacher\nPrep':
                        data_text = session.teacher_preparations or 'N/A'
                    elif label == 'PBA':
                        data_text = session.performance_assessment_questions or 'N/A'
                    else:
                        data_text = 'N/A'
                else:
                    data_text = 'N/A'  # Module doesn't have this session
                
                set_cell_text(current_row.cells[module_idx], data_text, 'Times New Roman', 8, center=True)
    
    # Add final enrichments section: Module:Section row + Activities row
    enrichments_section_row_index = num_rows - 2  # Second to last row
    enrichments_activities_row_index = num_rows - 1  # Last row
    
    # Module:Section row for Enrichments
    enrichments_section_row = table.rows[enrichments_section_row_index]
    set_cell_text(enrichments_section_row.cells[0], 'Module:\nSection', 'Rockwell', 8, bold=True)
    shade_cell_grey(enrichments_section_row.cells[0])
    
    for module_idx, module in enumerate(modules, 1):
        module_header = f"{module.name}:\nEnrichments"
        set_cell_text(enrichments_section_row.cells[module_idx], module_header, 'Rockwell', 9, bold=True)
        shade_cell_grey(enrichments_section_row.cells[module_idx])
    
    # Activities row (actual enrichments data)
    enrichments_activities_row = table.rows[enrichments_activities_row_index]
    set_cell_text(enrichments_activities_row.cells[0], 'Activities', 'Rockwell', 8, bold=True)
    shade_cell_grey(enrichments_activities_row.cells[0])
    
    # Enrichments data for each module
    for module_idx, module in enumerate(modules, 1):
        enrichments = module.enrichments.all()
        if enrichments:
            enrichment_texts = []
            for enrichment in enrichments:
                formatted_title = format_text_consistently(enrichment.title or '')
                enrichment_text = f"{enrichment.enrichment_number}. {formatted_title}"
                if enrichment.description:
                    formatted_description = format_text_consistently(enrichment.description)
                    enrichment_text += f": {formatted_description}"
                enrichment_texts.append(enrichment_text)
            data_text = '\n\n'.join(enrichment_texts)
        else:
            data_text = 'N/A'
        
        set_cell_text(enrichments_activities_row.cells[module_idx], data_text, 'Times New Roman', 8, center=True)
    
    print(f"DEBUG HLP: Table creation completed successfully, returning subdoc")
    print(f"DEBUG HLP: Final subdoc type: {type(subdoc)}")
    return subdoc

# Custom validator for correct answer field
def validate_answer_choice(form, field):
    """Validate that the answer is A, B, C, or D"""
    if field.data and field.data.upper() not in ['A', 'B', 'C', 'D', '']:
        raise ValidationError('Answer must be A, B, C, or D')

# Form for individual vocabulary words
class VocabularyWordForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    word = StringField('Term', validators=[Optional(), Length(min=1, max=100)])

# Main vocabulary worksheet form
class VocabularyWorksheetForm(FlaskForm):
    module_acronym = StringField('Module Acronym', 
                                validators=[DataRequired(), Length(min=1, max=20)],
                                render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    # Increased to 25 term fields for more vocabulary words
    words = FieldList(FormField(VocabularyWordForm), min_entries=25, max_entries=30)
    
    submit = SubmitField('Generate Worksheet')
    
    def validate_words(self, field):
        # Check if at least one word is provided
        has_words = any(word_data['word'] for word_data in field.data if word_data.get('word'))
        if not has_words:
            raise ValidationError('Please enter at least one vocabulary term.')

# ===== PBA WORKSHEET FORMS =====

# Form for individual assessments
class PBAAssessmentForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    assessment = StringField('Assessment', validators=[Optional(), Length(min=1, max=500)])

# Main PBA worksheet form
class PBAWorksheetForm(FlaskForm):
    module_acronym = StringField('Module Acronym', 
                                validators=[DataRequired(), Length(min=1, max=20)],
                                render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    session_number = StringField('Session Number', 
                                validators=[DataRequired(), Length(min=1, max=10)],
                                render_kw={"placeholder": "e.g., 1, 2, 3"})
    
    section_header = StringField('Section Header', 
                                validators=[DataRequired(), Length(min=1, max=100)],
                                render_kw={"placeholder": "e.g., Squares and Square Roots"})
    
    # 4 assessment fields
    assessments = FieldList(FormField(PBAAssessmentForm), min_entries=4, max_entries=4)
    
    submit = SubmitField('Generate PBA Worksheet')
    
    def validate_assessments(self, field):
        # Check if at least one assessment is provided
        has_assessments = any(assessment_data['assessment'] for assessment_data in field.data if assessment_data.get('assessment'))
        if not has_assessments:
            raise ValidationError('Please enter at least one assessment.')

# ===== POST TEST WORKSHEET FORMS =====

# Form for individual multiple choice questions
class PostTestQuestionForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    question_text = StringField('Question', validators=[Optional(), Length(min=1, max=500)])
    choice_a = StringField('Choice A', validators=[Optional(), Length(min=1, max=200)])
    choice_b = StringField('Choice B', validators=[Optional(), Length(min=1, max=200)])
    choice_c = StringField('Choice C', validators=[Optional(), Length(min=1, max=200)])
    choice_d = StringField('Choice D', validators=[Optional(), Length(min=1, max=200)])

# Main Post Test worksheet form
class PostTestWorksheetForm(FlaskForm):
    module_acronym = StringField('Module Acronym', 
                                validators=[DataRequired(), Length(min=1, max=20)],
                                render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    # Multiple choice questions (let's start with 10 questions)
    questions = FieldList(FormField(PostTestQuestionForm), min_entries=10, max_entries=15)
    
    submit = SubmitField('Generate Post Test Worksheet')
    
    def validate_questions(self, field):
        # Check if at least one question is provided
        has_questions = any(question_data['question_text'] for question_data in field.data if question_data.get('question_text'))
        if not has_questions:
            raise ValidationError('Please enter at least one question.')

# ===== PRE TEST WORKSHEET FORMS =====

# Form for individual multiple choice questions
class PreTestQuestionForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    question_text = StringField('Question', validators=[Optional(), Length(min=1, max=500)])
    choice_a = StringField('Choice A', validators=[Optional(), Length(min=1, max=200)])
    choice_b = StringField('Choice B', validators=[Optional(), Length(min=1, max=200)])
    choice_c = StringField('Choice C', validators=[Optional(), Length(min=1, max=200)])
    choice_d = StringField('Choice D', validators=[Optional(), Length(min=1, max=200)])

# Main Pre Test worksheet form
class PreTestWorksheetForm(FlaskForm):
    module_acronym = StringField('Module Acronym', 
                                validators=[DataRequired(), Length(min=1, max=20)],
                                render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    # Multiple choice questions (let's start with 10 questions)
    questions = FieldList(FormField(PreTestQuestionForm), min_entries=10, max_entries=15)
    
    submit = SubmitField('Generate Pre Test Worksheet')
    
    def validate_questions(self, field):
        # Check if at least one question is provided
        has_questions = any(question_data['question_text'] for question_data in field.data if question_data.get('question_text'))
        if not has_questions:
            raise ValidationError('Please enter at least one question.')

# ===== TEST WORKSHEET FORMS (Combined Pre-Test and Post-Test) =====

# Form for individual multiple choice questions
class TestQuestionForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    question_text = StringField('Question', validators=[Optional(), Length(min=1, max=500)])
    choice_a = StringField('Choice A', validators=[Optional(), Length(min=1, max=200)])
    choice_b = StringField('Choice B', validators=[Optional(), Length(min=1, max=200)])
    choice_c = StringField('Choice C', validators=[Optional(), Length(min=1, max=200)])
    choice_d = StringField('Choice D', validators=[Optional(), Length(min=1, max=200)])

# Main Test worksheet form (handles both Pre-Test and Post-Test)
class TestWorksheetForm(FlaskForm):
    test_type = RadioField('Test Type',
                          choices=[('pre', 'Pre-Test Worksheet'), ('post', 'Post-Test Worksheet')],
                          default='pre',
                          validators=[DataRequired()])
    
    module_acronym = StringField('Module Acronym', 
                                validators=[DataRequired(), Length(min=1, max=20)],
                                render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    # Multiple choice questions (10-15 questions)
    questions = FieldList(FormField(TestQuestionForm), min_entries=10, max_entries=15)
    
    submit = SubmitField('Generate Test Worksheet')
    
    def validate_questions(self, field):
        # Check if at least one question is provided
        has_questions = any(question_data['question_text'] for question_data in field.data if question_data.get('question_text'))
        if not has_questions:
            raise ValidationError('Please enter at least one question.')

# ===== GENERIC WORKSHEET FORMS =====

# Dynamic field form - represents one field in the worksheet
class DynamicFieldForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    field_type = HiddenField()  # Store the field type
    
    # Section Header fields
    section_title = StringField('Section Title', validators=[Optional(), Length(max=200)])
    
    # Question fields (shared)
    question_number = IntegerField('Question Number', validators=[Optional()])
    question_text = TextAreaField('Question Text', validators=[Optional(), Length(max=2000)])
    
    # Multiple Choice specific fields
    choice_count = SelectField('Number of Choices', 
                              choices=[('2', '2 choices'), ('3', '3 choices'), ('4', '4 choices')],
                              default='4')
    choice_a = StringField('Choice A', validators=[Optional(), Length(max=200)])
    choice_b = StringField('Choice B', validators=[Optional(), Length(max=200)])
    choice_c = StringField('Choice C', validators=[Optional(), Length(max=200)])
    choice_d = StringField('Choice D', validators=[Optional(), Length(max=200)])
    
    # Section Instructions field
    instructions_text = TextAreaField('Instructions', validators=[Optional(), Length(max=3000)])
    
    # Paragraph Text field
    paragraph_text = TextAreaField('Paragraph Text', validators=[Optional(), Length(max=3000)])
    
    # Math Problem field
    math_expression = StringField('Math Expression', validators=[Optional(), Length(max=1000)])
    
    # Image upload fields
    image_file = FileField('Image File', validators=[
        FileAllowed(['jpg', 'jpeg', 'png', 'gif'], 'Images only!')
    ])

# Worksheet Answer Key form - represents a single worksheet container
class WorksheetAnswerKeyForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    worksheet_title = StringField('Worksheet Title', 
                                 validators=[Optional(), Length(max=200)],
                                 render_kw={"placeholder": "e.g., Chapter 3 Review, Lab Answer Key"})
    
    # Dynamic fields for this specific worksheet
    dynamic_content = FieldList(FormField(DynamicFieldForm), min_entries=0)

# Main Generic Worksheet form
class GenericWorksheetForm(FlaskForm):
    # Fixed fields that always appear first
    module_acronym = StringField('Module Acronym', 
                                validators=[DataRequired(), Length(min=1, max=20)],
                                render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    worksheet_title = StringField('Worksheet Title', 
                                 validators=[DataRequired(), Length(min=1, max=100)],
                                 render_kw={"placeholder": "e.g., Heat Transfer Lab"})
    
    # Dynamic fields
    dynamic_fields = FieldList(FormField(DynamicFieldForm), min_entries=0)
    
    submit = SubmitField('Generate Generic Worksheet')

# ===== FAMILY BRIEFING FORM =====

class FamilyBriefingForm(FlaskForm):
    # Module name
    module_name = StringField('Module Acronym', 
                            validators=[DataRequired(), Length(min=1, max=100)],
                            render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    # Intro sentence
    introsentence = TextAreaField('Finish the sentence: During this module, your child will...', 
                                validators=[Optional(), Length(max=500)],
                                render_kw={"placeholder": "Enter 1-2 sentences from the original Parent Briefing document (e.g., explore the properties of matter and how it interacts with the environment.)"})
    
    # Learning Objectives (6)
    learningobjective1 = TextAreaField('Learning Objective 1', validators=[Optional(), Length(max=300)])
    learningobjective2 = TextAreaField('Learning Objective 2', validators=[Optional(), Length(max=300)])
    learningobjective3 = TextAreaField('Learning Objective 3', validators=[Optional(), Length(max=300)])
    learningobjective4 = TextAreaField('Learning Objective 4', validators=[Optional(), Length(max=300)])
    learningobjective5 = TextAreaField('Learning Objective 5', validators=[Optional(), Length(max=300)])
    learningobjective6 = TextAreaField('Learning Objective 6', validators=[Optional(), Length(max=300)])
    
    # Session Focus (7)
    activityname1 = StringField('Focus 1', validators=[Optional(), Length(max=300)])
    activityname2 = StringField('Focus 2', validators=[Optional(), Length(max=300)])
    activityname3 = StringField('Focus 3', validators=[Optional(), Length(max=300)])
    activityname4 = StringField('Focus 4', validators=[Optional(), Length(max=300)])
    activityname5 = StringField('Focus 5', validators=[Optional(), Length(max=300)])
    activityname6 = StringField('Focus 6', validators=[Optional(), Length(max=300)])
    activityname7 = StringField('Focus 7', validators=[Optional(), Length(max=300)])
    
    # Key Terminology (21 terms)
    term1 = StringField('Term 1', validators=[Optional(), Length(max=100)])
    term2 = StringField('Term 2', validators=[Optional(), Length(max=100)])
    term3 = StringField('Term 3', validators=[Optional(), Length(max=100)])
    term4 = StringField('Term 4', validators=[Optional(), Length(max=100)])
    term5 = StringField('Term 5', validators=[Optional(), Length(max=100)])
    term6 = StringField('Term 6', validators=[Optional(), Length(max=100)])
    term7 = StringField('Term 7', validators=[Optional(), Length(max=100)])
    term8 = StringField('Term 8', validators=[Optional(), Length(max=100)])
    term9 = StringField('Term 9', validators=[Optional(), Length(max=100)])
    term10 = StringField('Term 10', validators=[Optional(), Length(max=100)])
    term11 = StringField('Term 11', validators=[Optional(), Length(max=100)])
    term12 = StringField('Term 12', validators=[Optional(), Length(max=100)])
    term13 = StringField('Term 13', validators=[Optional(), Length(max=100)])
    term14 = StringField('Term 14', validators=[Optional(), Length(max=100)])
    term15 = StringField('Term 15', validators=[Optional(), Length(max=100)])
    term16 = StringField('Term 16', validators=[Optional(), Length(max=100)])
    term17 = StringField('Term 17', validators=[Optional(), Length(max=100)])
    term18 = StringField('Term 18', validators=[Optional(), Length(max=100)])
    term19 = StringField('Term 19', validators=[Optional(), Length(max=100)])
    term20 = StringField('Term 20', validators=[Optional(), Length(max=100)])
    term21 = StringField('Term 21', validators=[Optional(), Length(max=100)])
    
    # Key Concepts (3 concepts with explanations)
    keyconcept1_name = StringField('Key Concept 1 Name', validators=[Optional(), Length(max=300)])
    keyconcept1_explanation = TextAreaField('Key Concept 1 Explanation', validators=[Optional(), Length(max=600)])
    
    keyconcept2_name = StringField('Key Concept 2 Name', validators=[Optional(), Length(max=300)])
    keyconcept2_explanation = TextAreaField('Key Concept 2 Explanation', validators=[Optional(), Length(max=600)])
    
    keyconcept3_name = StringField('Key Concept 3 Name', validators=[Optional(), Length(max=300)])
    keyconcept3_explanation = TextAreaField('Key Concept 3 Explanation', validators=[Optional(), Length(max=600)])
    keyconcept4_name = StringField('Key Concept 4 Name', validators=[Optional(), Length(max=300)])
    keyconcept4_explanation = TextAreaField('Key Concept 4 Explanation', validators=[Optional(), Length(max=600)])
    keyconcept5_name = StringField('Key Concept 5 Name', validators=[Optional(), Length(max=300)])
    keyconcept5_explanation = TextAreaField('Key Concept 5 Explanation', validators=[Optional(), Length(max=600)])
    
    submit = SubmitField('Generate Family Briefing')

# ===== STUDENT LOGBOOK FORMS =====

class StudentModuleWorkbookForm(FlaskForm):
    # Module Information
    module_name = StringField('Module Acronym', 
                            validators=[DataRequired(), Length(min=1, max=100)],
                            render_kw={"placeholder": "e.g., APHY, ENSC, BIOE"})
    
    # Session Focus (7 sessions)
    focus_s1 = TextAreaField('Session 1 Focus', validators=[Optional(), Length(max=500)])
    focus_s2 = TextAreaField('Session 2 Focus', validators=[Optional(), Length(max=500)])
    focus_s3 = TextAreaField('Session 3 Focus', validators=[Optional(), Length(max=500)])
    focus_s4 = TextAreaField('Session 4 Focus', validators=[Optional(), Length(max=500)])
    focus_s5 = TextAreaField('Session 5 Focus', validators=[Optional(), Length(max=500)])
    focus_s6 = TextAreaField('Session 6 Focus', validators=[Optional(), Length(max=500)])
    focus_s7 = TextAreaField('Session 7 Focus', validators=[Optional(), Length(max=500)])
    
    # Session 1 Goals, Vocabulary, and Assessments
    s1_goal1 = StringField('Session 1 Goal 1', validators=[Optional(), Length(max=300)])
    s1_goal2 = StringField('Session 1 Goal 2', validators=[Optional(), Length(max=300)])
    s1_goal3 = StringField('Session 1 Goal 3', validators=[Optional(), Length(max=300)])
    s1_vocab1 = StringField('Session 1 Vocabulary 1', validators=[Optional(), Length(max=100)])
    s1_vocab2 = StringField('Session 1 Vocabulary 2', validators=[Optional(), Length(max=100)])
    s1_vocab3 = StringField('Session 1 Vocabulary 3', validators=[Optional(), Length(max=100)])
    s1_vocab4 = StringField('Session 1 Vocabulary 4', validators=[Optional(), Length(max=100)])
    s1_vocab5 = StringField('Session 1 Vocabulary 5', validators=[Optional(), Length(max=100)])
    s1_assessment1 = TextAreaField('Session 1 Assessment 1', validators=[Optional(), Length(max=500)])
    s1_assessment2 = TextAreaField('Session 1 Assessment 2', validators=[Optional(), Length(max=500)])
    s1_assessment3 = TextAreaField('Session 1 Assessment 3', validators=[Optional(), Length(max=500)])
    s1_assessment4 = TextAreaField('Session 1 Assessment 4', validators=[Optional(), Length(max=500)])
    
    # Session 2 Goals, Vocabulary, and Assessments
    s2_goal1 = StringField('Session 2 Goal 1', validators=[Optional(), Length(max=300)])
    s2_goal2 = StringField('Session 2 Goal 2', validators=[Optional(), Length(max=300)])
    s2_goal3 = StringField('Session 2 Goal 3', validators=[Optional(), Length(max=300)])
    s2_vocab1 = StringField('Session 2 Vocabulary 1', validators=[Optional(), Length(max=100)])
    s2_vocab2 = StringField('Session 2 Vocabulary 2', validators=[Optional(), Length(max=100)])
    s2_vocab3 = StringField('Session 2 Vocabulary 3', validators=[Optional(), Length(max=100)])
    s2_vocab4 = StringField('Session 2 Vocabulary 4', validators=[Optional(), Length(max=100)])
    s2_vocab5 = StringField('Session 2 Vocabulary 5', validators=[Optional(), Length(max=100)])
    s2_assessment1 = TextAreaField('Session 2 Assessment 1', validators=[Optional(), Length(max=500)])
    s2_assessment2 = TextAreaField('Session 2 Assessment 2', validators=[Optional(), Length(max=500)])
    s2_assessment3 = TextAreaField('Session 2 Assessment 3', validators=[Optional(), Length(max=500)])
    s2_assessment4 = TextAreaField('Session 2 Assessment 4', validators=[Optional(), Length(max=500)])
    
    # Session 3 Goals, Vocabulary, and Assessments
    s3_goal1 = StringField('Session 3 Goal 1', validators=[Optional(), Length(max=300)])
    s3_goal2 = StringField('Session 3 Goal 2', validators=[Optional(), Length(max=300)])
    s3_goal3 = StringField('Session 3 Goal 3', validators=[Optional(), Length(max=300)])
    s3_vocab1 = StringField('Session 3 Vocabulary 1', validators=[Optional(), Length(max=100)])
    s3_vocab2 = StringField('Session 3 Vocabulary 2', validators=[Optional(), Length(max=100)])
    s3_vocab3 = StringField('Session 3 Vocabulary 3', validators=[Optional(), Length(max=100)])
    s3_vocab4 = StringField('Session 3 Vocabulary 4', validators=[Optional(), Length(max=100)])
    s3_vocab5 = StringField('Session 3 Vocabulary 5', validators=[Optional(), Length(max=100)])
    s3_assessment1 = TextAreaField('Session 3 Assessment 1', validators=[Optional(), Length(max=500)])
    s3_assessment2 = TextAreaField('Session 3 Assessment 2', validators=[Optional(), Length(max=500)])
    s3_assessment3 = TextAreaField('Session 3 Assessment 3', validators=[Optional(), Length(max=500)])
    s3_assessment4 = TextAreaField('Session 3 Assessment 4', validators=[Optional(), Length(max=500)])
    
    # Session 4 Goals, Vocabulary, and Assessments
    s4_goal1 = StringField('Session 4 Goal 1', validators=[Optional(), Length(max=300)])
    s4_goal2 = StringField('Session 4 Goal 2', validators=[Optional(), Length(max=300)])
    s4_goal3 = StringField('Session 4 Goal 3', validators=[Optional(), Length(max=300)])
    s4_vocab1 = StringField('Session 4 Vocabulary 1', validators=[Optional(), Length(max=100)])
    s4_vocab2 = StringField('Session 4 Vocabulary 2', validators=[Optional(), Length(max=100)])
    s4_vocab3 = StringField('Session 4 Vocabulary 3', validators=[Optional(), Length(max=100)])
    s4_vocab4 = StringField('Session 4 Vocabulary 4', validators=[Optional(), Length(max=100)])
    s4_vocab5 = StringField('Session 4 Vocabulary 5', validators=[Optional(), Length(max=100)])
    s4_assessment1 = TextAreaField('Session 4 Assessment 1', validators=[Optional(), Length(max=500)])
    s4_assessment2 = TextAreaField('Session 4 Assessment 2', validators=[Optional(), Length(max=500)])
    s4_assessment3 = TextAreaField('Session 4 Assessment 3', validators=[Optional(), Length(max=500)])
    s4_assessment4 = TextAreaField('Session 4 Assessment 4', validators=[Optional(), Length(max=500)])
    
    # Session 5 Goals, Vocabulary, and Assessments
    s5_goal1 = StringField('Session 5 Goal 1', validators=[Optional(), Length(max=300)])
    s5_goal2 = StringField('Session 5 Goal 2', validators=[Optional(), Length(max=300)])
    s5_goal3 = StringField('Session 5 Goal 3', validators=[Optional(), Length(max=300)])
    s5_vocab1 = StringField('Session 5 Vocabulary 1', validators=[Optional(), Length(max=100)])
    s5_vocab2 = StringField('Session 5 Vocabulary 2', validators=[Optional(), Length(max=100)])
    s5_vocab3 = StringField('Session 5 Vocabulary 3', validators=[Optional(), Length(max=100)])
    s5_vocab4 = StringField('Session 5 Vocabulary 4', validators=[Optional(), Length(max=100)])
    s5_vocab5 = StringField('Session 5 Vocabulary 5', validators=[Optional(), Length(max=100)])
    s5_assessment1 = TextAreaField('Session 5 Assessment 1', validators=[Optional(), Length(max=500)])
    s5_assessment2 = TextAreaField('Session 5 Assessment 2', validators=[Optional(), Length(max=500)])
    s5_assessment3 = TextAreaField('Session 5 Assessment 3', validators=[Optional(), Length(max=500)])
    s5_assessment4 = TextAreaField('Session 5 Assessment 4', validators=[Optional(), Length(max=500)])
    
    # Session 6 Goals, Vocabulary, and Assessments
    s6_goal1 = StringField('Session 6 Goal 1', validators=[Optional(), Length(max=300)])
    s6_goal2 = StringField('Session 6 Goal 2', validators=[Optional(), Length(max=300)])
    s6_goal3 = StringField('Session 6 Goal 3', validators=[Optional(), Length(max=300)])
    s6_vocab1 = StringField('Session 6 Vocabulary 1', validators=[Optional(), Length(max=100)])
    s6_vocab2 = StringField('Session 6 Vocabulary 2', validators=[Optional(), Length(max=100)])
    s6_vocab3 = StringField('Session 6 Vocabulary 3', validators=[Optional(), Length(max=100)])
    s6_vocab4 = StringField('Session 6 Vocabulary 4', validators=[Optional(), Length(max=100)])
    s6_vocab5 = StringField('Session 6 Vocabulary 5', validators=[Optional(), Length(max=100)])
    s6_assessment1 = TextAreaField('Session 6 Assessment 1', validators=[Optional(), Length(max=500)])
    s6_assessment2 = TextAreaField('Session 6 Assessment 2', validators=[Optional(), Length(max=500)])
    s6_assessment3 = TextAreaField('Session 6 Assessment 3', validators=[Optional(), Length(max=500)])
    s6_assessment4 = TextAreaField('Session 6 Assessment 4', validators=[Optional(), Length(max=500)])
    
    # Session 7 Goals, Vocabulary, and Assessments
    s7_goal1 = StringField('Session 7 Goal 1', validators=[Optional(), Length(max=300)])
    s7_goal2 = StringField('Session 7 Goal 2', validators=[Optional(), Length(max=300)])
    s7_goal3 = StringField('Session 7 Goal 3', validators=[Optional(), Length(max=300)])
    s7_vocab1 = StringField('Session 7 Vocabulary 1', validators=[Optional(), Length(max=100)])
    s7_vocab2 = StringField('Session 7 Vocabulary 2', validators=[Optional(), Length(max=100)])
    s7_vocab3 = StringField('Session 7 Vocabulary 3', validators=[Optional(), Length(max=100)])
    s7_vocab4 = StringField('Session 7 Vocabulary 4', validators=[Optional(), Length(max=100)])
    s7_vocab5 = StringField('Session 7 Vocabulary 5', validators=[Optional(), Length(max=100)])
    s7_assessment1 = TextAreaField('Session 7 Assessment 1', validators=[Optional(), Length(max=500)])
    s7_assessment2 = TextAreaField('Session 7 Assessment 2', validators=[Optional(), Length(max=500)])
    s7_assessment3 = TextAreaField('Session 7 Assessment 3', validators=[Optional(), Length(max=500)])
    s7_assessment4 = TextAreaField('Session 7 Assessment 4', validators=[Optional(), Length(max=500)])
    
    submit = SubmitField('Generate Student Logbook')

# ===== RCA WORKSHEET FORMS =====

# Form for individual RCA questions (Research, Challenge, Application)
class RCAQuestionForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    question_text = StringField('Question', validators=[Optional(), Length(min=1, max=500)])
    choice_a = StringField('Choice A', validators=[Optional(), Length(min=1, max=200)])
    choice_b = StringField('Choice B', validators=[Optional(), Length(min=1, max=200)])
    choice_c = StringField('Choice C', validators=[Optional(), Length(min=1, max=200)])
    choice_d = StringField('Choice D', validators=[Optional(), Length(min=1, max=200)])

# Main RCA worksheet form
class RCAWorksheetForm(FlaskForm):
    module_acronym = StringField('Module Acronym', 
                                validators=[DataRequired(), Length(min=1, max=20)],
                                render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    session_number = StringField('Session Number', 
                                validators=[DataRequired(), Length(min=1, max=10)],
                                render_kw={"placeholder": "e.g., 1, 2, 3"})
    
    # Exactly 3 questions - one for Research, Challenge, and Application
    questions = FieldList(FormField(RCAQuestionForm), min_entries=3, max_entries=3)
    
    # Image upload field
    image_file = FileField('Image File', validators=[
        FileAllowed(['jpg', 'jpeg', 'png', 'gif'], 'Images only!')
    ])
    
    submit = SubmitField('Generate RCA Worksheet')
    
    def validate_questions(self, field):
        # Check if at least one question is provided
        has_questions = any(question_data['question_text'] for question_data in field.data if question_data.get('question_text'))
        if not has_questions:
            raise ValidationError('Please enter at least one question.')

# ===== MODULE GUIDE FORMS =====

# Form for individual standards
class ModuleGuideStandardForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    standard = StringField('Standard', validators=[Optional(), Length(min=1, max=600)])

# Form for individual vocabulary terms (reuse pattern from vocabulary template)
class ModuleGuideVocabForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    term = StringField('Term', validators=[Optional(), Length(min=1, max=100)])

# Form for individual career entries
class ModuleGuideCareersForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    career = StringField('Career', validators=[Optional(), Length(min=1, max=300)])

# Session Notes forms for Module Guide
class SessionGoalForm(FlaskForm):
    class Meta:
        csrf = False
    
    goal = StringField('Goal', validators=[Optional(), Length(min=1, max=600)])

class SessionPrepForm(FlaskForm):
    class Meta:
        csrf = False
    
    prep = StringField('Teacher Preparation', validators=[Optional(), Length(min=1, max=600)])

class SessionAssessmentForm(FlaskForm):
    class Meta:
        csrf = False
    
    assessment = StringField('Assessment', validators=[Optional(), Length(min=1, max=600)])

# Additional Resources form classes
class EnrichmentActivityForm(FlaskForm):
    class Meta:
        csrf = False
    
    activity = StringField('Enrichment Activity', validators=[Optional(), Length(min=1, max=600)])

class LocallySourcedMaterialForm(FlaskForm):
    class Meta:
        csrf = False
    
    material = StringField('Locally Sourced Material', validators=[Optional(), Length(min=1, max=600)])

class MaintenanceItemForm(FlaskForm):
    class Meta:
        csrf = False
    
    item = StringField('Maintenance Item', validators=[Optional(), Length(min=1, max=600)])

class AssemblyInstructionForm(FlaskForm):
    class Meta:
        csrf = False
    
    instruction = StringField('Assembly Instruction', validators=[Optional(), Length(min=1, max=600)])

class RecommendedWebsiteForm(FlaskForm):
    class Meta:
        csrf = False
    
    title = StringField('Website Title', validators=[Optional(), Length(min=1, max=100)])
    url = StringField('Website URL', validators=[Optional(), Length(min=1, max=300)])

class SessionNoteForm(FlaskForm):
    class Meta:
        csrf = False
    
    # Focus field
    focus = StringField('Session Focus', validators=[Optional(), Length(min=1, max=200)])
    
    # Goals - dynamic list (6-8 goals)
    goals = FieldList(FormField(SessionGoalForm), min_entries=6, max_entries=8)
    
    # Materials - fixed fields (15 materials like vocabulary)
    material1 = StringField('Material 1', validators=[Optional(), Length(max=100)])
    material2 = StringField('Material 2', validators=[Optional(), Length(max=100)])
    material3 = StringField('Material 3', validators=[Optional(), Length(max=100)])
    material4 = StringField('Material 4', validators=[Optional(), Length(max=100)])
    material5 = StringField('Material 5', validators=[Optional(), Length(max=100)])
    material6 = StringField('Material 6', validators=[Optional(), Length(max=100)])
    material7 = StringField('Material 7', validators=[Optional(), Length(max=100)])
    material8 = StringField('Material 8', validators=[Optional(), Length(max=100)])
    material9 = StringField('Material 9', validators=[Optional(), Length(max=100)])
    material10 = StringField('Material 10', validators=[Optional(), Length(max=100)])
    material11 = StringField('Material 11', validators=[Optional(), Length(max=100)])
    material12 = StringField('Material 12', validators=[Optional(), Length(max=100)])
    material13 = StringField('Material 13', validators=[Optional(), Length(max=100)])
    material14 = StringField('Material 14', validators=[Optional(), Length(max=100)])
    material15 = StringField('Material 15', validators=[Optional(), Length(max=100)])
    
    # Teacher Preparations - dynamic list (4-6 items)
    preparations = FieldList(FormField(SessionPrepForm), min_entries=4, max_entries=6)
    
    # Performance Based Assessments - dynamic list (4-6 items)
    assessments = FieldList(FormField(SessionAssessmentForm), min_entries=4, max_entries=6)

# Main Module Guide form (Section 1 - Teacher Tips)
class ModuleGuideForm(FlaskForm):
    module_acronym = StringField('Module Acronym', 
                               validators=[DataRequired(), Length(min=1, max=20)],
                               render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    teachertips_statement = TextAreaField('Teacher Tips Overview', 
                                        validators=[Optional(), Length(max=1000)],
                                        render_kw={"placeholder": "Enter the overview from the Teacher Tips document (e.g., In the Environmental Math Module, students will...)"})
    
    # Standards fields (starting with 6)
    standards = FieldList(FormField(ModuleGuideStandardForm), min_entries=6, max_entries=10)
    
    # Vocabulary fields (starting with 25)
    vocab_terms = FieldList(FormField(ModuleGuideVocabForm), min_entries=25, max_entries=30)
    
    # Careers fields (starting with 14)
    careers = FieldList(FormField(ModuleGuideCareersForm), min_entries=14, max_entries=20)
    
    # Session Notes - exactly 7 sessions
    sessions = FieldList(FormField(SessionNoteForm), min_entries=7, max_entries=7)
    
    # Additional Resources fields
    enrichment_activities = FieldList(FormField(EnrichmentActivityForm), min_entries=4, max_entries=10)
    locally_sourced_materials = FieldList(FormField(LocallySourcedMaterialForm), min_entries=4, max_entries=10)
    maintenance_items = FieldList(FormField(MaintenanceItemForm), min_entries=4, max_entries=10)
    assembly_instructions = FieldList(FormField(AssemblyInstructionForm), min_entries=4, max_entries=10)
    recommended_websites = FieldList(FormField(RecommendedWebsiteForm), min_entries=3, max_entries=10)
    
    submit = SubmitField('Generate Module Guide')
    
    def validate_standards(self, field):
        # Check if at least one standard is provided
        has_standards = any(standard_data['standard'] for standard_data in field.data if standard_data.get('standard'))
        if not has_standards:
            raise ValidationError('Please enter at least one standard.')

# ===== MODULE ANSWER KEY FORMS =====

# Form for individual multiple choice questions (used in pretest, posttest, and RCA)
class AnswerKeyQuestionForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    question_text = StringField('Question', validators=[Optional(), Length(min=1, max=600)])
    choice_a = StringField('Choice A', validators=[Optional(), Length(min=1, max=200)])
    choice_b = StringField('Choice B', validators=[Optional(), Length(min=1, max=200)])
    choice_c = StringField('Choice C', validators=[Optional(), Length(min=1, max=200)])
    choice_d = StringField('Choice D', validators=[Optional(), Length(min=1, max=200)])
    # Changed from SelectField to StringField to avoid nested form issues
    correct_answer = StringField('Correct Answer', 
                                validators=[Optional(), Length(max=1), validate_answer_choice],
                                render_kw={"placeholder": "A, B, C, or D", "maxlength": "1", "style": "text-transform: uppercase;"})

# Form for RCA sessions (sessions 2-5, each with 3 questions)
class RCASessionForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    # Each session has exactly 3 questions: Research, Challenge, Application
    questions = FieldList(FormField(AnswerKeyQuestionForm), min_entries=3, max_entries=3)

# Form for individual PBA assessment questions
class PBAAssessmentQuestionForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    question = StringField('Assessment Question', validators=[Optional(), Length(min=1, max=600)])
    correct_answer = TextAreaField('Correct Answer', validators=[Optional(), Length(max=800)])

# Form for PBA sessions
class PBASessionForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    session_number = IntegerField('Session Number', validators=[Optional()])
    activity_name = StringField('Activity Name', validators=[Optional(), Length(min=1, max=200)])
    assessment_questions = FieldList(FormField(PBAAssessmentQuestionForm), min_entries=4, max_entries=4)

# Form for vocabulary terms with separate term and definition fields
class VocabularyTermForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    term = StringField('Term', validators=[Optional(), Length(min=1, max=100)])
    definition = TextAreaField('Definition', validators=[Optional(), Length(max=300)])

# Form for portfolio checklist items
class PortfolioChecklistItemForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    product = StringField('Product', validators=[Optional(), Length(min=1, max=300)])
    session_number = StringField('Session Number', validators=[Optional(), Length(min=1, max=10)])

# Main Module Answer Key form
class ModuleAnswerKeyForm(FlaskForm):
    module_acronym = StringField('Module Acronym', 
                               validators=[DataRequired(), Length(min=1, max=20)],
                               render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    # Assessments Section
    # Pre-test questions (10 questions)
    pretest_questions = FieldList(FormField(AnswerKeyQuestionForm), min_entries=10, max_entries=10)
    
    # RCA sessions (sessions 2-5, each with 3 questions)
    rca_sessions = FieldList(FormField(RCASessionForm), min_entries=4, max_entries=4)
    
    # Post-test questions (10 questions)
    posttest_questions = FieldList(FormField(AnswerKeyQuestionForm), min_entries=10, max_entries=10)
    
    # Performance Based Assessments - Changed from 4 to 3 sessions
    pba_sessions = FieldList(FormField(PBASessionForm), min_entries=3, max_entries=3)
    
    # Vocabulary (at least 25 terms)
    vocabulary = FieldList(FormField(VocabularyTermForm), min_entries=25, max_entries=30)
    
    # Student Portfolio Checklist (at least 6 items)
    portfolio_checklist = FieldList(FormField(PortfolioChecklistItemForm), min_entries=6, max_entries=10)
    
    submit = SubmitField('Generate Module Answer Key')
    
    def validate_pretest_questions(self, field):
        # Check if at least one question is provided
        has_questions = any(q['question_text'] for q in field.data if q.get('question_text'))
        if not has_questions:
            raise ValidationError('Please enter at least one pre-test question.')
    
    def validate_posttest_questions(self, field):
        # Check if at least one question is provided
        has_questions = any(q['question_text'] for q in field.data if q.get('question_text'))
        if not has_questions:
            raise ValidationError('Please enter at least one post-test question.')
    
    def validate_vocabulary(self, field):
        # Check if at least one vocabulary term is provided
        has_terms = any(v['term'] for v in field.data if v.get('term'))
        if not has_terms:
            raise ValidationError('Please enter at least one vocabulary term.')

# ===== MODULE ANSWER KEY 2.0 FORM =====

# Main Module Answer Key 2.0 form - Streamlined version without Enrichment and Worksheet sections
class ModuleAnswerKey2Form(FlaskForm):
    module_acronym = StringField('Module Acronym', 
                               validators=[DataRequired(), Length(min=1, max=20)],
                               render_kw={"placeholder": "e.g., APHY, CHEM, BIOE", "data-autosave": "true"})
    
    # Assessments Section
    # Pre-test questions (10 questions)
    pretest_questions = FieldList(FormField(AnswerKeyQuestionForm), min_entries=10, max_entries=10)
    
    # RCA sessions (sessions 2-5, each with 3 questions)
    rca_sessions = FieldList(FormField(RCASessionForm), min_entries=4, max_entries=4)
    
    # Post-test questions (10 questions)
    posttest_questions = FieldList(FormField(AnswerKeyQuestionForm), min_entries=10, max_entries=10)
    
    # Performance Based Assessments (3 sessions, numbered 1-3)
    pba_sessions = FieldList(FormField(PBASessionForm), min_entries=3, max_entries=3)
    
    # Vocabulary (25-30 terms)
    vocabulary = FieldList(FormField(VocabularyTermForm), min_entries=25, max_entries=30)
    
    # Student Portfolio Checklist (6-10 items)
    portfolio_checklist = FieldList(FormField(PortfolioChecklistItemForm), min_entries=6, max_entries=10)
    
    submit = SubmitField('Generate Module Answer Key 2.0')
    
    def validate_pretest_questions(self, field):
        # Check if at least one question is provided
        has_questions = any(q['question_text'] for q in field.data if q.get('question_text'))
        if not has_questions:
            raise ValidationError('Please enter at least one pre-test question.')
    
    def validate_posttest_questions(self, field):
        # Check if at least one question is provided
        has_questions = any(q['question_text'] for q in field.data if q.get('question_text'))
        if not has_questions:
            raise ValidationError('Please enter at least one post-test question.')
    
    def validate_vocabulary(self, field):
        # Check if at least one vocabulary term is provided
        has_terms = any(v['term'] for v in field.data if v.get('term'))
        if not has_terms:
            raise ValidationError('Please enter at least one vocabulary term.')

# ===== MODULE ACTIVITY SHEET FORM =====

class ModuleActivitySheetForm(FlaskForm):
    module_acronym = StringField('Module Acronym', 
                               validators=[DataRequired(), Length(min=1, max=20)],
                               render_kw={"placeholder": "e.g., APHY, CHEM, BIOE"})
    
    # Session 1 - Always Pre-Test
    session1_activity = StringField('Session 1 Activity Name', 
                                  validators=[DataRequired(), Length(min=1, max=100)],
                                  render_kw={"placeholder": "e.g., Smart Cart and Track"})
    session1_is_pba = BooleanField('Session 1 includes Performance-Based Assessment')
    
    # Session 2 - Always RCA
    session2_activity = StringField('Session 2 Activity Name', 
                                  validators=[DataRequired(), Length(min=1, max=100)],
                                  render_kw={"placeholder": "e.g., Oscilloscope"})
    session2_is_pba = BooleanField('Session 2 includes Performance-Based Assessment')
    
    # Session 3 - Always RCA
    session3_activity = StringField('Session 3 Activity Name', 
                                  validators=[DataRequired(), Length(min=1, max=100)],
                                  render_kw={"placeholder": "e.g., Heat Expansion"})
    session3_is_pba = BooleanField('Session 3 includes Performance-Based Assessment')
    
    # Session 4 - Always RCA
    session4_activity = StringField('Session 4 Activity Name', 
                                  validators=[DataRequired(), Length(min=1, max=100)],
                                  render_kw={"placeholder": "e.g., Heat Experiment"})
    session4_is_pba = BooleanField('Session 4 includes Performance-Based Assessment')
    
    # Session 5 - Always RCA
    session5_activity = StringField('Session 5 Activity Name', 
                                  validators=[DataRequired(), Length(min=1, max=100)],
                                  render_kw={"placeholder": "e.g., Light Filter Experiment"})
    session5_is_pba = BooleanField('Session 5 includes Performance-Based Assessment')
    
    # Session 6 - Always Test Review
    session6_activity = StringField('Session 6 Activity Name', 
                                  validators=[DataRequired(), Length(min=1, max=100)],
                                  render_kw={"placeholder": "e.g., Laser Safety and Use"})
    session6_is_pba = BooleanField('Session 6 includes Performance-Based Assessment')
    
    # Session 7 - Always Post Test
    session7_activity = StringField('Session 7 Activity Name', 
                                  validators=[DataRequired(), Length(min=1, max=100)],
                                  render_kw={"placeholder": "e.g., Laser Experiments"})
    session7_is_pba = BooleanField('Session 7 includes Performance-Based Assessment')
    
    submit = SubmitField('Generate Module Activity Sheet')

# ===== HORIZONTAL LESSON PLAN FORMS =====

class HorizontalLessonPlanSessionForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    focus = TextAreaField('Session Focus', validators=[Optional(), Length(max=1000)])
    objectives = TextAreaField('Learning Objectives', validators=[Optional(), Length(max=2500)])
    materials = TextAreaField('Materials', validators=[Optional(), Length(max=2500)])
    teacher_prep = TextAreaField('Teacher Preparations', validators=[Optional(), Length(max=2500)])
    assessments = TextAreaField('Performance Based Assessments', validators=[Optional(), Length(max=2500)])

class HorizontalLessonPlanModuleForm(FlaskForm):
    class Meta:
        csrf = False  # Disable CSRF for nested forms
    
    module_name = StringField('Module Name', validators=[Optional(), Length(max=200)])
    
    # PDF upload for this module
    session_notes_pdf = FileField('Session Notes PDF', validators=[
        FileAllowed(['pdf'], 'PDF files only!')
    ])
    
    # 7 sessions for this module
    sessions = FieldList(FormField(HorizontalLessonPlanSessionForm), min_entries=7, max_entries=7)
    
    enrichment_activities = TextAreaField('Enrichment Activities', validators=[Optional(), Length(max=3000)])

class HorizontalLessonPlanForm(FlaskForm):
    # Basic information
    school_name = StringField('School Name', 
                             validators=[DataRequired(), Length(min=1, max=200)],
                             render_kw={"placeholder": "e.g., Jefferson Elementary School"})
    
    teacher_name = StringField('Teacher Name', 
                              validators=[DataRequired(), Length(min=1, max=100)],
                              render_kw={"placeholder": "e.g., Ms. Johnson"})
    
    term = StringField('Term', 
                      validators=[DataRequired(), Length(min=1, max=50)],
                      render_kw={"placeholder": "e.g., Fall 2024, Spring 2025"})
    
    # Up to 5 modules
    modules = FieldList(FormField(HorizontalLessonPlanModuleForm), min_entries=5, max_entries=5)
    
    submit = SubmitField('Generate Horizontal Lesson Plan')
    
    def validate(self, extra_validators=None):
        """Custom validation to ensure at least one module has data"""
        if not super().validate(extra_validators=extra_validators):
            return False
            
        # Check if this is an extraction request - if so, allow validation to pass if PDFs are uploaded
        action = request.form.get('action', 'generate') if request else 'generate'
        
        if action == 'extract':
            # For extraction, check if at least one PDF is uploaded
            has_pdf = False
            for module in self.modules:
                if module.session_notes_pdf.data and module.session_notes_pdf.data.filename:
                    has_pdf = True
                    break
            
            if not has_pdf:
                self.modules[0].session_notes_pdf.errors.append('Please upload at least one PDF file to extract data from.')
                return False
            
            return True
        
        # For other actions (generate/save_draft), check if at least one module has data
        has_module_data = False
        for module in self.modules:
            if module.module_name.data or any(session.focus.data or session.objectives.data 
                                            for session in module.sessions):
                has_module_data = True
                break
        
        if not has_module_data:
            self.modules[0].module_name.errors.append('Please provide data for at least one module.')
            return False
            
        return True

# Form classes for Curriculum Design Build Template
class ScienceStandardCoverageForm(FlaskForm):
    """Form for Science standard coverage percentages"""
    class Meta:
        csrf = False
    
    physical_sciences = StringField('Physical Sciences (%)', validators=[Optional()], 
                                   render_kw={"placeholder": "e.g., 100", "data-autosave": "true"})
    life_sciences = StringField('Life Sciences (%)', validators=[Optional()], 
                               render_kw={"placeholder": "e.g., 100", "data-autosave": "true"})
    earth_space_sciences = StringField('Earth & Space Sciences (%)', validators=[Optional()], 
                                     render_kw={"placeholder": "e.g., 100", "data-autosave": "true"})
    etas = StringField('ETAS (%)', validators=[Optional()], 
                      render_kw={"placeholder": "e.g., 100", "data-autosave": "true"})

class MathStandardCoverageForm(FlaskForm):
    """Form for Math standard coverage percentages"""
    class Meta:
        csrf = False
    
    rp = StringField('RP (%)', validators=[Optional()], 
                    render_kw={"placeholder": "e.g., 100", "data-autosave": "true"})
    ns = StringField('NS (%)', validators=[Optional()], 
                    render_kw={"placeholder": "e.g., 100", "data-autosave": "true"})
    ee = StringField('EE (%)', validators=[Optional()], 
                    render_kw={"placeholder": "e.g., 100", "data-autosave": "true"})
    f = StringField('F (%)', validators=[Optional()], 
                   render_kw={"placeholder": "e.g., 100", "data-autosave": "true"})
    g = StringField('G (%)', validators=[Optional()], 
                   render_kw={"placeholder": "e.g., 100", "data-autosave": "true"})
    sp = StringField('SP (%)', validators=[Optional()], 
                    render_kw={"placeholder": "e.g., 100", "data-autosave": "true"})

class CurriculumDesignBuildForm(FlaskForm):
    # Star Academy Design Details
    star_academy_model = StringField('Star Academy Model', 
                                   validators=[DataRequired(), Length(min=1, max=200)],
                                   render_kw={"placeholder": "e.g., Non-Accelerated Full Year, Accelerated Fall/Spring, etc.", "data-autosave": "true"})
    
    grade_level = StringField('Grade Level(s)', 
                            validators=[DataRequired(), Length(min=1, max=100)],
                            render_kw={"placeholder": "e.g., 7th Grade, 8th Grade, etc.", "data-autosave": "true"})
    
    soft_start = BooleanField('Soft/Mid-year Start?', 
                            render_kw={"data-autosave": "true"})
    
    site_name = StringField('Site Name', 
                          validators=[DataRequired(), Length(min=1, max=200)],
                          render_kw={"placeholder": "e.g., Union High School, TH Harris Middle School, etc.", "data-autosave": "true"})
    
    school_district = StringField('School District', 
                                validators=[DataRequired(), Length(min=1, max=200)],
                                render_kw={"placeholder": "e.g., Union Parish School District", "data-autosave": "true"})
    
    # Essential Star Academy Program Elements
    science_course_grades = StringField('Science Course Grade Level(s)', 
                                      validators=[DataRequired(), Length(min=1, max=100)],
                                      render_kw={"placeholder": "e.g., 7th/8th", "data-autosave": "true"})
    
    math_course_grades = StringField('Math Course Grade Level(s)', 
                                   validators=[DataRequired(), Length(min=1, max=100)],
                                   render_kw={"placeholder": "e.g., 7th/8th", "data-autosave": "true"})
    
    ss_course_grades = StringField('Social Studies Course Grade Level(s)', 
                                 validators=[DataRequired(), Length(min=1, max=100)],
                                 render_kw={"placeholder": "e.g., 7th/8th", "data-autosave": "true"})
    
    ela_course_grades = StringField('ELA Course Grade Level(s)', 
                                  validators=[DataRequired(), Length(min=1, max=100)],
                                  render_kw={"placeholder": "e.g., 7th/8th", "data-autosave": "true"})
    
    # Curriculum Elements
    science_rotations = SelectField('Science Rotations', 
                                  choices=[('Five', 'Five'), ('Ten', 'Ten')],
                                  validators=[DataRequired()],
                                  render_kw={"data-autosave": "true"})
    
    math_rotations = SelectField('Math Rotations', 
                               choices=[('Five', 'Five'), ('Ten', 'Ten')],
                               validators=[DataRequired()],
                               render_kw={"data-autosave": "true"})
    
    tier_one_component = StringField('Tier I Component (ELA)', 
                                   validators=[DataRequired(), Length(min=1, max=200)],
                                   render_kw={"placeholder": "e.g., Guidebooks", "data-autosave": "true"})
    
    # Curriculum Design: Science
    science_design_domain = StringField('Blended Science Domain', 
                                      validators=[DataRequired(), Length(min=1, max=100)],
                                      render_kw={"placeholder": "e.g., Physical, Life, Earth & Space, etc.", "data-autosave": "true"})
    
    
    # Dynamic Science Table - replaces hardcoded module structure
    science_table_title = StringField('Science Table Title (optional)', 
                                     validators=[Optional(), Length(max=200)],
                                     render_kw={"placeholder": "e.g., Science Curriculum Overview", "data-autosave": "true"})
    
    # Note: Table configuration and cell data will be handled via JavaScript like in generic worksheet
    
    
    
    # Curriculum Design: Math
    state_math_domains = StringField('# of state-specific math domains + State', 
                                   validators=[DataRequired(), Length(min=1, max=100)],
                                   render_kw={"placeholder": "e.g., five Louisiana", "data-autosave": "true"})
    
    ipls_additional_coverage = StringField('IPLs Additional Coverage Grade Level(s)', 
                                         validators=[Optional(), Length(max=100)],
                                         render_kw={"placeholder": "e.g., 5th, 6th, etc.", "data-autosave": "true"})
    
    ipls_critical_standards = StringField('IPLs Critical Standards Grade Level(s)', 
                                        validators=[Optional(), Length(max=100)],
                                        render_kw={"placeholder": "e.g., 5th, 6th, etc.", "data-autosave": "true"})
    
    
    # Dynamic Math Table - replaces hardcoded module structure
    math_table_title = StringField('Math Table Title (optional)', 
                                  validators=[Optional(), Length(max=200)],
                                  render_kw={"placeholder": "e.g., Math Curriculum Overview", "data-autosave": "true"})
    
    # Note: Table configuration and cell data will be handled via JavaScript like in generic worksheet
    
    
    
    
    # Curriculum Design: Social Studies
    tci_program_title = StringField('TCI Program Title', 
                                  validators=[DataRequired(), Length(min=1, max=200)],
                                  render_kw={"placeholder": "e.g., History Alive! US Through Modern Times", "data-autosave": "true"})
    
    ss_course_title = StringField('Social Studies Course Title', 
                                validators=[DataRequired(), Length(min=1, max=200)],
                                render_kw={"placeholder": "e.g., US History: Industrial Age through the Modern Era", "data-autosave": "true"})
    
    
    # Dynamic Social Studies Table - replaces hardcoded structure
    social_studies_table_title = StringField('Social Studies Table Title (optional)', 
                                            validators=[Optional(), Length(max=200)],
                                            render_kw={"placeholder": "e.g., Social Studies Curriculum Overview", "data-autosave": "true"})
    
    # Note: Table configuration and cell data will be handled via JavaScript like in generic worksheet
    
    # Social Studies standard coverage percentages (up to 2 grade levels)
    
    submit = SubmitField('Generate Curriculum Design Build')

class CorrelationReportForm(FlaskForm):
    state = SelectField('State', 
                       choices=[('', 'Select State...')],
                       validators=[DataRequired()])
    
    grade = SelectField('Grade Level', 
                       choices=[
                           ('', 'Select Grade...'),
                           ('7th Grade', '7th Grade'),
                           ('8th Grade', '8th Grade')
                       ],
                       validators=[DataRequired()])
    
    subject = SelectField('Subject', 
                         choices=[
                             ('', 'Select Subject...'),
                             ('Math', 'Math'),
                             ('Science', 'Science')
                         ],
                         validators=[DataRequired()])
    
    selected_modules = SelectMultipleField('Selected Modules',
                                         choices=[],
                                         validators=[DataRequired()])
    
    submit = SubmitField('Generate Correlation Report')

class StreamlinedHorizontalLessonPlanForm(FlaskForm):
    """Streamlined horizontal lesson plan form - database driven"""
    
    # Basic Information (simplified)
    school_name = StringField('School Name',
                             validators=[DataRequired(), Length(min=1, max=200)],
                             render_kw={"placeholder": "e.g., Jefferson Elementary School"})

    teacher_name = StringField('Teacher Name',
                              validators=[DataRequired(), Length(min=1, max=100)],
                              render_kw={"placeholder": "e.g., Ms. Johnson"})

    # School year auto-populated to 2025-2026 (user can edit if needed)
    school_year = StringField('School Year',
                             validators=[DataRequired(), Length(min=1, max=20)],
                             default='2025-2026')

    # Subject field for user input
    subject = StringField('Subject',
                         validators=[DataRequired(), Length(min=1, max=50)],
                         render_kw={"placeholder": "e.g., Math, Science, etc."})

    # Module selection (similar to correlation report)
    selected_modules = SelectMultipleField('Selected Modules (up to 10)',
                                          validators=[DataRequired()],
                                          coerce=int)

    submit = SubmitField('Generate Horizontal Lesson Plan')

    def validate_selected_modules(self, field):
        """Ensure max 10 modules selected"""
        if len(field.data) > 10:
            raise ValidationError('Please select no more than 10 modules.')

class IplReportForm(FlaskForm):
    """IPL (Individual Pacing List) Report Form"""
    
    module_acronym = StringField('Module Acronym', 
                                validators=[DataRequired(), Length(min=1, max=20)],
                                render_kw={"placeholder": "e.g., ASTR, BIOE, CHEM"})
    
    selected_modules = SelectMultipleField('Selected Modules',
                                          validators=[DataRequired()],
                                          coerce=int)
    
    submit = SubmitField('Generate Module IPL List')

@app.route('/dashboard')
@login_required
def dashboard():
    try:
        # Get user's recent drafts (this will be empty for new users, which is fine)
        recent_drafts = FormDraft.query.filter_by(
            user_id=current_user.id,
            is_current=True
        ).order_by(FormDraft.updated_at.desc()).limit(10).all()
        
        # Get user's recent documents (this will be empty for new users, which is fine)
        recent_docs = GeneratedDocument.query.filter_by(
            user_id=current_user.id
        ).order_by(GeneratedDocument.created_at.desc()).limit(10).all()
        
        # Stats
        total_drafts = FormDraft.query.filter_by(user_id=current_user.id, is_current=True).count()
        total_docs = GeneratedDocument.query.filter_by(user_id=current_user.id).count()
        
        # Calculate total downloads across all user documents
        user_docs = GeneratedDocument.query.filter_by(user_id=current_user.id).all()
        total_downloads = sum(doc.download_count for doc in user_docs)
        
        return render_template('dashboard.html', 
                             recent_drafts=recent_drafts,
                             recent_docs=recent_docs,
                             total_drafts=total_drafts,
                             total_docs=total_docs,
                             total_downloads=total_downloads)
                             
    except Exception as e:
        # If there's any database error, show a simple error message
        flash(f'Error loading dashboard: {str(e)}', 'error')
        return render_template('simple_dashboard.html')

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/autosave-vocabulary-draft', methods=['POST'])
@login_required
def autosave_vocabulary_draft():
    """AJAX endpoint for autosaving vocabulary draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'module_acronym': data.get('module_acronym', ''),
            'words': []
        }
        
        # Process vocabulary words
        words_data = data.get('words', [])
        for word_data in words_data:
            word_text = word_data.get('word', '').strip()
            if word_text:  # Only include non-empty words
                form_data['words'].append({'word': word_text})
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='vocabulary').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                # Update title if module acronym changed
                if form_data['module_acronym']:
                    draft.title = f"Vocabulary Worksheet - {form_data['module_acronym']}"
                    draft.module_acronym = form_data['module_acronym']
            else:
                return jsonify({'success': False, 'error': 'Draft not found'})
        else:
            # Create new draft
            title = f"Vocabulary Worksheet - {form_data['module_acronym']}" if form_data['module_acronym'] else "Vocabulary Worksheet - Untitled"
            draft = FormDraft(
                user_id=current_user.id,
                form_type='vocabulary',
                title=title,
                module_acronym=form_data['module_acronym'],
                form_data=form_data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'draft_id': draft.id,
            'message': 'Draft saved automatically',
            'timestamp': datetime.utcnow().strftime('%I:%M:%S %p')
        })
        
    except Exception as e:
        print(f"Error in autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/create-vocabulary', methods=['GET', 'POST'])
@login_required
def create_vocabulary():
    form = VocabularyWorksheetForm()
    
    if request.method == 'POST':
        print("Form submitted!")
        print(f"Form data: {request.form}")
        print(f"Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"Form errors: {form.errors}")
        
        # Check which action was requested
        action = request.form.get('action', 'generate')
        print(f"Action requested: {action}")
        
        if form.validate_on_submit():
            if action == 'save_draft':
                # Handle draft saving
                try:
                    # Prepare form data for JSON storage
                    form_data = {
                        'module_acronym': form.module_acronym.data,
                        'words': [{'word': word.word.data} for word in form.words if word.word.data]
                    }
                    
                    # Check if this is updating an existing draft
                    draft_id = request.form.get('draft_id')
                    if draft_id:
                        draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id).first()
                        if draft:
                            # Update existing draft
                            draft.form_data = form_data
                            draft.updated_at = datetime.utcnow()
                        else:
                            flash('Draft not found', 'error')
                            return render_template('create_vocabulary.html', form=form)
                    else:
                        # Create new draft
                        title = f"Vocabulary Worksheet - {form.module_acronym.data}" if form.module_acronym.data else "Vocabulary Worksheet - Untitled"
                        draft = FormDraft(
                            user_id=current_user.id,
                            form_type='vocabulary',
                            title=title,
                            module_acronym=form.module_acronym.data,
                            form_data=form_data
                        )
                        db.session.add(draft)
                    
                    db.session.commit()
                    flash('Draft saved successfully!', 'success')
                    return render_template('create_vocabulary.html', form=form, draft_id=draft.id)
                    
                except Exception as e:
                    print(f"Error saving draft: {e}")
                    flash(f'Error saving draft: {str(e)}', 'error')
                    return render_template('create_vocabulary.html', form=form)
            
            else:  # action == 'generate' or default
                # Handle document generation
                print("Form validation passed!")
                try:
                    print("Attempting to generate document...")
                    doc_path, filename = generate_vocabulary_worksheet(form)
                    print(f"Document generated at: {doc_path}")
                    
                    # Save document info to database
                    doc_record = GeneratedDocument(
                        user_id=current_user.id,
                        document_type='vocabulary',
                        filename=filename,
                        file_path=doc_path,
                        module_acronym=form.module_acronym.data,
                        file_size=os.path.getsize(doc_path)
                    )
                    db.session.add(doc_record)
                    db.session.commit()
                    
                    flash('Worksheet generated successfully!', 'success')
                    return redirect(url_for('my_documents'))
                except Exception as e:
                    print(f"Error generating document: {e}")
                    flash(f'Error generating worksheet: {str(e)}', 'error')
        else:
            if action == 'save_draft':
                flash('Please fix form errors before saving', 'error')
    
    return render_template('create_vocabulary.html', form=form)

@app.route('/create-pba', methods=['GET', 'POST'])
@login_required
def create_pba():
    form = PBAWorksheetForm()
    
    if request.method == 'POST':
        print("PBA form submitted!")
        print(f"Form data: {request.form}")
        print(f"Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"Form errors: {form.errors}")
    
        # Only handle document generation now - autosave handles saving
        if form.validate_on_submit():
            print("PBA form validation passed!")
            try:
                print("Attempting to generate PBA worksheet...")
                doc_path = generate_pba_worksheet(form)
                filename = os.path.basename(doc_path)
                
                print(f"PBA worksheet generated at: {doc_path}")
                
                # Save document info to database
                doc_record = GeneratedDocument(
                    user_id=current_user.id,
                    document_type='pba',
                    filename=filename,
                    file_path=doc_path,
                    module_acronym=form.module_acronym.data,
                    file_size=os.path.getsize(doc_path)
                )
                db.session.add(doc_record)
                db.session.commit()
                
                flash('PBA worksheet generated successfully!', 'success')
                return redirect(url_for('my_documents'))
            except Exception as e:
                print(f"Error generating PBA worksheet: {e}")
                flash(f'Error generating PBA worksheet: {str(e)}', 'error')
    
    return render_template('create_pba.html', form=form)

@app.route('/create-test', methods=['GET', 'POST'])
@login_required
def create_test():
    form = TestWorksheetForm()
    
    if request.method == 'POST':
        print("Test form submitted!")
        print(f"Form data: {request.form}")
        print(f"Test type selected: {form.test_type.data}")
        print(f"Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"Form errors: {form.errors}")
        
        # Only handle document generation now - autosave handles saving
        if form.validate_on_submit():
            print("Test form validation passed!")
            try:
                test_type = form.test_type.data
                if test_type == 'pre':
                    print("Attempting to generate pre-test worksheet...")
                    doc_path = generate_pretest_worksheet(form)
                    filename = os.path.basename(doc_path)
                    document_type = 'pretest'
                    flash('Pre-Test worksheet generated successfully!', 'success')
                else:  # test_type == 'post'
                    print("Attempting to generate post-test worksheet...")
                    doc_path = generate_posttest_worksheet(form)
                    filename = os.path.basename(doc_path)
                    document_type = 'posttest'
                    flash('Post-Test worksheet generated successfully!', 'success')
                
                print(f"Test worksheet generated at: {doc_path}")
                
                # Save document info to database
                doc_record = GeneratedDocument(
                    user_id=current_user.id,
                    document_type=document_type,
                    filename=filename,
                    file_path=doc_path,
                    module_acronym=form.module_acronym.data,
                    file_size=os.path.getsize(doc_path)
                )
                db.session.add(doc_record)
                db.session.commit()
                
                return redirect(url_for('my_documents'))
            except Exception as e:
                print(f"Error generating test worksheet: {e}")
                flash(f'Error generating test worksheet: {str(e)}', 'error')
    
    return render_template('create_test.html', form=form)

@app.route('/create-generic', methods=['GET', 'POST'])
@login_required
def create_generic():
    form = GenericWorksheetForm()
    
    if request.method == 'POST':
        print("Generic form submitted!")
        print(f"Form data: {request.form}")
        print(f"Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"Form errors: {form.errors}")
    
    if form.validate_on_submit():
        print("Generic form validation passed!")
        # Generate the document
        try:
            print("Attempting to generate generic worksheet...")
            doc_path, filename = generate_generic_worksheet(form)
            print(f"Generic worksheet generated at: {doc_path}")
            
            # Save document info to database
            doc_record = GeneratedDocument(
                user_id=current_user.id,
                document_type='generic',
                filename=filename,
                file_path=doc_path,
                module_acronym=form.module_acronym.data,
                file_size=os.path.getsize(doc_path)
            )
            db.session.add(doc_record)
            db.session.commit()
            
            flash('Generic worksheet generated successfully!', 'success')
            return redirect(url_for('my_documents'))
        except Exception as e:
            print(f"Error generating generic worksheet: {e}")
            flash(f'Error generating generic worksheet: {str(e)}', 'error')
    
    return render_template('create_generic.html', form=form)

@app.route('/create-familybriefing', methods=['GET', 'POST'])
@login_required
def create_familybriefing():
    form = FamilyBriefingForm()
    
    if request.method == 'POST':
        print("Family Briefing form submitted!")
        print(f"Form data: {request.form}")
        print(f"Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"Form errors: {form.errors}")
    
    if form.validate_on_submit():
        print("Family Briefing form validation passed!")
        # Generate the document
        try:
            print("Attempting to generate family briefing...")
            doc_path = generate_family_briefing(form)
            filename = os.path.basename(doc_path)
            
            print(f"Family briefing generated at: {doc_path}")
            
            # Save document info to database
            doc_record = GeneratedDocument(
                user_id=current_user.id,
                document_type='familybriefing',
                filename=filename,
                file_path=doc_path,
                module_acronym=form.module_name.data,  # Note: Family Briefing uses module_name, not module_acronym
                file_size=os.path.getsize(doc_path)
            )
            db.session.add(doc_record)
            db.session.commit()
            
            flash('Family Briefing generated successfully!', 'success')
            return redirect(url_for('my_documents'))
        except Exception as e:
            print(f"Error generating family briefing: {e}")
            flash(f'Error generating family briefing: {str(e)}', 'error')
    
    return render_template('create_familybriefing.html', form=form)

@app.route('/create-rca', methods=['GET', 'POST'])
@login_required
def create_rca():
    form = RCAWorksheetForm()
    
    if request.method == 'POST':
        print("RCA form submitted!")
        print(f"Form data: {request.form}")
        print(f"Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"Form errors: {form.errors}")
    
    if form.validate_on_submit():
        print("RCA form validation passed!")
        # Generate the document
        try:
            print("Attempting to generate RCA worksheet...")
            doc_path = generate_rca_worksheet(form)
            filename = os.path.basename(doc_path)
            
            print(f"RCA worksheet generated at: {doc_path}")
            
            # Save document info to database
            doc_record = GeneratedDocument(
                user_id=current_user.id,
                document_type='rca',
                filename=filename,
                file_path=doc_path,
                module_acronym=form.module_acronym.data,
                file_size=os.path.getsize(doc_path)
            )
            db.session.add(doc_record)
            db.session.commit()
            
            flash('RCA worksheet generated successfully!', 'success')
            return redirect(url_for('my_documents'))
        except Exception as e:
            print(f"Error generating RCA worksheet: {e}")
            flash(f'Error generating RCA worksheet: {str(e)}', 'error')
    
    return render_template('create_rca.html', form=form)

@app.route('/create-moduleGuide', methods=['GET', 'POST'])
@login_required
def create_module_guide():
    form = ModuleGuideForm()
    
    if request.method == 'POST':
        print("Module Guide form submitted!")
        print(f"Form data: {request.form}")
        print(f"Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"Form errors: {form.errors}")
    
    if form.validate_on_submit():
        print("Module Guide form validation passed!")
        # Generate the document
        try:
            print("Attempting to generate Module Guide...")
            doc_path = generate_module_guide(form)
            filename = os.path.basename(doc_path)
            
            print(f"Module Guide generated at: {doc_path}")
            
            # Save document info to database
            doc_record = GeneratedDocument(
                user_id=current_user.id,
                document_type='moduleguide',
                filename=filename,
                file_path=doc_path,
                module_acronym=form.module_acronym.data,
                file_size=os.path.getsize(doc_path)
            )
            db.session.add(doc_record)
            db.session.commit()
            
            flash('Module Guide generated successfully!', 'success')
            return redirect(url_for('my_documents'))
        except Exception as e:
            print(f"Error generating Module Guide: {e}")
            flash(f'Error generating Module Guide: {str(e)}', 'error')
    
    return render_template('create_moduleGuide.html', form=form)

@app.route('/create-moduleAnswerKey', methods=['GET', 'POST'])
@login_required
def create_module_answer_key():
    form = ModuleAnswerKeyForm()
    
    if request.method == 'POST':
        print("🔍 Module Answer Key form submitted!")
        print(f"🔍 Request method: {request.method}")
        print(f"🔍 Form data keys: {list(request.form.keys())}")
        print(f"🔍 Action value: {request.form.get('action', 'NOT FOUND')}")
        print(f"🔍 Module acronym: {request.form.get('module_acronym', 'NOT FOUND')}")
        print(f"🔍 Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"🔍 Form errors: {form.errors}")
        
        # Check if this is actually a generation request
        action = request.form.get('action')
        if action != 'generate':
            print(f"🔍 Skipping generation - action is '{action}', not 'generate'")
            return render_template('create_moduleAnswerKey.html', form=form)
        
        # Only handle document generation now - autosave handles saving
        if form.validate_on_submit():
            print("🔍 Module Answer Key form validation passed!")
            
            # CRITICAL FIX: Load autosaved draft data directly for generation
            print("🔍 Loading autosaved draft data for generation...")
            latest_draft = FormDraft.query.filter_by(
                user_id=current_user.id, 
                form_type='moduleanswerkey'
            ).order_by(FormDraft.updated_at.desc()).first()
            
            if latest_draft and latest_draft.form_data:
                print("🔍 Found draft data - merging with form data for generation")
                # Create a merged form that combines submitted data with autosaved data
                merged_form = ModuleAnswerKeyForm()
                
                # Use basic fields from submitted form (like module_acronym)
                merged_form.module_acronym.data = form.module_acronym.data
                
                # Load all data from the draft to ensure completeness
                draft_loaded = load_moduleanswerkey_draft_into_form(merged_form, current_user.id)
                if draft_loaded:
                    print(f"🔍 Successfully loaded draft - using merged form (worksheet sections removed for simplification)")
                    form = merged_form  # Use the merged form for generation
                else:
                    print("🔍 Failed to load draft - using submitted form")
            else:
                print("🔍 No draft data found - using submitted form")
            
            try:
                print("🔍 Attempting to generate Module Answer Key...")
                print(f"🔍 Form simplified - worksheet sections moved to separate template")
                doc_path = generate_module_answer_key(form)
                filename = os.path.basename(doc_path)
                
                print(f"🔍 Module Answer Key generated at: {doc_path}")
                
                # Save document info to database
                doc_record = GeneratedDocument(
                    user_id=current_user.id,
                    document_type='moduleanswerkey',
                    filename=filename,
                    file_path=doc_path,
                    module_acronym=form.module_acronym.data,
                    file_size=os.path.getsize(doc_path)
                )
                db.session.add(doc_record)
                db.session.commit()
                
                print(f"🔍 Database record created successfully")
                flash('Module Answer Key generated successfully!', 'success')
                return redirect(url_for('my_documents'))
            except Exception as e:
                print(f"🔍 Error generating Module Answer Key: {e}")
                import traceback
                traceback.print_exc()  # Print full traceback for debugging
                flash(f'Error generating Module Answer Key: {str(e)}', 'error')
                return render_template('create_moduleAnswerKey.html', form=form)
        else:
            print(f"🔍 Form validation failed!")
            print(f"🔍 Validation errors: {form.errors}")
    
    return render_template('create_moduleAnswerKey.html', form=form)

# ===== MODULE ANSWER KEY 2.0 ROUTES =====

@app.route('/create-module-answer-key2', methods=['GET', 'POST'])
@login_required
def create_module_answer_key2():
    form = ModuleAnswerKey2Form()
    
    if request.method == 'POST':
        print("🔍 Module Answer Key 2.0 form submitted!")
        print(f"🔍 Request method: {request.method}")
        print(f"🔍 Form data keys: {list(request.form.keys())}")
        print(f"🔍 Action value: {request.form.get('action', 'NOT FOUND')}")
        print(f"🔍 Module acronym: {request.form.get('module_acronym', 'NOT FOUND')}")
        print(f"🔍 Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"🔍 Form errors: {form.errors}")
        
        # Check if this is actually a generation request
        action = request.form.get('action')
        if action != 'generate':
            print(f"🔍 Skipping generation - action is '{action}', not 'generate'")
            return render_template('create_module_answer_key2.html', form=form)
        
        # Only handle document generation - autosave handles saving
        if form.validate_on_submit():
            print("🔍 Module Answer Key 2.0 form validation passed!")
            
            # Load autosaved draft data for generation reliability
            print("🔍 Loading autosaved draft data for generation...")
            latest_draft = FormDraft.query.filter_by(
                user_id=current_user.id, 
                form_type='module_answer_key2'
            ).order_by(FormDraft.updated_at.desc()).first()
            
            if latest_draft and latest_draft.form_data:
                print("🔍 Found draft data - merging with form data for generation")
                # Create a merged form that combines submitted data with autosaved data
                merged_form = ModuleAnswerKey2Form()
                
                # Use basic fields from submitted form (like module_acronym)
                merged_form.module_acronym.data = form.module_acronym.data
                
                # Load all data from the draft to ensure completeness
                draft_loaded = load_module_answer_key2_draft_into_form(merged_form, current_user.id)
                if draft_loaded:
                    print(f"🔍 Successfully loaded draft - using merged form for generation")
                    form = merged_form  # Use the merged form for generation
                else:
                    print("🔍 Failed to load draft - using submitted form")
            else:
                print("🔍 No draft data found - using submitted form")
            
            try:
                print("🔍 Attempting to generate Module Answer Key 2.0...")
                doc_path = generate_module_answer_key2(form)
                filename = os.path.basename(doc_path)
                
                print(f"🔍 Module Answer Key 2.0 generated at: {doc_path}")
                
                # Save document info to database
                doc_record = GeneratedDocument(
                    user_id=current_user.id,
                    document_type='module_answer_key2',
                    filename=filename,
                    file_path=doc_path,
                    module_acronym=form.module_acronym.data,
                    file_size=os.path.getsize(doc_path)
                )
                db.session.add(doc_record)
                db.session.commit()
                
                print(f"🔍 Database record created successfully")
                flash('Module Answer Key 2.0 generated successfully!', 'success')
                return redirect(url_for('my_documents'))
            except Exception as e:
                print(f"🔍 Error generating Module Answer Key 2.0: {e}")
                import traceback
                traceback.print_exc()  # Print full traceback for debugging
                flash(f'Error generating Module Answer Key 2.0: {str(e)}', 'error')
                return render_template('create_module_answer_key2.html', form=form)
        else:
            print(f"🔍 Form validation failed!")
            print(f"🔍 Validation errors: {form.errors}")
    
    return render_template('create_module_answer_key2.html', form=form)

@app.route('/create-moduleActivitySheet', methods=['GET', 'POST'])
@login_required
def create_module_activity_sheet():
    form = ModuleActivitySheetForm()
    
    if request.method == 'POST':
        print("Module Activity Sheet form submitted!")
        print(f"Form data: {request.form}")
        print(f"Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"Form errors: {form.errors}")
    
    if form.validate_on_submit():
        print("Module Activity Sheet form validation passed!")
        # Generate the document
        try:
            print("Attempting to generate Module Activity Sheet...")
            doc_path = generate_module_activity_sheet(form)
            filename = os.path.basename(doc_path)
            
            print(f"Module Activity Sheet generated at: {doc_path}")
            
            # Save document info to database
            doc_record = GeneratedDocument(
                user_id=current_user.id,
                document_type='moduleactivity',
                filename=filename,
                file_path=doc_path,
                module_acronym=form.module_acronym.data,
                file_size=os.path.getsize(doc_path)
            )
            db.session.add(doc_record)
            db.session.commit()
            
            flash('Module Activity Sheet generated successfully!', 'success')
            return redirect(url_for('my_documents'))
        except Exception as e:
            print(f"Error generating Module Activity Sheet: {e}")
            flash(f'Error generating Module Activity Sheet: {str(e)}', 'error')
    
    return render_template('create_moduleActivitySheet.html', form=form)

@app.route('/create-horizontal-lesson-plan', methods=['GET', 'POST'])
@login_required
def create_horizontal_lesson_plan():
    form = HorizontalLessonPlanForm()
    
    if request.method == 'POST':
        print("Horizontal Lesson Plan form submitted!")
        print(f"Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"Form errors: {form.errors}")
    
    if form.validate_on_submit():
        print("Horizontal Lesson Plan form validation passed!")
        
        # Check which action was clicked
        action = request.form.get('action', 'generate')
        print(f"Action clicked: {action}")
        
        if action == 'save_draft':
            # Handle draft saving
            try:
                # Prepare form data for JSON storage
                form_data = {
                    'school_name': form.school_name.data,
                    'teacher_name': form.teacher_name.data,
                    'term': form.term.data,
                    'modules': []
                }
                
                # Process modules
                for module in form.modules:
                    if module.module_name.data or any(session.focus.data or session.objectives.data 
                                                    for session in module.sessions):
                        module_data = {
                            'module_name': module.module_name.data,
                            'enrichment_activities': module.enrichment_activities.data,
                            'sessions': []
                        }
                        
                        # Process sessions
                        for session in module.sessions:
                            session_data = {
                                'focus': session.focus.data,
                                'objectives': session.objectives.data,
                                'materials': session.materials.data,
                                'teacher_prep': session.teacher_prep.data,
                                'assessments': session.assessments.data
                            }
                            module_data['sessions'].append(session_data)
                        
                        form_data['modules'].append(module_data)
                
                # Check if this is updating an existing draft
                draft_id = request.form.get('draft_id')
                if draft_id:
                    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id).first()
                    if draft:
                        # Update existing draft
                        draft.form_data = form_data
                        draft.updated_at = datetime.utcnow()
                    else:
                        flash('Draft not found', 'error')
                        return render_template('create_horizontal_lesson_plan.html', form=form)
                else:
                    # Create new draft
                    title = f"Horizontal Lesson Plan - {form.school_name.data}" if form.school_name.data else "Horizontal Lesson Plan - Untitled"
                    draft = FormDraft(
                        user_id=current_user.id,
                        form_type='horizontal_lesson_plan',
                        title=title,
                        form_data=form_data
                    )
                    db.session.add(draft)
                
                db.session.commit()
                flash('Draft saved successfully!', 'success')
                return render_template('create_horizontal_lesson_plan.html', form=form, draft_id=draft.id)
                
            except Exception as e:
                print(f"Error saving draft: {e}")
                flash(f'Error saving draft: {str(e)}', 'error')
                return render_template('create_horizontal_lesson_plan.html', form=form)
        
        elif action == 'extract':
            # Extract data from uploaded PDFs
            try:
                print("Extracting data from PDFs...")
                extracted_count = 0
                
                for i, module in enumerate(form.modules):
                    if module.session_notes_pdf.data and module.session_notes_pdf.data.filename:
                        print(f"Processing PDF for module {i+1}...")
                        try:
                            # Extract text from PDF
                            pdf_text = extract_pdf_text(module.session_notes_pdf.data)
                            print(f"Extracted text length: {len(pdf_text)}")
                            
                            # Parse session data using improved extraction
                            extracted_data = extract_session_data_from_text(pdf_text)
                            
                            print(f"DEBUG: Extracted data for module {i+1}:")
                            print(f"  Module name: {extracted_data.get('module_name', 'None')}")
                            print(f"  Enrichment: {extracted_data.get('enrichment', 'None')}")
                            print(f"  Sessions found: {len(extracted_data.get('sessions', []))}")
                            
                            # Populate form fields with extracted data
                            if extracted_data.get('module_name'):
                                module.module_name.data = extracted_data['module_name']
                                print(f"  Set module name: {extracted_data['module_name']}")
                            
                            if extracted_data.get('enrichment'):
                                module.enrichment_activities.data = extracted_data['enrichment']
                                print(f"  Set enrichment: {extracted_data['enrichment'][:50]}...")
                            
                            # Populate session data
                            sessions_data = extracted_data.get('sessions', [])
                            for j, session_data in enumerate(sessions_data[:7]):  # Max 7 sessions
                                if j < len(module.sessions):
                                    print(f"    Session {j+1}: focus='{session_data.get('focus', '')[:30]}...', objectives='{session_data.get('objectives', '')[:30]}...'")
                                    module.sessions[j].focus.data = session_data.get('focus', '')
                                    module.sessions[j].objectives.data = session_data.get('objectives', '')
                                    module.sessions[j].materials.data = session_data.get('materials', '')
                                    module.sessions[j].teacher_prep.data = session_data.get('teacher_prep', '')
                                    module.sessions[j].assessments.data = session_data.get('assessments', '')
                            
                            extracted_count += 1
                            
                        except Exception as e:
                            print(f"Error processing PDF for module {i+1}: {e}")
                            flash(f'Error processing PDF for module {i+1}: {str(e)}', 'warning')
                
                if extracted_count > 0:
                    flash(f'Data extracted from {extracted_count} PDF(s) successfully! Review and edit the data below.', 'success')
                else:
                    flash('No valid PDFs found to extract data from.', 'warning')
                
            except Exception as e:
                print(f"Error during extraction: {e}")
                flash(f'Error extracting data: {str(e)}', 'error')
                
            # Return to form with populated data
            return render_template('create_horizontal_lesson_plan.html', form=form)
            
        else:  # action == 'generate' or default
            # Generate the document
            try:
                print("Attempting to generate Horizontal Lesson Plan...")
                doc_path, filename = generate_horizontal_lesson_plan(form)
                
                print(f"Horizontal Lesson Plan generated at: {doc_path}")
                
                # Save document info to database
                doc_record = GeneratedDocument(
                    user_id=current_user.id,
                    document_type='horizontal_lesson_plan',
                    filename=filename,
                    file_path=doc_path,
                    module_acronym=None,  # This form doesn't have a single module acronym
                    file_size=os.path.getsize(doc_path)
                )
                db.session.add(doc_record)
                db.session.commit()
                
                flash('Horizontal Lesson Plan generated successfully!', 'success')
                return redirect(url_for('my_documents'))
            except Exception as e:
                print(f"Error generating Horizontal Lesson Plan: {e}")
                flash(f'Error generating Horizontal Lesson Plan: {str(e)}', 'error')
    
    return render_template('create_horizontal_lesson_plan.html', form=form)

@app.route('/load-moduleactivity-draft/<int:draft_id>')
@login_required
def load_moduleactivity_draft(draft_id):
    """Load module activity sheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='moduleactivity').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_module_activity_sheet'))
    
    try:
        # Create form and populate with draft data
        form = ModuleActivitySheetForm()
        form_data = draft.form_data
        
        # Populate form fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        
        # Populate session data
        sessions_data = form_data.get('sessions', [])
        session_dict = {session['session_number']: session for session in sessions_data}
        
        for i in range(1, 8):
            if i in session_dict:
                session_data = session_dict[i]
                getattr(form, f'session{i}_activity').data = session_data.get('activity', '')
                getattr(form, f'session{i}_is_pba').data = session_data.get('is_pba', False)
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_moduleActivitySheet.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading module activity draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_module_activity_sheet'))

@app.route('/delete-moduleactivity-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_moduleactivity_draft(draft_id):
    """Delete module activity sheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='moduleactivity').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

def generate_vocabulary_worksheet(form):
    """Generate a vocabulary worksheet using docxtpl"""
    # Use a master template that never gets touched
    master_template_path = 'templates/docx_templates/vocabulary_worksheet_master.docx'
    working_template_path = 'templates/docx_templates/vocabulary_worksheet.docx'
    
    print(f"Looking for vocabulary master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Vocabulary master DOCX template not found. Please create the master template first.")
    
    # Always copy from master to working template before processing
    print("Copying fresh vocabulary template from master...")
    shutil.copy2(master_template_path, working_template_path)
    
    print("Loading vocabulary working template...")
    
    # Create a temporary copy of the working template
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(working_template_path, temp_file.name)
        temp_template_path = temp_file.name
    
    try:
        # Load the temporary template
        doc = DocxTemplate(temp_template_path)
        
        # Prepare context data with XML escaping
        words_data = [
            {'word': escape_xml(word_data['word'])}
            for word_data in form.words.data
            if word_data.get('word')  # Only include non-empty words
        ]
        
        context = {
            'date': datetime.now().strftime('%B %d, %Y'),
            'words': words_data
        }
        
        print(f"Vocabulary context data: {context}")
        print(f"Number of words: {len(words_data)}")
        
        # Render the document
        print("Rendering vocabulary document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"Vocabulary {escape_xml(form.module_acronym.data).replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving vocabulary document to: {output_path}")
        doc.save(output_path)
        
        print("Vocabulary document saved successfully!")
        return output_path, filename
        
    finally:
        # Clean up the temporary file
        try:
            os.unlink(temp_template_path)
        except:
            pass  # Ignore cleanup errors

def generate_pba_worksheet(form):
    """Generate a PBA worksheet using docxtpl"""
    # Use a master template that never gets touched
    master_template_path = 'templates/docx_templates/pba_worksheet_master.docx'
    working_template_path = 'templates/docx_templates/pba_worksheet.docx'
    
    print(f"Looking for PBA master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("PBA master DOCX template not found. Please create the master template first.")
    
    # Always copy from master to working template before processing
    print("Copying fresh PBA template from master...")
    shutil.copy2(master_template_path, working_template_path)
    
    print("Loading PBA working template...")
    
    # Create a temporary copy of the working template
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(working_template_path, temp_file.name)
        temp_template_path = temp_file.name
    
    try:
        # Load the temporary template
        doc = DocxTemplate(temp_template_path)
        
        # Prepare context data with XML escaping
        assessments_data = [
            {'assessment': escape_xml(assessment_data['assessment'])}
            for assessment_data in form.assessments.data
            if assessment_data.get('assessment')  # Only include non-empty assessments
        ]
        
        context = {
            'module_acronym': escape_xml(form.module_acronym.data),
            'session_number': escape_xml(form.session_number.data),
            'section_header': escape_xml(form.section_header.data),
            'assessments': assessments_data
        }
        
        print(f"PBA context data: {context}")
        print(f"Number of assessments: {len(assessments_data)}")
        
        # Render the document
        print("Rendering PBA document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"PBA (Session {escape_xml(form.session_number.data)}) {escape_xml(form.module_acronym.data).replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving PBA document to: {output_path}")
        doc.save(output_path)
        
        print("PBA document saved successfully!")
        return output_path
        
    finally:
        # Clean up the temporary file
        try:
            os.unlink(temp_template_path)
        except:
            pass  # Ignore cleanup errors

def generate_posttest_worksheet(form):
    """Generate a post test worksheet using docxtpl"""
    # Use a master template that never gets touched
    master_template_path = 'templates/docx_templates/post_test_worksheet_master.docx'
    working_template_path = 'templates/docx_templates/post_test_worksheet.docx'
    
    print(f"Looking for post test master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Post test master DOCX template not found. Please create the master template first.")
    
    # Always copy from master to working template before processing
    print("Copying fresh post test template from master...")
    shutil.copy2(master_template_path, working_template_path)
    
    print("Loading post test working template...")
    
    # Create a temporary copy of the working template
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(working_template_path, temp_file.name)
        temp_template_path = temp_file.name
    
    try:
        # Load the temporary template
        doc = DocxTemplate(temp_template_path)
        
        # Prepare context data with new structure for template and XML escaping
        questions_data = []
        for question_data in form.questions.data:
            if question_data.get('question_text'):  # Only include questions with text
                # Create choice array for template indexing with XML escaping
                questions_data.append({
                    'question': escape_xml(question_data['question_text']),
                    'choice': [
                        escape_xml(question_data['choice_a']),
                        escape_xml(question_data['choice_b']),
                        escape_xml(question_data['choice_c']),
                        escape_xml(question_data['choice_d'])
                    ]
                })
        
        context = {
            'module_acronym': escape_xml(form.module_acronym.data),
            'questions': questions_data
        }
        
        print(f"Post test context data: {context}")
        print(f"Number of questions: {len(questions_data)}")
        
        # Render the document
        print("Rendering post test document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"Post-Test WS {escape_xml(form.module_acronym.data).replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving post test document to: {output_path}")
        doc.save(output_path)
        
        print("Post test document saved successfully!")
        return output_path
        
    finally:
        # Clean up the temporary file
        try:
            os.unlink(temp_template_path)
        except:
            pass  # Ignore cleanup errors

def generate_pretest_worksheet(form):
    """Generate a pre test worksheet using docxtpl"""
    # Use a master template that never gets touched
    master_template_path = 'templates/docx_templates/pre_test_worksheet_master.docx'
    working_template_path = 'templates/docx_templates/pre_test_worksheet.docx'
    
    print(f"Looking for pre test master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Pre test master DOCX template not found. Please create the master template first.")
    
    # PROTECTION: Verify master template integrity before proceeding
    master_stat = os.stat(master_template_path)
    print(f"Master template size: {master_stat.st_size} bytes, modified: {datetime.fromtimestamp(master_stat.st_mtime)}")
    
    # Always copy from master to working template before processing
    print("Copying fresh pre test template from master...")
    try:
        shutil.copy2(master_template_path, working_template_path)
        print(f"✓ Successfully copied master to working template")
    except Exception as e:
        raise Exception(f"Failed to copy master template: {e}")
    
    print("Loading pre test working template...")
    
    # Create a temporary copy of the working template - NEVER touch master directly
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        try:
            shutil.copy2(working_template_path, temp_file.name)
            temp_template_path = temp_file.name
            print(f"✓ Created temporary template at: {temp_template_path}")
        except Exception as e:
            raise Exception(f"Failed to create temporary template: {e}")
    
    try:
        # Load the temporary template - NEVER the master
        print(f"Loading DocxTemplate from temporary file: {temp_template_path}")
        doc = DocxTemplate(temp_template_path)
        
        # Prepare context data with new structure for template and XML escaping
        questions_data = []
        for question_data in form.questions.data:
            if question_data.get('question_text'):  # Only include questions with text
                # Create choice array for template indexing with XML escaping
                questions_data.append({
                    'question': escape_xml(question_data['question_text']),
                    'choice': [
                        escape_xml(question_data['choice_a']),
                        escape_xml(question_data['choice_b']),
                        escape_xml(question_data['choice_c']),
                        escape_xml(question_data['choice_d'])
                    ]
                })
        
        context = {
            'module_acronym': escape_xml(form.module_acronym.data),
            'questions': questions_data
        }
        
        print(f"Pre test context data: {context}")
        print(f"Number of questions: {len(questions_data)}")
        
        # Render the document
        print("Rendering pre test document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"Pre-Test WS {escape_xml(form.module_acronym.data).replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving pre test document to: {output_path}")
        doc.save(output_path)
        
        print("Pre test document saved successfully!")
        
        # PROTECTION: Verify master template wasn't accidentally modified
        master_stat_after = os.stat(master_template_path)
        if master_stat.st_mtime != master_stat_after.st_mtime:
            print("⚠️  WARNING: Master template modification time changed during processing!")
        else:
            print("✓ Master template integrity verified - no accidental changes")
            
        return output_path
        
    finally:
        # Clean up the temporary file
        try:
            if os.path.exists(temp_template_path):
                os.unlink(temp_template_path)
                print(f"✓ Cleaned up temporary file: {temp_template_path}")
        except Exception as e:
            print(f"Warning: Could not clean up temporary file {temp_template_path}: {e}")

def generate_generic_worksheet(form):
    """Generate a generic worksheet using docxtpl with dynamic content"""
    # Import request for accessing raw form data (needed for table and multi-column problems)
    from flask import request
    import re
    
    def latex_to_omml(latex_text):
        """Convert LaTeX expressions to OMML (Office Math Markup Language) XML"""
        if not latex_text.strip():
            return None
        
        namespace = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        
        def create_text_run(text):
            """Create a text run element"""
            # Escape XML special characters
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return f'<m:r><m:t>{text}</m:t></m:r>'
        
        def create_fraction(numerator, denominator):
            """Create OMML for fraction"""
            return f'''<m:f>
                <m:num>{latex_to_omml_content(numerator)}</m:num>
                <m:den>{latex_to_omml_content(denominator)}</m:den>
            </m:f>'''
        
        def create_superscript(base, exponent):
            """Create OMML for superscript"""
            return f'''<m:sSup>
                <m:e>{latex_to_omml_content(base)}</m:e>
                <m:sup>{latex_to_omml_content(exponent)}</m:sup>
            </m:sSup>'''
        
        def create_subscript(base, subscript):
            """Create OMML for subscript"""
            return f'''<m:sSub>
                <m:e>{latex_to_omml_content(base)}</m:e>
                <m:sub>{latex_to_omml_content(subscript)}</m:sub>
            </m:sSub>'''
        
        def create_square_root(radicand):
            """Create OMML for square root"""
            return f'''<m:rad>
                <m:radPr><m:degHide m:val="1"/></m:radPr>
                <m:deg/>
                <m:e>{latex_to_omml_content(radicand)}</m:e>
            </m:rad>'''
        
        def create_nary(symbol, sub_content, sup_content, body_content):
            """Create OMML for n-ary operations (integrals, sums)"""
            sub_part = f'<m:sub>{latex_to_omml_content(sub_content)}</m:sub>' if sub_content else '<m:sub/>'
            sup_part = f'<m:sup>{latex_to_omml_content(sup_content)}</m:sup>' if sup_content else '<m:sup/>'
            
            return f'''<m:nary>
                <m:naryPr><m:chr m:val="{symbol}"/></m:naryPr>
                {sub_part}
                {sup_part}
                <m:e>{latex_to_omml_content(body_content)}</m:e>
            </m:nary>'''
        
        def find_matching_brace(text, start_pos):
            """Find the matching closing brace for an opening brace at start_pos"""
            if start_pos >= len(text) or text[start_pos] != '{':
                return -1
            
            brace_count = 0
            for i in range(start_pos, len(text)):
                if text[i] == '{':
                    brace_count += 1
                elif text[i] == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        return i
            return -1
        
        def latex_to_omml_content(text):
            """Convert LaTeX content to OMML, handling nested expressions properly"""
            if not text.strip():
                return create_text_run('')
            
            # Process from left to right, handling the FIRST (leftmost) LaTeX command found
            # This ensures we handle outer expressions before inner ones
            
            # Find positions of all LaTeX commands
            frac_pos = text.find('\\frac{')
            sqrt_pos = text.find('\\sqrt{')
            
            # Determine which command comes first (leftmost)
            first_cmd_pos = float('inf')
            first_cmd_type = None
            
            if frac_pos != -1:
                first_cmd_pos = frac_pos
                first_cmd_type = 'frac'
            
            if sqrt_pos != -1 and sqrt_pos < first_cmd_pos:
                first_cmd_pos = sqrt_pos
                first_cmd_type = 'sqrt'
            
            # Process the first (leftmost) command found
            if first_cmd_type == 'frac':
                # Find numerator
                num_brace_start = frac_pos + 5  # Position of first '{'
                num_brace_end = find_matching_brace(text, num_brace_start)
                
                if num_brace_end != -1:
                    # Find denominator
                    den_brace_start = num_brace_end + 1
                    if den_brace_start < len(text) and text[den_brace_start] == '{':
                        den_brace_end = find_matching_brace(text, den_brace_start)
                        
                        if den_brace_end != -1:
                            before = text[:frac_pos]
                            numerator = text[num_brace_start + 1:num_brace_end]
                            denominator = text[den_brace_start + 1:den_brace_end]
                            after = text[den_brace_end + 1:]
                            
                            frac_omml = create_fraction(numerator, denominator)
                            
                            result = ''
                            if before.strip():
                                result += latex_to_omml_content(before)
                            result += frac_omml
                            if after.strip():
                                result += latex_to_omml_content(after)
                            return result
            
            elif first_cmd_type == 'sqrt':
                brace_start = sqrt_pos + 5  # Position of '{'
                brace_end = find_matching_brace(text, brace_start)
                
                if brace_end != -1:
                    before = text[:sqrt_pos]
                    radicand = text[brace_start + 1:brace_end]  # Content between braces
                    after = text[brace_end + 1:]
                    
                    sqrt_omml = create_square_root(radicand)
                    
                    result = ''
                    if before.strip():
                        result += latex_to_omml_content(before)
                    result += sqrt_omml
                    if after.strip():
                        result += latex_to_omml_content(after)
                    return result
            
            # Handle integrals with bounds
            int_match = re.search(r'\\int_{([^{}]*(?:\{[^{}]*\}[^{}]*)*)}(?:\^{([^{}]*(?:\{[^{}]*\}[^{}]*)*)})?(?:\s+([^\\]*))?', text)
            if int_match:
                before = text[:int_match.start()]
                after = text[int_match.end():]
                sub_content = int_match.group(1)
                sup_content = int_match.group(2) if int_match.group(2) else ''
                body_content = int_match.group(3).strip() if int_match.group(3) else 'f(x)dx'
                
                int_omml = create_nary('∫', sub_content, sup_content, body_content)
                
                result = ''
                if before.strip():
                    result += latex_to_omml_content(before)
                result += int_omml
                if after.strip():
                    result += latex_to_omml_content(after)
                return result
            
            # Handle simple integrals
            if '\\int' in text:
                text = text.replace('\\int', '∫')
            
            # Handle summations with bounds
            sum_match = re.search(r'\\sum_{([^{}]*(?:\{[^{}]*\}[^{}]*)*)}(?:\^{([^{}]*(?:\{[^{}]*\}[^{}]*)*)})?(?:\s+([^\\]*))?', text)
            if sum_match:
                before = text[:sum_match.start()]
                after = text[sum_match.end():]
                sub_content = sum_match.group(1)
                sup_content = sum_match.group(2) if sum_match.group(2) else ''
                body_content = sum_match.group(3).strip() if sum_match.group(3) else 'x_i'
                
                sum_omml = create_nary('∑', sub_content, sup_content, body_content)
                
                result = ''
                if before.strip():
                    result += latex_to_omml_content(before)
                result += sum_omml
                if after.strip():
                    result += latex_to_omml_content(after)
                return result
            
            # Handle simple summations
            if '\\sum' in text:
                text = text.replace('\\sum', '∑')
            
            # Handle superscripts
            sup_match = re.search(r'([a-zA-Z0-9]+)\^{([^{}]+)}', text)
            if not sup_match:
                sup_match = re.search(r'([a-zA-Z0-9]+)\^([0-9a-zA-Z])', text)
            
            if sup_match:
                before = text[:sup_match.start()]
                after = text[sup_match.end():]
                sup_omml = create_superscript(sup_match.group(1), sup_match.group(2))
                
                result = ''
                if before.strip():
                    result += latex_to_omml_content(before)
                result += sup_omml
                if after.strip():
                    result += latex_to_omml_content(after)
                return result
            
            # Handle subscripts
            sub_match = re.search(r'([a-zA-Z0-9]+)_{([^{}]+)}', text)
            if not sub_match:
                sub_match = re.search(r'([a-zA-Z0-9]+)_([0-9a-zA-Z])', text)
            
            if sub_match:
                before = text[:sub_match.start()]
                after = text[sub_match.end():]
                sub_omml = create_subscript(sub_match.group(1), sub_match.group(2))
                
                result = ''
                if before.strip():
                    result += latex_to_omml_content(before)
                result += sub_omml
                if after.strip():
                    result += latex_to_omml_content(after)
                return result
            
            # Handle common Greek letters
            greek_map = {
                'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ', 'epsilon': 'ε',
                'theta': 'θ', 'lambda': 'λ', 'mu': 'μ', 'pi': 'π', 'sigma': 'σ', 'phi': 'φ'
            }
            for latex_name, unicode_char in greek_map.items():
                text = re.sub(f'\\\\{latex_name}\\b', unicode_char, text)
            
            # Return as text run if no special formatting found
            return create_text_run(text)
        
        # Process the main LaTeX expression
        content = latex_to_omml_content(latex_text.strip())
        
        # Wrap in oMath element
        return f'''<m:oMath xmlns:m="{namespace}">
            {content}
        </m:oMath>'''

    def process_equations_in_text(text, paragraph):
        """Process [EQUATION]...[/EQUATION] markers in text and add Word equations as OMML objects"""
        if not text or '[EQUATION]' not in text:
            return text
        
        from lxml import etree
        
        # Find all equation markers
        equation_pattern = r'\[EQUATION\](.*?)\[/EQUATION\]'
        equations = re.findall(equation_pattern, text)
        
        if not equations:
            return text
        
        # Split text by equation markers
        parts = re.split(equation_pattern, text)
        
        # Clear the paragraph and rebuild with equations
        paragraph.clear()
        
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Regular text
                if part.strip():
                    run = paragraph.add_run(part)
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(11)
            else:  # Equation part
                try:
                    # Convert LaTeX to OMML
                    omml_xml = latex_to_omml(part.strip())
                    
                    if omml_xml:
                        # Parse and insert OMML element
                        omml_element = etree.fromstring(omml_xml)
                        paragraph._element.append(omml_element)
                    else:
                        # Fallback to text if conversion fails
                        fallback_run = paragraph.add_run(f" {part} ")
                        fallback_run.font.name = 'Cambria Math'
                        fallback_run.font.size = Pt(11)
                        
                except Exception as e:
                    print(f"Warning: Error processing equation '{part}': {e}")
                    # Fallback to text if OMML processing fails
                    fallback_run = paragraph.add_run(f" {part} ")
                    fallback_run.font.name = 'Cambria Math'
                    fallback_run.font.size = Pt(11)
        
        return "processed"  # Indicate text was processed
    
    # Use the generic master template (formerly homework template)
    master_template_path = 'templates/docx_templates/generic_worksheet_master.docx'
    working_template_path = 'templates/docx_templates/generic_worksheet.docx'
    
    print(f"Looking for generic master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Generic master DOCX template not found. Please create the master template first.")
    
    # Always copy from master to working template before processing
    print("Copying fresh generic template from master...")
    shutil.copy2(master_template_path, working_template_path)
    
    print("Loading generic working template...")
    
    # Create a temporary copy of the working template
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(working_template_path, temp_file.name)
        temp_template_path = temp_file.name
    
    try:
        # Load the temporary template
        doc = DocxTemplate(temp_template_path)
        
        # Create a subdocument to hold all dynamic content
        subdoc = doc.new_subdoc()
        
        question_counter = 1
        
        for i, field_data in enumerate(form.dynamic_fields.data):
            field_type = field_data.get('field_type')
            
            print(f"Processing field {i}: type = {field_type}")
            
            if field_type == 'section_header':
                title = (field_data.get('section_title') or '').strip()
                if title:
                    # Add header paragraph with direct formatting
                    p = subdoc.add_paragraph()
                    result = process_equations_in_text(title, p)
                    if result != "processed":
                        # No equations, add text normally
                        p.add_run(title)
                    
                    # Apply header formatting: Segoe UI 14pt Bold
                    for run in p.runs:
                        run.font.name = 'Segoe UI'
                        run.font.size = Pt(14)
                        run.font.bold = True
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(6)
            
            elif field_type == 'section_instructions':
                inst_text = (field_data.get('instructions_text') or '').strip()
                if inst_text:
                    # Add instruction paragraph with direct formatting
                    p = subdoc.add_paragraph()
                    result = process_equations_in_text(inst_text, p)
                    if result != "processed":
                        p.add_run(inst_text)
                    
                    # Apply instructions formatting: Segoe UI 11pt Italic
                    for run in p.runs:
                        run.font.name = 'Segoe UI'
                        run.font.size = Pt(11)
                        run.font.italic = True
                    p.paragraph_format.space_after = Pt(6)
            
            elif field_type == 'paragraph_text':
                para_text = (field_data.get('paragraph_text') or '').strip()
                if para_text:
                    # Add body text paragraph with direct formatting
                    p = subdoc.add_paragraph()
                    result = process_equations_in_text(para_text, p)
                    if result != "processed":
                        p.add_run(para_text)
                    
                    # Apply body text formatting: Segoe UI 11pt Regular
                    for run in p.runs:
                        run.font.name = 'Segoe UI'
                        run.font.size = Pt(11)
                    p.paragraph_format.space_after = Pt(6)
            
            elif field_type == 'table':
                # For table fields, we need to access raw request data since WTForms 
                # doesn't know about our dynamic table structure
                
                # Extract table configuration from raw form data
                field_prefix = f'dynamic_fields-{i}-'
                table_title = (request.form.get(f'{field_prefix}table_title') or '').strip()
                table_rows = int(request.form.get(f'{field_prefix}table_rows', 3))
                table_cols = int(request.form.get(f'{field_prefix}table_cols', 3))
                
                print(f"DEBUG: Extracted from raw form - title='{table_title}', rows={table_rows}, cols={table_cols}")
                
                # Add table title if provided
                if table_title:
                    p = subdoc.add_paragraph(table_title)
                    for run in p.runs:
                        run.font.name = 'Segoe UI'
                        run.font.size = Pt(12)
                        run.font.bold = True
                    p.paragraph_format.space_after = Pt(6)
                
                # Create the table
                table = subdoc.add_table(rows=table_rows, cols=table_cols)
                table.style = 'Table Grid'  # Use built-in table style with borders
                
                # Populate table cells with data
                for row_idx in range(table_rows):
                    for col_idx in range(table_cols):
                        cell_key = f'{field_prefix}table_cell_{row_idx}_{col_idx}'
                        cell_value = (request.form.get(cell_key) or '').strip()
                        
                        print(f"DEBUG: Looking for raw form key '{cell_key}', found value: '{cell_value}'")
                        
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_value
                        
                        # Format cell content
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Segoe UI'
                                run.font.size = Pt(11)
                                # Make header row bold
                                if row_idx == 0:
                                    run.font.bold = True
                
                # Add spacing after table
                subdoc.add_paragraph().paragraph_format.space_after = Pt(12)
            
            elif field_type == 'multi_column_problems':
                # For multi-column problems, access raw request data
                field_prefix = f'dynamic_fields-{i}-'
                problem_rows = int(request.form.get(f'{field_prefix}problem_rows', 3))
                problem_cols = int(request.form.get(f'{field_prefix}problem_cols', 2))
                start_number = int(request.form.get(f'{field_prefix}start_number', 1))
                answer_style = request.form.get(f'{field_prefix}answer_style', 'line')
                show_answers = request.form.get(f'{field_prefix}show_answers', 'false') == 'true'
                
                print(f"DEBUG: Multi-column problems - rows={problem_rows}, cols={problem_cols}, start={start_number}, style={answer_style}, answers={show_answers}")
                
                # Create borderless table for problems
                table = subdoc.add_table(rows=problem_rows, cols=problem_cols)
                table.style = None  # Remove table style for borderless appearance
                
                # Set column widths to be equal
                for col in table.columns:
                    col.width = Inches(6.5 / problem_cols)  # Distribute across page width
                
                problem_number = start_number
                for row_idx in range(problem_rows):
                    for col_idx in range(problem_cols):
                        cell = table.cell(row_idx, col_idx)
                        
                        # Get problem and answer data from raw form
                        cell_index = row_idx * problem_cols + col_idx
                        problem_key = f'{field_prefix}problem_{cell_index}'
                        answer_key = f'{field_prefix}answer_{cell_index}'
                        
                        problem_text = (request.form.get(problem_key) or '').strip()
                        answer_text = (request.form.get(answer_key) or '').strip()
                        
                        print(f"DEBUG: Problem {problem_number}: '{problem_text}', Answer: '{answer_text}'")
                        
                        if problem_text:
                            # Add problem number and text with equation support
                            p = cell.add_paragraph()
                            full_text = f"{problem_number}. {problem_text}"
                            result = process_equations_in_text(full_text, p)
                            if result != "processed":
                                run = p.add_run(full_text)
                                run.font.name = 'Segoe UI'
                                run.font.size = Pt(11)
                            
                            # Add answer space based on style
                            if answer_style == 'line' and not show_answers:
                                answer_p = cell.add_paragraph()
                                answer_run = answer_p.add_run("   _____________")
                                answer_run.font.name = 'Segoe UI'
                                answer_run.font.size = Pt(11)
                            elif answer_style == 'box' and not show_answers:
                                answer_p = cell.add_paragraph()
                                answer_run = answer_p.add_run("   □")
                                answer_run.font.name = 'Segoe UI'
                                answer_run.font.size = Pt(14)
                            elif answer_style == 'equals' and not show_answers:
                                answer_p = cell.add_paragraph()
                                answer_run = answer_p.add_run("   = _______")
                                answer_run.font.name = 'Segoe UI'
                                answer_run.font.size = Pt(11)
                            elif show_answers and answer_text:
                                # Show the actual answer for answer keys
                                answer_p = cell.add_paragraph()
                                if answer_style == 'equals':
                                    answer_run = answer_p.add_run(f"   = {answer_text}")
                                else:
                                    answer_run = answer_p.add_run(f"   {answer_text}")
                                answer_run.font.name = 'Segoe UI'
                                answer_run.font.size = Pt(11)
                                answer_run.font.bold = True  # Bold answers in answer keys
                            
                            # Add some spacing in the cell
                            cell.add_paragraph().paragraph_format.space_after = Pt(6)
                        
                        problem_number += 1
                
                # Add spacing after the problem table
                subdoc.add_paragraph().paragraph_format.space_after = Pt(12)
            
            # Note: Image field type removed from frontend
            
            elif field_type in ['multiple_choice', 'fill_in_blank', 'text_entry', 'math_problem']:
                question_text = (field_data.get('question_text') or '').strip()
                math_expression = (field_data.get('math_expression') or '').strip()
                
                if question_text or math_expression:
                    # Add question or math problem with direct formatting
                    if field_type == 'math_problem' and math_expression:
                        p = subdoc.add_paragraph(f"{question_counter}. {math_expression}")
                    else:
                        p = subdoc.add_paragraph(f"{question_counter}. {question_text}")
                    
                    # Apply question formatting: Segoe UI 11pt Regular
                    for run in p.runs:
                        run.font.name = 'Segoe UI'
                        run.font.size = Pt(11)
                    p.paragraph_format.space_after = Pt(6)
                    
                    if field_type == 'multiple_choice':
                        # Create a 2x4 table for the choices (2 rows, 4 columns)
                        table = subdoc.add_table(rows=2, cols=4)
                        
                        # Set column widths to be reasonable
                        table.autofit = False
                        for row in table.rows:
                            # Letters get smaller width, choices get larger width
                            row.cells[0].width = Inches(0.3)  # A.
                            row.cells[1].width = Inches(3)  # Choice A text
                            row.cells[2].width = Inches(0.3)  # C.
                            row.cells[3].width = Inches(3)  # Choice C text
                        
                        # Remove table borders for a cleaner look
                        from docx.oxml.shared import qn
                        for row in table.rows:
                            for cell in row.cells:
                                # Set cell margins
                                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                        
                        # Get all choices
                        choices = [
                            ('A', field_data.get('choice_a', '')),
                            ('B', field_data.get('choice_b', '')),
                            ('C', field_data.get('choice_c', '')),
                            ('D', field_data.get('choice_d', ''))
                        ]
                        
                        # Populate the table cells in 2x4 format
                        # Row 0: A. | Choice A text | C. | Choice C text
                        # Row 1: B. | Choice B text | D. | Choice D text
                        
                        for i, (letter, choice_text) in enumerate(choices):
                            if choice_text.strip():
                                if i == 0:  # A
                                    row, letter_col, text_col = 0, 0, 1
                                elif i == 1:  # B  
                                    row, letter_col, text_col = 1, 0, 1
                                elif i == 2:  # C
                                    row, letter_col, text_col = 0, 2, 3
                                else:  # D
                                    row, letter_col, text_col = 1, 2, 3
                                
                                # Add letter (A., B., C., D.)
                                letter_cell = table.cell(row, letter_col)
                                letter_cell.text = f"{letter}."
                                for paragraph in letter_cell.paragraphs:
                                    paragraph.paragraph_format.space_before = Pt(6)
                                    for run in paragraph.runs:
                                        run.font.name = 'Segoe UI'
                                        run.font.size = Pt(11)
                                        run.font.bold = False
                                
                                # Add choice text
                                text_cell = table.cell(row, text_col)
                                text_cell.text = choice_text.strip()
                                for paragraph in text_cell.paragraphs:
                                    paragraph.paragraph_format.space_before = Pt(6)
                                    for run in paragraph.runs:
                                        run.font.name = 'Segoe UI'
                                        run.font.size = Pt(11)
                        
                        # Add spacing after the table
                        subdoc.add_paragraph().paragraph_format.space_after = Pt(6)
                    
                    elif field_type == 'fill_in_blank':
                        p = subdoc.add_paragraph("")
                        p.paragraph_format.left_indent = Pt(36)
                    
                    elif field_type == 'text_entry':
                        # Add blank lines for text entry
                        for _ in range(3):
                            p = subdoc.add_paragraph()
                            p.paragraph_format.space_after = Pt(12)
                    
                    # Add spacing after question
                    subdoc.add_paragraph()
                    question_counter += 1
        
        # Initialize context with subdoc as dynamic_content
        context = {
            'module_acronym': escape_xml(form.module_acronym.data),
            'worksheet_title': escape_xml(form.worksheet_title.data),
            'date': datetime.now().strftime('%B %d, %Y'),
            'dynamic_content': subdoc
        }
        
        print(f"Built dynamic content as subdocument")
        
        # Render the document
        print("Rendering generic document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"{escape_xml(form.worksheet_title.data).replace(' ', '_')} {escape_xml(form.module_acronym.data).replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving generic document to: {output_path}")
        doc.save(output_path)
        
        print("Generic document saved successfully!")
        return output_path, filename
        
    finally:
        # Clean up the temporary files
        try:
            os.unlink(temp_template_path)
        except:
            pass  # Ignore cleanup errors
        

def generate_family_briefing(form):
    """Generate a family briefing document using docxtpl"""
    # Use a master template that never gets touched
    master_template_path = 'templates/docx_templates/family_briefing_master.docx'
    working_template_path = 'templates/docx_templates/family_briefing.docx'
    
    print(f"Looking for family briefing master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Family briefing master DOCX template not found. Please create the master template first.")
    
    # Always copy from master to working template before processing
    print("Copying fresh family briefing template from master...")
    shutil.copy2(master_template_path, working_template_path)
    
    print("Loading family briefing working template...")
    
    # Create a temporary copy of the working template
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(working_template_path, temp_file.name)
        temp_template_path = temp_file.name
    
    try:
        # Load the temporary template
        doc = DocxTemplate(temp_template_path)
        
        # Build the context dictionary
        context = {
            'module_name': escape_xml(form.module_name.data) if form.module_name.data else '',
            'introsentence': escape_xml(form.introsentence.data) if form.introsentence.data else '',
            
            # Learning Objectives
            'learningobjective1': escape_xml(form.learningobjective1.data) if form.learningobjective1.data else '',
            'learningobjective2': escape_xml(form.learningobjective2.data) if form.learningobjective2.data else '',
            'learningobjective3': escape_xml(form.learningobjective3.data) if form.learningobjective3.data else '',
            'learningobjective4': escape_xml(form.learningobjective4.data) if form.learningobjective4.data else '',
            'learningobjective5': escape_xml(form.learningobjective5.data) if form.learningobjective5.data else '',
            'learningobjective6': escape_xml(form.learningobjective6.data) if form.learningobjective6.data else '',
            
            # Session Focus
            'activityname1': escape_xml(form.activityname1.data) if form.activityname1.data else '',
            'activityname2': escape_xml(form.activityname2.data) if form.activityname2.data else '',
            'activityname3': escape_xml(form.activityname3.data) if form.activityname3.data else '',
            'activityname4': escape_xml(form.activityname4.data) if form.activityname4.data else '',
            'activityname5': escape_xml(form.activityname5.data) if form.activityname5.data else '',
            'activityname6': escape_xml(form.activityname6.data) if form.activityname6.data else '',
            'activityname7': escape_xml(form.activityname7.data) if form.activityname7.data else '',
            
            # Key Terminology
            'term1': escape_xml(form.term1.data) if form.term1.data else '',
            'term2': escape_xml(form.term2.data) if form.term2.data else '',
            'term3': escape_xml(form.term3.data) if form.term3.data else '',
            'term4': escape_xml(form.term4.data) if form.term4.data else '',
            'term5': escape_xml(form.term5.data) if form.term5.data else '',
            'term6': escape_xml(form.term6.data) if form.term6.data else '',
            'term7': escape_xml(form.term7.data) if form.term7.data else '',
            'term8': escape_xml(form.term8.data) if form.term8.data else '',
            'term9': escape_xml(form.term9.data) if form.term9.data else '',
            'term10': escape_xml(form.term10.data) if form.term10.data else '',
            'term11': escape_xml(form.term11.data) if form.term11.data else '',
            'term12': escape_xml(form.term12.data) if form.term12.data else '',
            'term13': escape_xml(form.term13.data) if form.term13.data else '',
            'term14': escape_xml(form.term14.data) if form.term14.data else '',
            'term15': escape_xml(form.term15.data) if form.term15.data else '',
            'term16': escape_xml(form.term16.data) if form.term16.data else '',
            'term17': escape_xml(form.term17.data) if form.term17.data else '',
            'term18': escape_xml(form.term18.data) if form.term18.data else '',
            'term19': escape_xml(form.term19.data) if form.term19.data else '',
            'term20': escape_xml(form.term20.data) if form.term20.data else '',
            'term21': escape_xml(form.term21.data) if form.term21.data else '',
            
            # Key Concepts - Using simple variable names
            'keyconcept1': escape_xml(form.keyconcept1_name.data) if form.keyconcept1_name.data else '',
            'keyconcept1_explanation': escape_xml(form.keyconcept1_explanation.data) if form.keyconcept1_explanation.data else '',
            'keyconcept2': escape_xml(form.keyconcept2_name.data) if form.keyconcept2_name.data else '',
            'keyconcept2_explanation': escape_xml(form.keyconcept2_explanation.data) if form.keyconcept2_explanation.data else '',
            'keyconcept3': escape_xml(form.keyconcept3_name.data) if form.keyconcept3_name.data else '',
            'keyconcept3_explanation': escape_xml(form.keyconcept3_explanation.data) if form.keyconcept3_explanation.data else '',
            'keyconcept4': escape_xml(form.keyconcept4_name.data) if form.keyconcept4_name.data else '',
            'keyconcept4_explanation': escape_xml(form.keyconcept4_explanation.data) if form.keyconcept4_explanation.data else '',
            'keyconcept5': escape_xml(form.keyconcept5_name.data) if form.keyconcept5_name.data else '',
            'keyconcept5_explanation': escape_xml(form.keyconcept5_explanation.data) if form.keyconcept5_explanation.data else ''
        }
        
        print(f"Family briefing context data: {context}")
        
        # Render the document
        print("Rendering family briefing document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"Family Briefing {escape_xml(form.module_name.data).replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving family briefing document to: {output_path}")
        doc.save(output_path)
        
        print("Family briefing document saved successfully!")
        return output_path
        
    finally:
        # Clean up the temporary file
        try:
            os.unlink(temp_template_path)
        except:
            pass  # Ignore cleanup errors

def generate_student_module_workbook(form):
    """Generate a student logbook using docxtpl"""
    # Use a master template that never gets touched
    master_template_path = 'templates/docx_templates/student_module_workbook_master.docx'
    working_template_path = 'templates/docx_templates/student_module_workbook.docx'
    
    print(f"Looking for student logbook master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Student Logbook master DOCX template not found. Please create the master template first.")
    
    # Always copy from master to working template before processing
    print("Copying fresh student logbook template from master...")
    shutil.copy2(master_template_path, working_template_path)
    
    print("Loading student logbook working template...")
    
    # Create a temporary copy of the working template
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(working_template_path, temp_file.name)
        temp_template_path = temp_file.name
    
    try:
        # Load the temporary template
        doc = DocxTemplate(temp_template_path)
        
        # Build the context dictionary based on your template structure
        context = {
            # Module information
            'module_name': escape_xml(form.module_name.data) if form.module_name.data else '',
            
            # Session focus
            'focus': {
                's1': escape_xml(form.focus_s1.data) if form.focus_s1.data else '',
                's2': escape_xml(form.focus_s2.data) if form.focus_s2.data else '',
                's3': escape_xml(form.focus_s3.data) if form.focus_s3.data else '',
                's4': escape_xml(form.focus_s4.data) if form.focus_s4.data else '',
                's5': escape_xml(form.focus_s5.data) if form.focus_s5.data else '',
                's6': escape_xml(form.focus_s6.data) if form.focus_s6.data else '',
                's7': escape_xml(form.focus_s7.data) if form.focus_s7.data else '',
            }
        }
        
        # Add session data matching your template structure
        for s in range(1, 8):
            # Collect goals (filter out empty ones)
            goals = []
            for g in range(1, 4):
                goal_data = getattr(form, f's{s}_goal{g}').data
                if goal_data and goal_data.strip():
                    goals.append(escape_xml(goal_data))
            
            # Collect vocabulary (filter out empty ones)
            vocabulary = []
            for v in range(1, 6):
                vocab_data = getattr(form, f's{s}_vocab{v}').data
                if vocab_data and vocab_data.strip():
                    vocabulary.append(escape_xml(vocab_data))
            
            # Collect assessments for this specific session (filter out empty ones)
            assessments = []
            for a in range(1, 5):
                assessment_data = getattr(form, f's{s}_assessment{a}').data
                if assessment_data and assessment_data.strip():
                    assessments.append(escape_xml(assessment_data))
            
            # Add both naming patterns to match your template
            context[f'session{s}'] = {
                'goals': goals,
                'vocabulary': vocabulary
            }
            context[f's{s}'] = {
                'assessments': assessments  # Each session now has its own assessments
            }
        
        print(f"Student logbook context data: {context}")
        
        # Render the document
        doc.render(context)
        
        # Save the document
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate filename based on module name
        module_name_safe = 'Module'
        if form.module_name.data:
            module_name_safe = escape_xml(form.module_name.data).replace(' ', '_').replace('/', '_').replace('\\', '_')
        
        filename = f"{module_name_safe}_Student_Logbook.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving student logbook document to: {output_path}")
        doc.save(output_path)
        
        print("Student logbook document saved successfully!")
        return output_path
        
    finally:
        # Clean up the temporary file
        try:
            os.unlink(temp_template_path)
        except:
            pass  # Ignore cleanup errors

def generate_rca_worksheet(form):
    """Generate an RCA worksheet using docxtpl"""
    # Use a master template that never gets touched
    master_template_path = 'templates/docx_templates/rca_worksheet_master.docx'
    working_template_path = 'templates/docx_templates/rca_worksheet.docx'
    
    print(f"Looking for RCA master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("RCA master DOCX template not found. Please create the master template first.")
    
    # PROTECTION: Verify master template integrity before proceeding
    master_stat = os.stat(master_template_path)
    print(f"Master template size: {master_stat.st_size} bytes, modified: {datetime.fromtimestamp(master_stat.st_mtime)}")
    
    # Always copy from master to working template before processing
    print("Copying fresh RCA template from master...")
    try:
        shutil.copy2(master_template_path, working_template_path)
        print(f"✓ Successfully copied master to working template")
    except Exception as e:
        raise Exception(f"Failed to copy master template: {e}")
    
    print("Loading RCA working template...")
    
    # Create a temporary copy of the working template - NEVER touch master directly
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        try:
            shutil.copy2(working_template_path, temp_file.name)
            temp_template_path = temp_file.name
            print(f"✓ Created temporary template at: {temp_template_path}")
        except Exception as e:
            raise Exception(f"Failed to create temporary template: {e}")
    
    temp_image_path = None
    try:
        # Load the temporary template - NEVER the master
        print(f"Loading DocxTemplate from temporary file: {temp_template_path}")
        doc = DocxTemplate(temp_template_path)
        
        # Prepare context data with same structure as test templates
        questions_data = []
        
        # Always include exactly 3 questions (Research, Challenge, Application)
        # Even if some are empty, to match template expectations
        for i in range(3):
            if i < len(form.questions.data):
                question_data = form.questions.data[i]
                questions_data.append({
                    'question': escape_xml(question_data.get('question_text', '')),
                    'choice': [
                        escape_xml(question_data.get('choice_a', '')),
                        escape_xml(question_data.get('choice_b', '')),
                        escape_xml(question_data.get('choice_c', '')),
                        escape_xml(question_data.get('choice_d', ''))
                    ]
                })
            else:
                # Add empty question if form doesn't have enough
                questions_data.append({
                    'question': '',
                    'choice': ['', '', '', '']
                })
        
        # Start with basic context
        context = {
            'module_acronym': escape_xml(form.module_acronym.data),
            'questions': questions_data
        }
        
        # Handle image upload if provided
        if form.image_file.data:
            print("Processing uploaded image...")
            
            # Save the uploaded file temporarily
            image_file = form.image_file.data
            temp_dir = os.path.join('temp', 'uploads')
            os.makedirs(temp_dir, exist_ok=True)
            
            # Create a unique filename
            filename = secure_filename(image_file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            temp_filename = f"{timestamp}_{filename}"
            temp_image_path = os.path.join(temp_dir, temp_filename)
            
            # Save the file
            image_file.save(temp_image_path)
            print(f"✓ Image saved temporarily at: {temp_image_path}")
            
            # Create InlineImage for the template
            # Using width of 5 inches and maintaining aspect ratio
            image_obj = InlineImage(doc, temp_image_path, width=Inches(5))
            
            # Add image to context
            context['image'] = image_obj
        
        print(f"RCA context data: {context}")
        print(f"Number of questions: {len(questions_data)}")
        
        # Debug logging to help identify triplication issue
        for i, q in enumerate(questions_data):
            print(f"Question {i}: {q['question'][:50] if q['question'] else '[empty]'}...")
        
        # Render the document
        print("Rendering RCA document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"RCA {escape_xml(form.session_number.data)} {escape_xml(form.module_acronym.data).replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving RCA document to: {output_path}")
        doc.save(output_path)
        
        print("RCA document saved successfully!")
        
        # PROTECTION: Verify master template wasn't accidentally modified
        master_stat_after = os.stat(master_template_path)
        if master_stat.st_mtime != master_stat_after.st_mtime:
            print("⚠️  WARNING: Master template modification time changed during processing!")
        else:
            print("✓ Master template integrity verified - no accidental changes")
            
        return output_path
        
    finally:
        # Clean up the temporary files
        try:
            if os.path.exists(temp_template_path):
                os.unlink(temp_template_path)
                print(f"✓ Cleaned up temporary file: {temp_template_path}")
        except Exception as e:
            print(f"Warning: Could not clean up temporary file {temp_template_path}: {e}")
            
        # Clean up temporary image if it exists
        if temp_image_path and os.path.exists(temp_image_path):
            try:
                os.unlink(temp_image_path)
                print(f"✓ Cleaned up temporary image: {temp_image_path}")
            except Exception as e:
                print(f"Warning: Could not clean up temporary image {temp_image_path}: {e}")

def generate_module_guide(form):
    """Generate a Module Guide using docxtpl"""
    # Use a master template that never gets touched
    master_template_path = 'templates/docx_templates/module_guide_master.docx'
    working_template_path = 'templates/docx_templates/module_guide.docx'
    
    print(f"Looking for Module Guide master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Module Guide master DOCX template not found. Please create the master template first.")
    
    # PROTECTION: Verify master template integrity before proceeding
    master_stat = os.stat(master_template_path)
    print(f"Master template size: {master_stat.st_size} bytes, modified: {datetime.fromtimestamp(master_stat.st_mtime)}")
    
    # Always copy from master to working template before processing
    print("Copying fresh Module Guide template from master...")
    try:
        shutil.copy2(master_template_path, working_template_path)
        print(f"✓ Successfully copied master to working template")
    except Exception as e:
        raise Exception(f"Failed to copy master template: {e}")
    
    print("Loading Module Guide working template...")
    
    # Create a temporary copy of the working template - NEVER touch master directly
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        try:
            shutil.copy2(working_template_path, temp_file.name)
            temp_template_path = temp_file.name
            print(f"✓ Created temporary template at: {temp_template_path}")
        except Exception as e:
            raise Exception(f"Failed to create temporary template: {e}")
    
    try:
        # Load the temporary template - NEVER the master
        print(f"Loading DocxTemplate from temporary file: {temp_template_path}")
        doc = DocxTemplate(temp_template_path)
        
        # Prepare context data for Teacher Tips section
        standards_data = [
            escape_xml(standard_data['standard'])
            for standard_data in form.standards.data
            if standard_data.get('standard')  # Only include non-empty standards
        ]
        
        # Prepare vocabulary terms data
        vocab_terms_data = [
            escape_xml(vocab_data['term'])
            for vocab_data in form.vocab_terms.data
            if vocab_data.get('term')  # Only include non-empty terms
        ]
        
        # Prepare careers data
        careers_data = [
            escape_xml(career_data['career'])
            for career_data in form.careers.data
            if career_data.get('career')  # Only include non-empty careers
        ]
        
        # Build context with individual numbered vocabulary terms
        context = {
            'teachertips': {
                'overview': escape_xml(form.teachertips_statement.data) if form.teachertips_statement.data else '',
                'standard': standards_data  # Array for template indexing
            },
            # Add empty vocab and careers objects to prevent undefined errors in template
            'vocab': {},
            'careers': {}
        }
        
        # Add individual vocabulary terms (term1, term2, etc.)
        for i, term in enumerate(vocab_terms_data):
            context[f'term{i+1}'] = term
        
        # Fill remaining vocabulary terms with empty strings (up to 25 like Family Briefing)
        for i in range(len(vocab_terms_data), 25):
            context[f'term{i+1}'] = ''
            
        # Add individual careers (career1, career2, etc.)
        for i, career in enumerate(careers_data):
            context[f'career{i+1}'] = career
            
        # Fill remaining careers with empty strings (up to 14)
        for i in range(len(careers_data), 14):
            context[f'career{i+1}'] = ''
        
        # Process Session Notes data
        sessions_data = []
        for session_form in form.sessions.data:
            session_obj = {
                'focus': escape_xml(session_form.get('focus', '')),
                'goals': [],
                'prep': [],
                'assessments': []
            }
            
            # Process goals
            if 'goals' in session_form:
                for goal_data in session_form['goals']:
                    if goal_data.get('goal'):
                        session_obj['goals'].append(escape_xml(goal_data['goal']))
            
            # Add material fields (material1 through material15)
            for i in range(1, 16):
                material_key = f'material{i}'
                material_value = (session_form.get(material_key) or '').strip()
                session_obj[material_key] = escape_xml(material_value) if material_value else ''
            
            # Also add materials as a list for easier template processing
            materials_list = []
            for i in range(1, 16):
                material_key = f'material{i}'
                material_value = (session_form.get(material_key) or '').strip()
                if material_value:  # Only add non-empty materials
                    materials_list.append(escape_xml(material_value))
            session_obj['materials'] = materials_list
            
            # Process preparations
            if 'preparations' in session_form:
                for prep_data in session_form['preparations']:
                    if prep_data.get('prep'):
                        session_obj['prep'].append(escape_xml(prep_data['prep']))
            
            # Process assessments
            if 'assessments' in session_form:
                for assessment_data in session_form['assessments']:
                    if assessment_data.get('assessment'):
                        session_obj['assessments'].append(escape_xml(assessment_data['assessment']))
            
            sessions_data.append(session_obj)
        
        # Add sessions to context
        context['sessions'] = sessions_data
        
        # Process Additional Resources data
        # Enrichment Activities
        enrichment_activities_data = [
            escape_xml(activity_data['activity'])
            for activity_data in form.enrichment_activities.data
            if activity_data.get('activity')  # Only include non-empty activities
        ]
        
        # Locally Sourced Materials
        locally_sourced_materials_data = [
            escape_xml(material_data['material'])
            for material_data in form.locally_sourced_materials.data
            if material_data.get('material')  # Only include non-empty materials
        ]
        
        # Maintenance Items
        maintenance_items_data = [
            escape_xml(item_data['item'])
            for item_data in form.maintenance_items.data
            if item_data.get('item')  # Only include non-empty items
        ]
        
        # Assembly Instructions
        assembly_instructions_data = [
            escape_xml(instruction_data['instruction'])
            for instruction_data in form.assembly_instructions.data
            if instruction_data.get('instruction')  # Only include non-empty instructions
        ]
        
        # Recommended Websites
        recommended_websites_data = [
            {
                'title': website_data['title'],
                'url': website_data['url']
            }
            for website_data in form.recommended_websites.data
            if website_data.get('title') and website_data.get('url')  # Only include websites with both title and URL
        ]
        
        # Add Additional Resources to context following template variable names
        context['enrichment'] = {
            'activities': enrichment_activities_data
        }
        context['locally'] = {
            'sourced': locally_sourced_materials_data
        }
        context['maintenance'] = {
            'list': maintenance_items_data
        }
        context['assembly'] = {
            'instructions': assembly_instructions_data
        }
        context['recommended'] = {
            'websites': recommended_websites_data
        }
        
        print(f"Module Guide context data: {context}")
        print(f"Number of standards: {len(standards_data)}")
        print(f"Number of vocabulary terms: {len(vocab_terms_data)}")
        print(f"Number of careers: {len(careers_data)}")
        print(f"Number of sessions: {len(sessions_data)}")
        print(f"Number of enrichment activities: {len(enrichment_activities_data)}")
        print(f"Number of locally sourced materials: {len(locally_sourced_materials_data)}")
        print(f"Number of maintenance items: {len(maintenance_items_data)}")
        print(f"Number of assembly instructions: {len(assembly_instructions_data)}")
        print(f"Number of recommended websites: {len(recommended_websites_data)}")
        
        # Render the document
        print("Rendering Module Guide document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"Module Guide {form.module_acronym.data.replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving Module Guide document to: {output_path}")
        doc.save(output_path)
        
        print("Module Guide document saved successfully!")
        
        # PROTECTION: Verify master template wasn't accidentally modified
        master_stat_after = os.stat(master_template_path)
        if master_stat.st_mtime != master_stat_after.st_mtime:
            print("⚠️  WARNING: Master template modification time changed during processing!")
        else:
            print("✓ Master template integrity verified - no accidental changes")
            
        return output_path
        
    finally:
        # Clean up the temporary file
        try:
            if os.path.exists(temp_template_path):
                os.unlink(temp_template_path)
                print(f"✓ Cleaned up temporary file: {temp_template_path}")
        except Exception as e:
            print(f"Warning: Could not clean up temporary file {temp_template_path}: {e}")

def generate_module_answer_key(form):
    """Generate a Module Answer Key using docxtpl"""
    # Use a master template that never gets touched
    master_template_path = 'templates/docx_templates/module_ak_master.docx'
    working_template_path = 'templates/docx_templates/module_ak.docx'
    
    print(f"Looking for Module Answer Key master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Module Answer Key master DOCX template not found. Please create the master template first.")
    
    # PROTECTION: Verify master template integrity before proceeding
    master_stat = os.stat(master_template_path)
    print(f"Master template size: {master_stat.st_size} bytes, modified: {datetime.fromtimestamp(master_stat.st_mtime)}")
    
    # Always copy from master to working template before processing
    print("Copying fresh Module Answer Key template from master...")
    try:
        shutil.copy2(master_template_path, working_template_path)
        print(f"✓ Successfully copied master to working template")
    except Exception as e:
        raise Exception(f"Failed to copy master template: {e}")
    
    print("Loading Module Answer Key working template...")
    
    # Create a temporary copy of the working template - NEVER touch master directly
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        try:
            shutil.copy2(working_template_path, temp_file.name)
            temp_template_path = temp_file.name
            print(f"✓ Created temporary template at: {temp_template_path}")
        except Exception as e:
            raise Exception(f"Failed to create temporary template: {e}")
    
    # Track saved images for cleanup
    saved_images = []
    
    # Track processed image fields to avoid duplicates across both sections
    processed_enrichment_image_fields = set()
    processed_worksheet_image_fields = set()
    
    # Debug: Print all image files available in request
    image_files_in_request = [key for key in request.files.keys() if 'image_file' in key]
    print(f"Available image files in request: {image_files_in_request}")
    
    try:
        # Load the temporary template - NEVER the master
        print(f"Loading DocxTemplate from temporary file: {temp_template_path}")
        doc = DocxTemplate(temp_template_path)
        
        # Prepare context data matching the template structure
        
        # 1. Assessments Section
        # Pre-test questions
        pretest_questions_data = []
        for question_data in form.pretest_questions.data:
            if question_data.get('question_text'):
                pretest_questions_data.append({
                    'question': escape_xml(question_data['question_text']),
                    'choice': [
                        escape_xml(question_data['choice_a']),
                        escape_xml(question_data['choice_b']),
                        escape_xml(question_data['choice_c']),
                        escape_xml(question_data['choice_d'])
                    ],
                    'correct_answer': question_data.get('correct_answer', '').upper()
                })
        
        # RCA sessions (sessions 2-5, each with 3 questions)
        rca_sessions_data = []
        for i, session_data in enumerate(form.rca_sessions.data):
            session_obj = {
                'session_number': i + 2,  # Sessions 2-5
            }
            
            # Handle both form data structure (when coming from form submission)
            # and autosave data structure (when coming from draft loading)
            questions = []
            
            if 'questions' in session_data and isinstance(session_data['questions'], list):
                # Autosave/draft structure: questions as array
                questions = session_data['questions']
                print(f"🔍 RCA Session {i+1}: Using autosave structure with {len(questions)} questions")
            else:
                # Form submission structure: individual question objects
                questions = session_data.get('questions', [])
                print(f"🔍 RCA Session {i+1}: Using form structure with {len(questions)} questions")
            
            # Add research question (first question)
            if len(questions) > 0 and questions[0].get('question_text'):
                session_obj['research_question'] = {
                    'text': escape_xml(questions[0]['question_text']),
                    'choice': [
                        escape_xml(questions[0].get('choice_a', '')),
                        escape_xml(questions[0].get('choice_b', '')),
                        escape_xml(questions[0].get('choice_c', '')),
                        escape_xml(questions[0].get('choice_d', ''))
                    ],
                    'correct_answer': questions[0].get('correct_answer', '').upper()
                }
                print(f"🔍   Research question: {questions[0]['question_text'][:50]}...")
            
            # Add challenge question (second question)
            if len(questions) > 1 and questions[1].get('question_text'):
                session_obj['challenge_question'] = {
                    'text': escape_xml(questions[1]['question_text']),
                    'choice': [
                        escape_xml(questions[1].get('choice_a', '')),
                        escape_xml(questions[1].get('choice_b', '')),
                        escape_xml(questions[1].get('choice_c', '')),
                        escape_xml(questions[1].get('choice_d', ''))
                    ],
                    'correct_answer': questions[1].get('correct_answer', '').upper()
                }
                print(f"🔍   Challenge question: {questions[1]['question_text'][:50]}...")
            
            # Add application question (third question)
            if len(questions) > 2 and questions[2].get('question_text'):
                session_obj['application_question'] = {
                    'text': escape_xml(questions[2]['question_text']),
                    'choice': [
                        escape_xml(questions[2].get('choice_a', '')),
                        escape_xml(questions[2].get('choice_b', '')),
                        escape_xml(questions[2].get('choice_c', '')),
                        escape_xml(questions[2].get('choice_d', ''))
                    ],
                    'correct_answer': questions[2].get('correct_answer', '').upper()
                }
                print(f"🔍   Application question: {questions[2]['question_text'][:50]}...")
            
            # Only add sessions that have at least one question
            if any(key in session_obj for key in ['research_question', 'challenge_question', 'application_question']):
                rca_sessions_data.append(session_obj)
                print(f"🔍 Added RCA Session {i+1} with {len([k for k in ['research_question', 'challenge_question', 'application_question'] if k in session_obj])} questions")
        
        # Post-test questions
        posttest_questions_data = []
        for question_data in form.posttest_questions.data:
            if question_data.get('question_text'):
                posttest_questions_data.append({
                    'question': escape_xml(question_data['question_text']),
                    'choice': [
                        escape_xml(question_data['choice_a']),
                        escape_xml(question_data['choice_b']),
                        escape_xml(question_data['choice_c']),
                        escape_xml(question_data['choice_d'])
                    ],
                    'correct_answer': question_data.get('correct_answer', '').upper()
                })
        
        # 2. Performance Based Assessments - Changed from 4 to 3 sessions
        pba_sessions_data = []
        for i, session_data in enumerate(form.pba_sessions.data):
            if session_data.get('activity_name'):
                session_obj = {
                    'session_number': session_data.get('session_number', i + 1),  # Default to index + 1 if not provided
                    'activity_name': escape_xml(session_data['activity_name']),
                    'assessment_questions': []
                }
                
                # Process assessment questions for this session
                for question_data in session_data.get('assessment_questions', []):
                    if question_data.get('question'):
                        session_obj['assessment_questions'].append({
                            'question': escape_xml(question_data['question']),
                            'correct_answer': escape_xml(question_data.get('correct_answer', ''))
                        })
                
                if session_obj['assessment_questions']:  # Only add sessions with questions
                    pba_sessions_data.append(session_obj)
        
        # 3. Vocabulary
        vocabulary_data = []
        for term_data in form.vocabulary.data:
            if term_data.get('term'):
                vocabulary_data.append({
                    'term': escape_xml(term_data['term']),
                    'definition': escape_xml(term_data.get('definition', ''))
                })
        
        # 4. Student Portfolio Checklist
        portfolio_checklist_data = []
        for item_data in form.portfolio_checklist.data:
            if item_data.get('product'):
                portfolio_checklist_data.append({
                    'product': escape_xml(item_data['product']),
                    'session_number': escape_xml(item_data.get('session_number', ''))
                })
        
        # 5. Enrichment Activities & Worksheet Answer Keys - REMOVED FOR SIMPLIFICATION
        # These complex sections have been moved to separate "Academic Worksheet Builder" template
        
        # Create empty subdocuments for template compatibility
        enrichment_subdoc = doc.new_subdoc()
        worksheet_keys_subdoc = doc.new_subdoc()
        
        # Skip complex enrichment and worksheet processing - moved to separate template

        # Build the complete context (simplified - no enrichment or worksheet sections)
        context = {
            'module_acronym': escape_xml(form.module_acronym.data),
            'pretest_questions': pretest_questions_data,
            'rca_sessions': rca_sessions_data,
            'posttest_questions': posttest_questions_data,
            'pba_sessions': pba_sessions_data,
            'vocabulary': vocabulary_data,
            'portfolio_checklist': portfolio_checklist_data,
            'enrichment_dynamic_content': enrichment_subdoc,
            'worksheet_answer_keys': worksheet_keys_subdoc
        }
        
        print(f"Module Answer Key context data prepared")
        print(f"Number of pre-test questions: {len(pretest_questions_data)}")
        print(f"Number of RCA sessions: {len(rca_sessions_data)}")
        print(f"Number of post-test questions: {len(posttest_questions_data)}")
        print(f"Number of PBA sessions: {len(pba_sessions_data)}")
        print(f"Number of vocabulary terms: {len(vocabulary_data)}")
        print(f"Number of portfolio checklist items: {len(portfolio_checklist_data)}")
        
        # Render the document
        print("Rendering Module Answer Key document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"Module Answer Key {escape_xml(form.module_acronym.data).replace(' ', '_')}.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving Module Answer Key document to: {output_path}")
        doc.save(output_path)
        
        print("Module Answer Key document saved successfully!")
        
        # PROTECTION: Verify master template wasn't accidentally modified
        master_stat_after = os.stat(master_template_path)
        if master_stat.st_mtime != master_stat_after.st_mtime:
            print("⚠️  WARNING: Master template modification time changed during processing!")
        else:
            print("✓ Master template integrity verified - no accidental changes")
            
        return output_path
        
    finally:
        # Clean up the temporary file
        try:
            if os.path.exists(temp_template_path):
                os.unlink(temp_template_path)
                print(f"✓ Cleaned up temporary file: {temp_template_path}")
        except Exception as e:
            print(f"Warning: Could not clean up temporary file {temp_template_path}: {e}")


def generate_module_answer_key2(form):
    """Generate a Module Answer Key 2.0 using docxtpl - streamlined version without complex sections"""
    # Use the dedicated Module Answer Key 2.0 template (same Jinja syntax, no enrichment/worksheet sections)
    master_template_path = 'templates/docx_templates/module_answer_key2_master.docx'
    working_template_path = 'templates/docx_templates/module_answer_key2.docx'
    
    print(f"🔍 Looking for Module Answer Key 2.0 master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Module Answer Key 2.0 master DOCX template not found. Please create the master template first.")
    
    # PROTECTION: Verify master template integrity before proceeding
    master_stat = os.stat(master_template_path)
    print(f"🔍 Master template size: {master_stat.st_size} bytes, modified: {datetime.fromtimestamp(master_stat.st_mtime)}")
    
    # Always copy from master to working template before processing
    print("🔍 Copying fresh Module Answer Key 2.0 template from master...")
    try:
        shutil.copy2(master_template_path, working_template_path)
        print(f"✓ Successfully copied master to working template")
    except Exception as e:
        raise Exception(f"Failed to copy master template: {e}")
    
    print("🔍 Loading Module Answer Key 2.0 working template...")
    
    # Create a temporary copy of the working template - NEVER touch master directly
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        try:
            shutil.copy2(working_template_path, temp_file.name)
            temp_template_path = temp_file.name
            print(f"✓ Created temporary template at: {temp_template_path}")
        except Exception as e:
            raise Exception(f"Failed to create temporary template: {e}")
    
    try:
        # Load the temporary template - NEVER the master
        print(f"🔍 Loading DocxTemplate from temporary file: {temp_template_path}")
        doc = DocxTemplate(temp_template_path)
        
        # Prepare context data matching the template structure (simplified version)
        
        # 1. Pre-test questions - Fixed to match template expectation of {{ question.question }}
        pretest_questions_data = []
        for question_data in form.pretest_questions.data:
            if question_data.get('question_text'):
                pretest_questions_data.append({
                    'question': escape_xml(question_data['question_text']),  # Template uses {{ question.question }}
                    'choice': [
                        escape_xml(question_data.get('choice_a', '')),
                        escape_xml(question_data.get('choice_b', '')),
                        escape_xml(question_data.get('choice_c', '')),
                        escape_xml(question_data.get('choice_d', ''))
                    ],
                    'correct_answer': question_data.get('correct_answer', '').upper()
                })
        
        # 2. RCA sessions (sessions 2-5, each with 3 questions)
        rca_sessions_data = []
        for i, session_data in enumerate(form.rca_sessions.data):
            session_obj = {
                'session_number': i + 2,  # Sessions 2-5
            }
            
            # Handle both form data structure and autosave data structure
            questions = []
            
            if 'questions' in session_data and isinstance(session_data['questions'], list):
                # Autosave/draft structure: questions as array
                questions = session_data['questions']
                print(f"🔍 RCA Session {i+1}: Using autosave structure with {len(questions)} questions")
            else:
                # Form submission structure: individual question objects
                questions = session_data.get('questions', [])
                print(f"🔍 RCA Session {i+1}: Using form structure with {len(questions)} questions")
            
            # Add research question (first question)
            if len(questions) > 0 and questions[0].get('question_text'):
                session_obj['research_question'] = {
                    'text': escape_xml(questions[0]['question_text']),
                    'choice': [
                        escape_xml(questions[0].get('choice_a', '')),
                        escape_xml(questions[0].get('choice_b', '')),
                        escape_xml(questions[0].get('choice_c', '')),
                        escape_xml(questions[0].get('choice_d', ''))
                    ],
                    'correct_answer': questions[0].get('correct_answer', '').upper()
                }
                print(f"🔍   Research question: {questions[0]['question_text'][:50]}...")
            
            # Add challenge question (second question)
            if len(questions) > 1 and questions[1].get('question_text'):
                session_obj['challenge_question'] = {
                    'text': escape_xml(questions[1]['question_text']),
                    'choice': [
                        escape_xml(questions[1].get('choice_a', '')),
                        escape_xml(questions[1].get('choice_b', '')),
                        escape_xml(questions[1].get('choice_c', '')),
                        escape_xml(questions[1].get('choice_d', ''))
                    ],
                    'correct_answer': questions[1].get('correct_answer', '').upper()
                }
                print(f"🔍   Challenge question: {questions[1]['question_text'][:50]}...")
            
            # Add application question (third question)
            if len(questions) > 2 and questions[2].get('question_text'):
                session_obj['application_question'] = {
                    'text': escape_xml(questions[2]['question_text']),
                    'choice': [
                        escape_xml(questions[2].get('choice_a', '')),
                        escape_xml(questions[2].get('choice_b', '')),
                        escape_xml(questions[2].get('choice_c', '')),
                        escape_xml(questions[2].get('choice_d', ''))
                    ],
                    'correct_answer': questions[2].get('correct_answer', '').upper()
                }
                print(f"🔍   Application question: {questions[2]['question_text'][:50]}...")
            
            # Only add sessions that have at least one question
            if any(key in session_obj for key in ['research_question', 'challenge_question', 'application_question']):
                rca_sessions_data.append(session_obj)
                print(f"🔍 Added RCA Session {i+1} with {len([k for k in ['research_question', 'challenge_question', 'application_question'] if k in session_obj])} questions")
        
        # 3. Post-test questions - Fixed to match template expectation of {{ question.question }}
        posttest_questions_data = []
        for question_data in form.posttest_questions.data:
            if question_data.get('question_text'):
                posttest_questions_data.append({
                    'question': escape_xml(question_data['question_text']),  # Template uses {{ question.question }}
                    'choice': [
                        escape_xml(question_data.get('choice_a', '')),
                        escape_xml(question_data.get('choice_b', '')),
                        escape_xml(question_data.get('choice_c', '')),
                        escape_xml(question_data.get('choice_d', ''))
                    ],
                    'correct_answer': question_data.get('correct_answer', '').upper()
                })
        
        # 4. Performance Based Assessments (3 sessions, numbered 1-3)
        pba_sessions_data = []
        for i, session_data in enumerate(form.pba_sessions.data):
            if session_data.get('activity_name'):
                session_obj = {
                    'session_number': session_data.get('session_number', i + 1),  # Default to index + 1 if not provided
                    'activity_name': escape_xml(session_data['activity_name']),
                    'assessment_questions': []
                }
                
                # Process assessment questions for this session
                for question_data in session_data.get('assessment_questions', []):
                    if question_data.get('question'):
                        session_obj['assessment_questions'].append({
                            'question': escape_xml(question_data['question']),
                            'correct_answer': escape_xml(question_data.get('correct_answer', ''))
                        })
                
                if session_obj['assessment_questions']:  # Only add sessions with questions
                    pba_sessions_data.append(session_obj)
        
        # 5. Vocabulary
        vocabulary_data = []
        for term_data in form.vocabulary.data:
            if term_data.get('term'):
                vocabulary_data.append({
                    'term': escape_xml(term_data['term']),
                    'definition': escape_xml(term_data.get('definition', ''))
                })
        
        # 6. Student Portfolio Checklist
        portfolio_checklist_data = []
        for item_data in form.portfolio_checklist.data:
            if item_data.get('product'):
                portfolio_checklist_data.append({
                    'product': escape_xml(item_data['product']),
                    'session_number': escape_xml(item_data.get('session_number', ''))
                })
        
        # Build the complete context (simplified - no enrichment or worksheet sections)
        context = {
            'module_acronym': escape_xml(form.module_acronym.data),
            'pretest_questions': pretest_questions_data,
            'rca_sessions': rca_sessions_data,
            'posttest_questions': posttest_questions_data,
            'pba_sessions': pba_sessions_data,
            'vocabulary': vocabulary_data,
            'portfolio_checklist': portfolio_checklist_data
        }
        
        print(f"🔍 Module Answer Key 2.0 context data prepared")
        print(f"🔍 Number of pre-test questions: {len(pretest_questions_data)}")
        print(f"🔍 Number of RCA sessions: {len(rca_sessions_data)}")
        print(f"🔍 Number of post-test questions: {len(posttest_questions_data)}")
        print(f"🔍 Number of PBA sessions: {len(pba_sessions_data)}")
        print(f"🔍 Number of vocabulary terms: {len(vocabulary_data)}")
        print(f"🔍 Number of portfolio checklist items: {len(portfolio_checklist_data)}")
        
        # Debug: Print first few items to verify structure
        if vocabulary_data:
            print(f"🔍 First vocabulary item: {vocabulary_data[0]}")
        if portfolio_checklist_data:
            print(f"🔍 First portfolio item: {portfolio_checklist_data[0]}")
        
        # Render the document
        print("🔍 Rendering Module Answer Key 2.0 document...")
        try:
            doc.render(context)
        except Exception as template_error:
            print(f"🚨 Template rendering error: {str(template_error)}")
            print(f"🔍 This suggests a Jinja2 syntax error in the DOCX template file")
            print("🔍 Check the module_answer_key2_master.docx file for missing {% endfor %} tags")
            raise Exception("DOCX template syntax error: " + str(template_error) + ". Please check the module_answer_key2_master.docx template for missing {% endfor %} tags in loops.")
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"Module Answer Key 2.0 {escape_xml(form.module_acronym.data).replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"🔍 Saving Module Answer Key 2.0 document to: {output_path}")
        doc.save(output_path)
        
        print("✅ Module Answer Key 2.0 document saved successfully!")
        
        # PROTECTION: Verify master template wasn't accidentally modified
        master_stat_after = os.stat(master_template_path)
        if master_stat.st_mtime != master_stat_after.st_mtime:
            print("⚠️  WARNING: Master template modification time changed during processing!")
        else:
            print("✓ Master template integrity verified - no accidental changes")
            
        return output_path
        
    finally:
        # Clean up the temporary file
        try:
            if os.path.exists(temp_template_path):
                os.unlink(temp_template_path)
                print(f"✓ Cleaned up temporary file: {temp_template_path}")
        except Exception as e:
            print(f"Warning: Could not clean up temporary file {temp_template_path}: {e}")

def generate_module_activity_sheet(form):
    """Generate a Module Activity Sheet using docxtpl"""
    # Use a master template that never gets touched
    master_template_path = 'templates/docx_templates/module_activity_sheet_master.docx'
    working_template_path = 'templates/docx_templates/module_activity_sheet.docx'
    
    print(f"Looking for Module Activity Sheet master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Module Activity Sheet master DOCX template not found. Please create the master template first.")
    
    # PROTECTION: Verify master template integrity before proceeding
    master_stat = os.stat(master_template_path)
    print(f"Master template size: {master_stat.st_size} bytes, modified: {datetime.fromtimestamp(master_stat.st_mtime)}")
    
    # Always copy from master to working template before processing
    print("Copying fresh Module Activity Sheet template from master...")
    try:
        shutil.copy2(master_template_path, working_template_path)
        print(f"✓ Successfully copied master to working template")
    except Exception as e:
        raise Exception(f"Failed to copy master template: {e}")
    
    print("Loading Module Activity Sheet working template...")
    
    # Create a temporary copy of the working template - NEVER touch master directly
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        try:
            shutil.copy2(working_template_path, temp_file.name)
            temp_template_path = temp_file.name
            print(f"✓ Created temporary template at: {temp_template_path}")
        except Exception as e:
            raise Exception(f"Failed to create temporary template: {e}")
    
    try:
        # Load the temporary template - NEVER the master
        print(f"Loading DocxTemplate from temporary file: {temp_template_path}")
        doc = DocxTemplate(temp_template_path)
        
        # Build context with session data
        context = {
            'module_acronym': escape_xml(form.module_acronym.data),
        }
        
        # Process each session
        for i in range(1, 8):
            activity = getattr(form, f'session{i}_activity').data
            is_pba = getattr(form, f'session{i}_is_pba').data
            
            # Escape XML special characters for DOCX compatibility
            activity = escape_xml(activity)
            
            # Create session object matching template syntax
            session_data = {
                'activity': activity,
                'is_pba': is_pba,
                'has_pba': is_pba  # Same as is_pba for assessment logic
            }
            
            context[f'session{i}'] = session_data
        
        print(f"Module Activity Sheet context data: {context}")
        
        # Render the document
        print("Rendering Module Activity Sheet document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"Module Activity Sheet {escape_xml(form.module_acronym.data).replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving Module Activity Sheet document to: {output_path}")
        doc.save(output_path)
        
        print("Module Activity Sheet document saved successfully!")
        
        # PROTECTION: Verify master template wasn't accidentally modified
        master_stat_after = os.stat(master_template_path)
        if master_stat.st_mtime != master_stat_after.st_mtime:
            print("⚠️  WARNING: Master template modification time changed during processing!")
        else:
            print("✓ Master template integrity verified - no accidental changes")
            
        return output_path
        
    finally:
        # Clean up the temporary file
        try:
            if os.path.exists(temp_template_path):
                os.unlink(temp_template_path)
                print(f"✓ Cleaned up temporary file: {temp_template_path}")
        except Exception as e:
            print(f"Warning: Could not clean up temporary file {temp_template_path}: {e}")

def extract_pdf_text(pdf_file):
    """Extract text from uploaded PDF file"""
    try:
        # Reset file pointer to beginning
        pdf_file.seek(0)
        
        # Create a PDF reader object
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Extract text from all pages
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return None

def extract_session_data_from_text(text):
    """Extract session data from the Session Notes PDF with improved accuracy"""
    extracted_data = {
        'module_name': '',
        'enrichment': '',
        'sessions': []
    }
    
    try:
        # Clean up the text while preserving line breaks for structure
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\s*\n\s*', '\n', text)  # Clean up line breaks
        
        # Extract module name - try multiple strategies
        module_name = extract_module_name_advanced(text)
        if module_name:
            extracted_data['module_name'] = module_name
        
        # Split text by sessions using the exact pattern "Session [number]-"
        session_splits = re.split(r'Session\s*(\d+)\s*-', text, flags=re.IGNORECASE)
        
        # Process each session (skip index 0 which is content before first session)
        for i in range(1, len(session_splits), 2):
            if i + 1 < len(session_splits):
                session_num = int(session_splits[i])
                if session_num <= 7:  # Only process sessions 1-7
                    session_content = session_splits[i + 1]
                    
                    # Stop at the next session or enrichments section, but be more careful about boundaries
                    # Look for the start of the next session, but make sure it's actually a session header
                    next_session_match = re.search(r'(?:\n|^)\s*Session\s*\d+\s*-', session_content, re.IGNORECASE)
                    enrichment_match = re.search(r'(?:\n|^)\s*Enrichment', session_content, re.IGNORECASE)
                    
                    if next_session_match:
                        session_content = session_content[:next_session_match.start()]
                    elif enrichment_match:
                        session_content = session_content[:enrichment_match.start()]
                    
                    # Debug output for session content
                    print(f"DEBUG: Processing Session {session_num}, content length: {len(session_content)}")
                    if 'Performance' in session_content:
                        perf_index = session_content.lower().find('performance')
                        print(f"DEBUG: 'Performance' found at position {perf_index} in session content")
                        print(f"DEBUG: Text around Performance: '{session_content[max(0, perf_index-50):perf_index+150]}'")
                    else:
                        print(f"DEBUG: 'Performance' NOT found in session content")
                    
                    session_data = extract_session_fields_improved(session_content)
                    extracted_data['sessions'].append(session_data)
        
        # Extract enrichment activities (everything after "Enrichments")
        enrichment_match = re.search(r'Enrichments?\s*(.*?)(?=\Z)', text, re.IGNORECASE | re.DOTALL)
        if enrichment_match:
            enrichment_text = enrichment_match.group(1).strip()
            # Clean up the enrichment text
            enrichment_text = re.sub(r'^\d+\.?\s*', '', enrichment_text, flags=re.MULTILINE)  # Remove numbering
            extracted_data['enrichment'] = clean_text(enrichment_text)
    
    except Exception as e:
        print(f"Error extracting session data: {e}")
    
    return extracted_data

def extract_module_name_advanced(text):
    """Advanced module name extraction using multiple strategies"""
    
    # Strategy 1: Look for text immediately after "Session Notes" 
    match = re.search(r'Session\s+Notes\s+([^\n\r]+)', text, re.IGNORECASE)
    if match:
        potential_name = match.group(1).strip()
        potential_name = remove_header_artifacts(potential_name)
        if potential_name and len(potential_name) > 3 and len(potential_name) < 50:
            return potential_name
    
    # Strategy 2: Look for text before "Session Notes"
    match = re.search(r'([^\n\r]+)\s+Session\s+Notes', text, re.IGNORECASE)
    if match:
        potential_name = match.group(1).strip()
        potential_name = remove_header_artifacts(potential_name)
        if potential_name and len(potential_name) > 3 and len(potential_name) < 50:
            return potential_name
    
    # Strategy 3: Look at the very beginning, skip common artifacts
    lines = text.split('\n')[:15]  # Check first 15 lines
    
    for line in lines:
        line = line.strip()
        # Skip obvious artifacts but look for substantial content
        if (line and 
            not re.search(r'Session\s+Notes', line, re.IGNORECASE) and
            not re.search(r'star\s*ACADEMY', line, re.IGNORECASE) and
            not re.search(r'NOLA\s+EDUCATION', line, re.IGNORECASE) and
            not re.search(r'www\.', line, re.IGNORECASE) and
            not re.search(r'Session\s+\d+', line, re.IGNORECASE) and
            not re.search(r'Focus:', line, re.IGNORECASE) and
            not re.search(r'Objectives:', line, re.IGNORECASE) and
            len(line) > 3 and len(line) < 50):
            
            # Clean up the line
            clean_line = remove_header_artifacts(line)
            clean_line = re.sub(r'[^\w\s&\'-]', '', clean_line).strip()
            
            # Check if it looks like a valid module name
            if (clean_line and 
                len(clean_line) > 3 and 
                not clean_line.lower() in ['session', 'notes', 'page', 'may', 'be', 'photocopied']):
                return clean_line
    
    # Strategy 4: Look for common module name patterns in the text
    common_patterns = [
        r'\b(Organism\s+Reproduction)\b',
        r'\b(Animals?)\b(?!\s+having)',  # "Animals" but not "animals having"
        r'\b(Plants?)\b(?!\s+reproduce)',
        r'\b(Chemistry)\b',
        r'\b(Physics)\b',
        r'\b(Biology)\b',
        r'\b(Environmental?\s+Science)\b',
        r'\b(Earth\s+Science)\b',
        r'\b(Sports?\s+Statistics)\b',
    ]
    
    for pattern in common_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)
    
    return None

def extract_session_fields_improved(session_text):
    """Extract individual fields from a session text block with improved accuracy"""
    session_data = {
        'focus': '',
        'objectives': '',
        'materials': '',
        'teacher_prep': '',
        'assessments': ''
    }
    
    # More precise patterns based on the actual document structure
    field_patterns = {
        'focus': [
            r'Focus:\s*(.*?)(?=\s*Objectives:)',
            r'Focus:\s*(.*?)(?=\s*(?:Objectives|Materials|Teacher|Performance))',
        ],
        'objectives': [
            r'Objectives:\s*(.*?)(?=\s*Materials:)',
            r'Objectives:\s*(.*?)(?=\s*(?:Materials|Teacher|Performance))',
        ],
        'materials': [
            r'Materials:\s*(.*?)(?=\s*Teacher\s*Preparations?:)',
            r'Materials:\s*(.*?)(?=\s*(?:Teacher|Performance))',
        ],
        'teacher_prep': [
            r'Teacher\s*Preparations?:\s*(.*?)(?=\s*Performance\s*(?:Based\s*)?(?:Assessment\s*)?Questions?:)',
            r'Teacher\s*Preparations?:\s*(.*?)(?=\s*(?:Performance|Assessment))',
        ],
        'assessments': [
            r'Performance\s*Assessment\s*Questions?\s*:?\s*(.*?)(?=\s*(?:Session\s*\d+|Enrichments?|$))',
            r'Performance\s*Based\s*Assessment\s*Questions?\s*:?\s*(.*?)(?=\s*(?:Session\s*\d+|Enrichments?|$))',
            r'Performance\s*(?:Based\s*)?(?:Assessment\s*)?Questions?\s*:?\s*(.*?)(?=\s*(?:Session\s*\d+|Enrichments?|$))',
            r'Assessment\s*Questions?\s*:?\s*(.*?)(?=\s*(?:Session\s*\d+|Enrichments?|$))',
            r'PBA\s*:?\s*(.*?)(?=\s*(?:Session\s*\d+|Enrichments?|$))',
        ]
    }
    
    # Extract each field using the patterns
    for field, patterns in field_patterns.items():
        extracted = False
        for i, pattern in enumerate(patterns):
            match = re.search(pattern, session_text, re.IGNORECASE | re.DOTALL)
            if match:
                extracted_text = match.group(1).strip()
                
                # Debug output for assessments
                if field == 'assessments':
                    print(f"DEBUG: Assessments found with pattern {i+1}: '{extracted_text[:100]}...'")
                
                # Clean up the extracted text based on field type
                if field == 'focus':
                    session_data[field] = clean_focus_text(extracted_text)
                elif field == 'objectives':
                    session_data[field] = clean_objectives_text(extracted_text)
                elif field == 'teacher_prep':
                    session_data[field] = clean_teacher_prep_text(extracted_text)
                elif field == 'assessments':
                    session_data[field] = clean_assessment_text(extracted_text)
                else:
                    session_data[field] = clean_text(extracted_text)
                extracted = True
                break
        
        # Debug output for failed extractions
        if not extracted and field == 'assessments':
            print(f"DEBUG: No assessments found. Session text sample: '{session_text[:200]}...'")
        
        # If no data was extracted for teacher_prep or assessments, set to N/A
        if not extracted:
            if field in ['teacher_prep', 'assessments']:
                session_data[field] = 'N/A'
    
    return session_data

def clean_focus_text(text):
    """Clean session focus text with proper title case"""
    if not text:
        return ''
    
    # Remove header/footer artifacts first
    text = remove_header_artifacts(text)
    
    # Basic cleanup
    text = clean_text(text)
    
    # Apply title case while preserving acronyms and special terms
    words = text.split()
    title_words = []
    
    for word in words:
        # Keep short words like &, and, or lowercase unless at start
        if word.lower() in ['&', 'and', 'or', 'the', 'of', 'in', 'on', 'at', 'to', 'for', 'with'] and title_words:
            title_words.append(word.lower())
        else:
            # Capitalize first letter of each word
            title_words.append(word.capitalize())
    
    return ' '.join(title_words)

def remove_header_artifacts(text):
    """Remove common header/footer artifacts from PDF text"""
    if not text:
        return text
    
    # List of common header/footer patterns to remove
    artifacts = [
        r'Nola\s+Education,?\s*Llc',
        r'NOLA\s+EDUCATION,?\s*LLC',
        r'www\.staracademyprogram\.com',
        r'This\s+Page\s+May\s+Be\s+Photocopied\s+for\s+Use\s+Only\s+Within\s+the.*?Star\s+Academy',
        r'This\s+page\s+may\s+be\s+photocopied\s+for\s+use\s+only\s+within\s+the.*?Star\s+Academy',
        r'star\s*ACADEMY',
        r'Session\s+Notes\s+Animals',
        r'⎢.*?\.com',
        r'www\..*?\.com',
        r'\|\s*www\.staracademyprogram\.com',
        r'NOLA\s+EDUCATION,?\s*LLC\s*\|\s*www\.staracademyprogram\.com',
        r'NOLA\s+EDUCATION,?\s*LLC\s*[|\u007c\u2502\u2500-\u257F]*\s*www\.staracademyprogram\.com',
    ]
    
    # Remove each artifact pattern
    for pattern in artifacts:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    # Clean up extra whitespace
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    
    return text

def clean_objectives_text(text):
    """Clean objectives text and ensure proper punctuation"""
    if not text:
        return ''
    
    # Handle "N/A" case
    if re.match(r'^\s*n/?a\s*$', text, re.IGNORECASE):
        return 'N/A'
    
    # Remove header artifacts first
    text = remove_header_artifacts(text)
    
    # Basic cleanup
    text = clean_text(text)
    
    # Split into sentences (by common patterns)
    sentences = re.split(r'(?<=[.!?])\s+|(?<=[a-z])\s+(?=[A-Z])', text)
    
    cleaned_sentences = []
    for sentence in sentences:
        sentence = sentence.strip()
        if sentence:
            # Ensure sentence ends with a period if it doesn't have punctuation
            if not re.search(r'[.!?]$', sentence):
                sentence += '.'
            cleaned_sentences.append(sentence)
    
    return ' '.join(cleaned_sentences)

def clean_teacher_prep_text(text):
    """Clean teacher preparation text and ensure proper punctuation"""
    if not text:
        return 'N/A'
    
    # Handle "N/A" case
    if re.match(r'^\s*n/?a\s*$', text, re.IGNORECASE):
        return 'N/A'
    
    # Basic cleanup
    text = clean_text(text)
    
    # Ensure it ends with a period if it doesn't have punctuation
    if text and not re.search(r'[.!?]$', text):
        text += '.'
    
    return text if text else 'N/A'

def clean_assessment_text(text):
    """Clean assessment text while preserving list structure"""
    if not text:
        return 'N/A'
    
    # Handle "N/A" case
    if re.match(r'^\s*n/?a\s*$', text, re.IGNORECASE):
        return 'N/A'
    
    # Remove any "Enrichments" text that might have leaked in
    text = re.sub(r'\s*Enrichments?\.?\s*.*$', '', text, flags=re.IGNORECASE | re.DOTALL)
    
    # Clean up but preserve numbered lists
    text = re.sub(r'\s+', ' ', text.strip())
    
    # Fix common PDF extraction issues
    text = fix_pdf_text_issues(text)
    
    # If it contains numbered items, format them nicely and add punctuation
    if re.search(r'\d+\.', text):
        # Split by numbers and rejoin with proper formatting
        parts = re.split(r'(\d+\.)', text)
        formatted_parts = []
        for i, part in enumerate(parts):
            if re.match(r'\d+\.', part):
                formatted_parts.append(f"\n{part}")
            elif part.strip():
                sentence = part.strip()
                
                # Fix spacing before punctuation
                sentence = re.sub(r'\s+([.!?])', r'\1', sentence)
                
                # Add period if missing
                if not re.search(r'[.!?]$', sentence):
                    sentence += '.'
                formatted_parts.append(f" {sentence}")
        text = ''.join(formatted_parts).strip()
    else:
        # Single sentence - fix spacing and add period if missing
        text = re.sub(r'\s+([.!?])', r'\1', text)
        if not re.search(r'[.!?]$', text):
            text += '.'
    
    return text if text else 'N/A'

def fix_pdf_text_issues(text):
    """Fix common PDF text extraction issues"""
    if not text:
        return text
    
    # Fix broken words (common PDF extraction issues)
    text = re.sub(r'cephalizatio\s*n', 'cephalization', text, flags=re.IGNORECASE)
    text = re.sub(r'segmentatio\s*n', 'segmentation', text, flags=re.IGNORECASE)
    text = re.sub(r'classificatio\s*n', 'classification', text, flags=re.IGNORECASE)
    text = re.sub(r'organizatio\s*n', 'organization', text, flags=re.IGNORECASE)
    text = re.sub(r'observatio\s*n', 'observation', text, flags=re.IGNORECASE)
    
    # Fix spacing issues around punctuation
    text = re.sub(r'\s+([.!?,:;])', r'\1', text)
    
    return text

def clean_text(text):
    """Clean and normalize extracted text"""
    if not text:
        return ''
    
    # Remove extra whitespace and normalize
    text = re.sub(r'\s+', ' ', text.strip())
    
    # Remove common PDF artifacts
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
    
    # Fix spacing issues around punctuation (remove spaces before punctuation)
    text = re.sub(r'\s+([.!?,:;])', r'\1', text)
    
    return text

def generate_horizontal_lesson_plan(form):
    """Generate a horizontal lesson plan using docxtpl"""
    # Use the horizontal lesson plan master template
    master_template_path = 'templates/docx_templates/horizontal_lesson_plan_master.docx'
    
    print(f"Looking for horizontal lesson plan master template at: {master_template_path}")
    
    # Check if master template exists
    if not os.path.exists(master_template_path):
        raise FileNotFoundError("Horizontal Lesson Plan master DOCX template not found. Please create the master template first.")
    
    # Create a temporary copy of the master template
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(master_template_path, temp_file.name)
        temp_template_path = temp_file.name
    
    try:
        # Load the temporary template
        doc = DocxTemplate(temp_template_path)
        
        # Initialize the context with basic fields matching your template exactly
        context = {
            'school': {
                'name': escape_xml(form.school_name.data)
            },
            'teacher': {
                'name': escape_xml(form.teacher_name.data)
            },
            'term': escape_xml(form.term.data)
        }
        
        # Process each module using your template's flat naming convention
        for i, module in enumerate(form.modules):
            module_num = i + 1  # m1, m2, m3, m4, m5
            
            if module.module_name.data or any(session.focus.data or session.objectives.data 
                                           for session in module.sessions):
                
                # Add module name: {{ m1.name }}, {{ m2.name }}, etc.
                context[f'm{module_num}'] = {
                    'name': escape_xml(module.module_name.data or f"Module {module_num}"),
                    'enrichment': escape_xml(module.enrichment_activities.data or '')
                }
                
                # Process sessions for this module: {{ m1.s1.focus }}, {{ m1.s2.focus }}, etc.
                for j, session in enumerate(module.sessions):
                    session_num = j + 1  # s1, s2, s3, s4, s5, s6, s7
                    
                    # Create session data with your template's field names
                    session_data = {
                        'focus': escape_xml(session.focus.data or ''),
                        'objectives': escape_xml(session.objectives.data or ''),
                        'materials': escape_xml(session.materials.data or ''),
                        'prep': escape_xml(session.teacher_prep.data or ''),  # 'prep' not 'teacher_prep'
                        'pba': escape_xml(session.assessments.data or '')    # 'pba' not 'assessments'
                    }
                    
                    # Add session to module: context['m1']['s1'] = session_data
                    context[f'm{module_num}'][f's{session_num}'] = session_data
        
        print(f"Horizontal Lesson Plan context prepared")
        # Count modules with data by checking for m1, m2, m3, m4, m5 keys
        module_count = sum(1 for key in context.keys() if key.startswith('m') and key[1:].isdigit())
        print(f"Number of modules with data: {module_count}")
        
        # Render the document
        print("Rendering horizontal lesson plan document...")
        doc.render(context)
        
        # Save to output directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"Horizontal Lesson Plan {escape_xml(form.school_name.data).replace(' ', '_')} {escape_xml(form.teacher_name.data).replace(' ', '_')}_v2.0.docx"
        output_path = os.path.join(output_dir, filename)
        
        print(f"Saving horizontal lesson plan document to: {output_path}")
        doc.save(output_path)
        
        print("Horizontal lesson plan document saved successfully!")
        return output_path, filename
        
    finally:
        # Clean up the temporary file
        try:
            os.unlink(temp_template_path)
        except:
            pass  # Ignore cleanup errors

@app.route('/debug/db-status')
def debug_db_status():
    """Debug route to check database status - DEVELOPMENT ONLY"""
    # SECURITY: Only allow in development
    if app.config.get('IS_PRODUCTION', False):
        return "Debug routes disabled in production", 403
    
    try:
        # Don't call db.create_all() - just check existing tables
        total_users = User.query.count()
        admin_count = User.query.filter_by(is_admin=True).count()
        
        # Test database connection without modifications
        test_query = db.session.execute(db.text('SELECT 1')).scalar()
        
        return f"""
        <h1>Database Status - OK</h1>
        <ul>
        <li>Database URL: {app.config.get('SQLALCHEMY_DATABASE_URI', 'Not set')[:50]}...</li>
        <li>Total users: {total_users}</li>
        <li>Admin users: {admin_count}</li>
        <li>Database connection: {test_query}</li>
        </ul>
        <p><a href="/setup">Go to Setup</a> | <a href="/">Home</a></p>
        """
        
    except Exception as e:
        return f"""
        <h1>Database Status - ERROR</h1>
        <p><strong>Error:</strong> {str(e)}</p>
        <p><strong>Database URL:</strong> {app.config.get('SQLALCHEMY_DATABASE_URI', 'Not set')[:50]}...</p>
        <p><a href="/">Home</a></p>
        """

@app.route('/migrate-db')
def migrate_db():
    """Safely migrate database - DEVELOPMENT ONLY"""
    # SECURITY: Only allow in development
    if app.config.get('IS_PRODUCTION', False):
        return "Debug routes disabled in production", 403
    
    try:
        # Only create missing tables (safe operation) - don't drop existing ones
        db.create_all()
        
        # Test that we can query existing users
        total_users = User.query.count()
        
        return f"""
        <h1>✅ Database Migration Complete</h1>
        <p><strong>Status:</strong> Missing tables created successfully</p>
        <p><strong>Existing users preserved:</strong> {total_users} users found</p>
        <p>Your login should now work!</p>
        <p><a href="/login">Try Login Again</a> | <a href="/">Home</a></p>
        """
        
    except Exception as e:
        return f"""
        <h1>Database Migration - ERROR</h1>
        <p><strong>Error:</strong> {str(e)}</p>
        <p><a href="/debug/db-status">Check Database Status</a></p>
        """

@app.route('/create-admin-simple', methods=['GET', 'POST'])
def create_admin_simple():
    """Simple admin creation with better error handling - DEVELOPMENT ONLY"""
    # SECURITY: Only allow in development or when no admin exists
    if app.config.get('IS_PRODUCTION', False):
        # In production, only allow if no admin exists
        existing_admin = User.query.filter_by(is_admin=True).first()
        if existing_admin:
            return "Admin creation disabled - admin already exists", 403
    
    try:
        # Don't call db.create_all() - use migrations instead
        # Tables should already exist via flask db upgrade
        
        if request.method == 'GET':
            # Check if admin already exists
            existing_admin = User.query.filter_by(is_admin=True).first()
            if existing_admin:
                return f"""
                <h2>Admin Already Exists</h2>
                <p>Email: {existing_admin.email}</p>
                <p>Username: {existing_admin.username}</p>
                <p><a href="/login">Go to Login</a></p>
                """
            
            return """
            <h2>Create Admin User</h2>
            <form method="POST">
                <p><label>Email: <input type="email" name="email" required></label></p>
                <p><label>Username: <input type="text" name="username" required></label></p>
                <p><label>Password: <input type="password" name="password" required></label></p>
                <p><label>First Name: <input type="text" name="first_name"></label></p>
                <p><label>Last Name: <input type="text" name="last_name"></label></p>
                <p><button type="submit">Create Admin</button></p>
            </form>
            """
        
        # POST request - create admin
        email = request.form.get('email')
        username = request.form.get('username') 
        password = request.form.get('password')
        first_name = request.form.get('first_name', '')
        last_name = request.form.get('last_name', '')
        
        if not email or not username or not password:
            return "<h2>Error: Email, username and password required</h2><a href='/create-admin-simple'>Try Again</a>"
        
        # Check if user exists
        if User.query.filter_by(email=email).first():
            return f"<h2>Error: Email {email} already exists</h2><a href='/create-admin-simple'>Try Again</a>"
        
        # Create admin user
        admin = User(
            email=email,
            username=username,
            first_name=first_name,
            last_name=last_name,
            is_admin=True,
            is_active=True
        )
        admin.set_password(password)
        
        db.session.add(admin)
        db.session.commit()
        
        return f"""
        <h2>Admin Created Successfully!</h2>
        <p>Email: {email}</p>
        <p>Username: {username}</p>
        <p><a href="/login">Go to Login</a></p>
        """
        
    except Exception as e:
        return f"""
        <h2>Error Creating Admin</h2>
        <p>{str(e)}</p>
        <p><a href="/debug/db-status">Check Database Status</a></p>
        """

# New routes for save/load/download functionality
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_vocabulary'))
    
    try:
        # Create form and populate with draft data
        form = VocabularyWorksheetForm()
        form_data = draft.form_data
        
        # Populate form fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        
        # Populate vocabulary words
        words_data = form_data.get('words', [])
        for i, word_data in enumerate(words_data):
            if i < len(form.words):
                form.words[i].word.data = word_data.get('word', '')
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_vocabulary.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_vocabulary'))

@app.route('/vocabulary-drafts')
@login_required
def vocabulary_drafts():
    """List user's vocabulary worksheet drafts"""
    drafts = FormDraft.query.filter_by(
        user_id=current_user.id, 
        form_type='vocabulary',
        is_current=True
    ).order_by(FormDraft.updated_at.desc()).all()
    
    return render_template('vocabulary_drafts.html', drafts=drafts)

@app.route('/delete-vocabulary-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_vocabulary_draft(draft_id):
    """Delete vocabulary worksheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='vocabulary').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

@app.route('/download/<int:doc_id>')
@login_required
def download_document(doc_id):
    """Download a generated document with improved error handling"""
    try:
        document = GeneratedDocument.query.filter_by(id=doc_id, user_id=current_user.id).first()
        
        if not document:
            print(f"⚠️  Download attempt for non-existent document ID: {doc_id} by user: {current_user.id}")
            flash('Document not found. It may have been removed or you may not have permission to access it.', 'error')
            return redirect(url_for('my_documents'))
        
        # Enhanced file existence check
        if not document.file_path:
            print(f"❌ Document {doc_id} has no file path")
            flash(f'Document "{document.filename}" has no file path. Please contact support.', 'error')
            # Remove the broken record
            db.session.delete(document)
            db.session.commit()
            return redirect(url_for('my_documents'))
        
        if not os.path.exists(document.file_path):
            print(f"❌ File not found on disk: {document.file_path}")
            flash(f'The file for "{document.filename}" is no longer available. The document record has been removed.', 'warning')
            # Remove the broken record
            db.session.delete(document)
            db.session.commit()
            return redirect(url_for('my_documents'))
        
        # Verify file is readable
        try:
            file_size = os.path.getsize(document.file_path)
            if file_size == 0:
                print(f"❌ File is empty: {document.file_path}")
                flash(f'The file for "{document.filename}" appears to be corrupted (0 bytes).', 'error')
                return redirect(url_for('my_documents'))
        except OSError as e:
            print(f"❌ Cannot access file: {document.file_path}, error: {e}")
            flash(f'Cannot access the file for "{document.filename}". Please try again later.', 'error')
            return redirect(url_for('my_documents'))
        
        # Track download
        document.increment_download()
        db.session.commit()
        
        # Activity logging
        ActivityLog.log_activity('document_downloaded', current_user.id, 
                                {'document_type': document.document_type, 'filename': document.filename}, 
                                request)
        db.session.commit()
        
        print(f"✅ Successful download: {document.filename} by user {current_user.id}")
        
        # Return file for download
        return send_file(document.file_path, as_attachment=True, download_name=document.filename)
        
    except Exception as e:
        print(f"🚨 Download error for doc_id {doc_id}: {str(e)}")
        flash('An unexpected error occurred while downloading the document. Please try again or contact support.', 'error')
        return redirect(url_for('my_documents'))

@app.route('/my-documents')
@login_required
def my_documents():
    """List user's generated documents"""
    documents = GeneratedDocument.query.filter_by(
        user_id=current_user.id
    ).order_by(GeneratedDocument.created_at.desc()).all()
    
    return render_template('my_documents.html', documents=documents)

# Routes for vocabulary draft management (moved here to avoid conflicts)
@app.route('/load-vocabulary-draft/<int:draft_id>')
@login_required
def load_vocabulary_draft(draft_id):
    """Load vocabulary worksheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='vocabulary').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_vocabulary'))
    
    try:
        # Create form and populate with draft data
        form = VocabularyWorksheetForm()
        form_data = draft.form_data
        
        # Populate form fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        
        # Populate vocabulary words
        words_data = form_data.get('words', [])
        for i, word_data in enumerate(words_data):
            if i < len(form.words):
                form.words[i].word.data = word_data.get('word', '')
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_vocabulary.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_vocabulary'))

@app.route('/drafts')
@login_required
def drafts():
    """List all user's drafts (all document types)"""
    drafts = FormDraft.query.filter_by(
        user_id=current_user.id,
        is_current=True
    ).order_by(FormDraft.updated_at.desc()).all()
    
    return render_template('drafts.html', drafts=drafts)

@app.route('/delete-document/<int:doc_id>', methods=['POST'])
@login_required
def delete_document(doc_id):
    """Delete a generated document"""
    document = GeneratedDocument.query.filter_by(id=doc_id, user_id=current_user.id).first()
    
    if not document:
        flash('Document not found', 'error')
        return redirect(url_for('my_documents'))
    
    try:
        # Delete the physical file if it exists
        if document.file_path and os.path.exists(document.file_path):
            os.remove(document.file_path)
            print(f"✓ Deleted file: {document.file_path}")
        
        # Store filename for success message
        filename = document.filename
        
        # Delete the database record
        db.session.delete(document)
        db.session.commit()
        
        flash(f'Document "{filename}" removed successfully!', 'success')
        
    except Exception as e:
        print(f"Error deleting document: {e}")
        flash(f'Error removing document: {str(e)}', 'error')
    
    return redirect(url_for('my_documents'))

@app.route('/autosave-test-draft', methods=['POST'])
@login_required
def autosave_test_draft():
    """AJAX endpoint for autosaving test draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'module_acronym': data.get('module_acronym', ''),
            'test_type': data.get('test_type', 'pre'),
            'questions': []
        }
        
        # Process test questions
        questions_data = data.get('questions', [])
        for question_data in questions_data:
            question_text = question_data.get('question_text', '').strip()
            choice_a = question_data.get('choice_a', '').strip()
            choice_b = question_data.get('choice_b', '').strip()
            choice_c = question_data.get('choice_c', '').strip()
            choice_d = question_data.get('choice_d', '').strip()
            
            # Only include questions with some content
            if question_text or choice_a or choice_b or choice_c or choice_d:
                form_data['questions'].append({
                    'question_text': question_text,
                    'choice_a': choice_a,
                    'choice_b': choice_b,
                    'choice_c': choice_c,
                    'choice_d': choice_d
                })
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='test').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                # Update title if module acronym or test type changed
                test_type_display = "Pre-Test" if form_data['test_type'] == 'pre' else "Post-Test"
                if form_data['module_acronym']:
                    draft.title = f"{test_type_display} Worksheet - {form_data['module_acronym']}"
                    draft.module_acronym = form_data['module_acronym']
            else:
                return jsonify({'success': False, 'error': 'Draft not found'})
        else:
            # Create new draft
            test_type_display = "Pre-Test" if form_data['test_type'] == 'pre' else "Post-Test"
            title = f"{test_type_display} Worksheet - {form_data['module_acronym']}" if form_data['module_acronym'] else f"{test_type_display} Worksheet - Untitled"
            draft = FormDraft(
                user_id=current_user.id,
                form_type='test',
                title=title,
                module_acronym=form_data['module_acronym'],
                form_data=form_data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'draft_id': draft.id,
            'message': 'Draft saved automatically',
            'timestamp': datetime.utcnow().strftime('%I:%M:%S %p')
        })
        
    except Exception as e:
        print(f"Error in test autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/load-test-draft/<int:draft_id>')
@login_required
def load_test_draft(draft_id):
    """Load test worksheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='test').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_test'))
    
    try:
        # Create form and populate with draft data
        form = TestWorksheetForm()
        form_data = draft.form_data
        
        # Populate form fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        form.test_type.data = form_data.get('test_type', 'pre')
        
        # Populate test questions
        questions_data = form_data.get('questions', [])
        for i, question_data in enumerate(questions_data):
            if i < len(form.questions):
                form.questions[i].question_text.data = question_data.get('question_text', '')
                form.questions[i].choice_a.data = question_data.get('choice_a', '')
                form.questions[i].choice_b.data = question_data.get('choice_b', '')
                form.questions[i].choice_c.data = question_data.get('choice_c', '')
                form.questions[i].choice_d.data = question_data.get('choice_d', '')
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_test.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading test draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_test'))

@app.route('/delete-test-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_test_draft(draft_id):
    """Delete test worksheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='test').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

@app.route('/load-pba-draft/<int:draft_id>')
@login_required
def load_pba_draft(draft_id):
    """Load PBA worksheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='pba').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_pba'))
    
    try:
        # Create form and populate with draft data
        form = PBAWorksheetForm()
        form_data = draft.form_data
        
        # Populate basic fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        form.session_number.data = form_data.get('session_number', '')
        form.section_header.data = form_data.get('section_header', '')
        
        # Populate assessment fields
        assessments_data = form_data.get('assessments', [])
        for i, assessment_data in enumerate(assessments_data):
            if i < len(form.assessments):
                form.assessments[i].assessment.data = assessment_data.get('assessment', '')
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_pba.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading PBA draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_pba'))

@app.route('/delete-pba-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_pba_draft(draft_id):
    """Delete PBA worksheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='pba').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

@app.route('/autosave-pba-draft', methods=['POST'])
@login_required
def autosave_pba_draft():
    """AJAX endpoint for autosaving PBA draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'module_acronym': data.get('module_acronym', ''),
            'session_number': data.get('session_number', ''),
            'section_header': data.get('section_header', ''),
            'assessments': []
        }
        
        # Process assessment fields
        assessments_data = data.get('assessments', [])
        for assessment_data in assessments_data:
            assessment_text = assessment_data.get('assessment', '').strip()
            # Only include non-empty assessments
            if assessment_text:
                form_data['assessments'].append({
                    'assessment': assessment_text
                })
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='pba').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                # Update title based on form data
                if form_data['module_acronym'] or form_data['session_number']:
                    draft.title = f"PBA Worksheet - {form_data['module_acronym']} Session {form_data['session_number']}"
                    draft.module_acronym = form_data['module_acronym']
            else:
                return jsonify({'success': False, 'error': 'Draft not found'})
        else:
            # Create new draft
            title_parts = []
            if form_data['module_acronym']:
                title_parts.append(form_data['module_acronym'])
            if form_data['session_number']:
                title_parts.append(f"Session {form_data['session_number']}")
            
            title = f"PBA Worksheet - {' '.join(title_parts)}" if title_parts else "PBA Worksheet - Untitled"
            
            draft = FormDraft(
                user_id=current_user.id,
                form_type='pba',
                title=title,
                module_acronym=form_data['module_acronym'],
                form_data=form_data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'draft_id': draft.id,
            'message': 'Draft saved automatically',
            'timestamp': datetime.utcnow().strftime('%I:%M:%S %p')
        })
        
    except Exception as e:
        print(f"Error in PBA autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/autosave-generic-draft', methods=['POST'])
@login_required  
def autosave_generic_draft():
    """AJAX endpoint for autosaving generic worksheet draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'module_acronym': data.get('module_acronym', ''),
            'dynamic_fields': data.get('dynamic_fields', [])
        }
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='generic').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                # Update title based on form data
                if form_data['module_acronym']:
                    draft.title = f"Generic Worksheet - {form_data['module_acronym']}"
                    draft.module_acronym = form_data['module_acronym']
            else:
                return jsonify({'success': False, 'error': 'Draft not found'})
        else:
            # Create new draft
            title = f"Generic Worksheet - {form_data['module_acronym']}" if form_data['module_acronym'] else "Generic Worksheet - Untitled"
            draft = FormDraft(
                user_id=current_user.id,
                form_type='generic',
                title=title,
                module_acronym=form_data['module_acronym'],
                form_data=form_data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'draft_id': draft.id,
            'message': 'Draft saved automatically',
            'timestamp': datetime.utcnow().strftime('%I:%M:%S %p')
        })
        
    except Exception as e:
        print(f"Error in generic autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/autosave-moduleactivity-draft', methods=['POST'])
@login_required
def autosave_moduleactivity_draft():
    """AJAX endpoint for autosaving module activity sheet draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'module_acronym': data.get('module_acronym', ''),
            'sessions': []
        }
        
        # Process session data (7 sessions)
        for i in range(1, 8):
            session_activity = data.get(f'session{i}_activity', '').strip()
            session_is_pba = data.get(f'session{i}_is_pba', False)
            
            if session_activity or session_is_pba:
                form_data['sessions'].append({
                    'session_number': i,
                    'activity': session_activity,
                    'is_pba': session_is_pba
                })
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='moduleactivity').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                # Update title based on form data
                if form_data['module_acronym']:
                    draft.title = f"Module Activity Sheet - {form_data['module_acronym']}"
                    draft.module_acronym = form_data['module_acronym']
            else:
                return jsonify({'success': False, 'error': 'Draft not found'})
        else:
            # Create new draft
            title = f"Module Activity Sheet - {form_data['module_acronym']}" if form_data['module_acronym'] else "Module Activity Sheet - Untitled"
            draft = FormDraft(
                user_id=current_user.id,
                form_type='moduleactivity',
                title=title,
                module_acronym=form_data['module_acronym'],
                form_data=form_data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'draft_id': draft.id,
            'message': 'Draft saved automatically',
            'timestamp': datetime.utcnow().strftime('%I:%M:%S %p')
        })
        
    except Exception as e:
        print(f"Error in module activity autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/autosave-moduleguide-draft', methods=['POST'])
@login_required
def autosave_moduleguide_draft():
    """AJAX endpoint for autosaving module guide draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'module_acronym': data.get('module_acronym', ''),
            'teachertips_statement': data.get('teachertips_statement', ''),
            'standards': [],
            'vocab_terms': [],
            'careers': [],
            'sessions': [],
            'enrichment_activities': [],
            'locally_sourced_materials': [],
            'maintenance_items': [],
            'assembly_instructions': [],
            'recommended_websites': []
        }
        
        # Process standards
        standards_data = data.get('standards', [])
        for standard in standards_data:
            # Check if standard is a valid dictionary (not None)
            if standard and isinstance(standard, dict):
                if standard.get('standard'):
                    form_data['standards'].append({
                        'standard': standard['standard']
                    })
        
        # Process vocabulary terms
        vocab_data = data.get('vocab_terms', [])
        for term in vocab_data:
            # Check if term is a valid dictionary (not None)
            if term and isinstance(term, dict):
                if term.get('term'):
                    form_data['vocab_terms'].append({
                        'term': term['term']
                    })
        
        # Process careers
        careers_data = data.get('careers', [])
        for career in careers_data:
            # Check if career is a valid dictionary (not None)
            if career and isinstance(career, dict):
                if career.get('career'):
                    form_data['careers'].append({
                        'career': career['career']
                    })
        
        # Process sessions (7 sessions with complex structure)
        sessions_data = data.get('sessions', [])
        for session in sessions_data:
            # Check if session is a valid dictionary (not None)
            if session and isinstance(session, dict):
                if session.get('focus') or session.get('goals') or session.get('materials') or session.get('preparations') or session.get('assessments'):
                    session_data = {
                        'focus': session.get('focus', ''),
                        'goals': [goal for goal in session.get('goals', []) if goal],
                        'materials': [material for material in session.get('materials', []) if material],
                        'preparations': [prep for prep in session.get('preparations', []) if prep],
                        'assessments': [assessment for assessment in session.get('assessments', []) if assessment]
                    }
                    form_data['sessions'].append(session_data)
        
        # Process enrichment activities
        enrichment_data = data.get('enrichment_activities', [])
        for activity in enrichment_data:
            # Check if activity is a valid dictionary (not None)
            if activity and isinstance(activity, dict):
                if activity.get('activity'):
                    form_data['enrichment_activities'].append({
                        'activity': activity['activity']
                    })
        
        # Process locally sourced materials
        materials_data = data.get('locally_sourced_materials', [])
        for material in materials_data:
            # Check if material is a valid dictionary (not None)
            if material and isinstance(material, dict):
                if material.get('material'):
                    form_data['locally_sourced_materials'].append({
                        'material': material['material']
                    })
        
        # Process maintenance items
        maintenance_data = data.get('maintenance_items', [])
        for item in maintenance_data:
            # Check if item is a valid dictionary (not None)
            if item and isinstance(item, dict):
                if item.get('item'):
                    form_data['maintenance_items'].append({
                        'item': item['item']
                    })
        
        # Process assembly instructions
        assembly_data = data.get('assembly_instructions', [])
        for instruction in assembly_data:
            # Check if instruction is a valid dictionary (not None)
            if instruction and isinstance(instruction, dict):
                if instruction.get('instruction'):
                    form_data['assembly_instructions'].append({
                        'instruction': instruction['instruction']
                    })
        
        # Process recommended websites
        websites_data = data.get('recommended_websites', [])
        for website in websites_data:
            # Check if website is a valid dictionary (not None)
            if website and isinstance(website, dict):
                if website.get('title') or website.get('url'):
                    form_data['recommended_websites'].append({
                        'title': website.get('title', ''),
                        'url': website.get('url', '')
                    })
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='moduleguide').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                # Update title based on form data
                if form_data['module_acronym']:
                    draft.title = f"Module Guide - {form_data['module_acronym']}"
                    draft.module_acronym = form_data['module_acronym']
            else:
                return jsonify({'success': False, 'error': 'Draft not found'})
        else:
            # Create new draft
            title = f"Module Guide - {form_data['module_acronym']}" if form_data['module_acronym'] else "Module Guide - Untitled"
            draft = FormDraft(
                user_id=current_user.id,
                form_type='moduleguide',
                title=title,
                module_acronym=form_data['module_acronym'],
                form_data=form_data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'draft_id': draft.id,
            'message': 'Draft saved automatically',
            'timestamp': datetime.utcnow().strftime('%I:%M:%S %p')
        })
        
    except Exception as e:
        print(f"Error in module guide autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/load-moduleguide-draft/<int:draft_id>')
@login_required
def load_moduleguide_draft(draft_id):
    """Load module guide draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='moduleguide').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_module_guide'))
    
    try:
        # Create form and populate with draft data
        form = ModuleGuideForm()
        form_data = draft.form_data
        
        # Populate basic fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        form.teachertips_statement.data = form_data.get('teachertips_statement', '')
        
        # Populate standards
        standards_data = form_data.get('standards', [])
        for i, standard_data in enumerate(standards_data):
            if i < len(form.standards):
                form.standards[i].standard.data = standard_data.get('standard', '')
        
        # Populate vocabulary terms
        vocab_data = form_data.get('vocab_terms', [])
        for i, vocab_item in enumerate(vocab_data):
            if i < len(form.vocab_terms):
                form.vocab_terms[i].term.data = vocab_item.get('term', '')
        
        # Populate careers
        careers_data = form_data.get('careers', [])
        for i, career_data in enumerate(careers_data):
            if i < len(form.careers):
                form.careers[i].career.data = career_data.get('career', '')
        
        # Populate sessions (complex nested structure)
        sessions_data = form_data.get('sessions', [])
        for i, session_data in enumerate(sessions_data):
            if i < len(form.sessions):
                session_form = form.sessions[i]
                session_form.focus.data = session_data.get('focus', '')
                
                # Populate goals
                goals = session_data.get('goals', [])
                for j, goal in enumerate(goals):
                    if j < len(session_form.goals):
                        session_form.goals[j].goal.data = goal
                
                # Populate materials (15 material fields)
                materials = session_data.get('materials', [])
                for j, material in enumerate(materials):
                    if j < 15:  # 15 material fields
                        material_field = getattr(session_form, f'material{j+1}')
                        material_field.data = material
                
                # Populate preparations
                preparations = session_data.get('preparations', [])
                for j, prep in enumerate(preparations):
                    if j < len(session_form.preparations):
                        session_form.preparations[j].prep.data = prep
                
                # Populate assessments
                assessments = session_data.get('assessments', [])
                for j, assessment in enumerate(assessments):
                    if j < len(session_form.assessments):
                        session_form.assessments[j].assessment.data = assessment
        
        # Populate enrichment activities
        enrichment_data = form_data.get('enrichment_activities', [])
        for i, activity_data in enumerate(enrichment_data):
            if i < len(form.enrichment_activities):
                form.enrichment_activities[i].activity.data = activity_data.get('activity', '')
        
        # Populate locally sourced materials
        materials_data = form_data.get('locally_sourced_materials', [])
        for i, material_data in enumerate(materials_data):
            if i < len(form.locally_sourced_materials):
                form.locally_sourced_materials[i].material.data = material_data.get('material', '')
        
        # Populate maintenance items
        maintenance_data = form_data.get('maintenance_items', [])
        for i, item_data in enumerate(maintenance_data):
            if i < len(form.maintenance_items):
                form.maintenance_items[i].item.data = item_data.get('item', '')
        
        # Populate assembly instructions
        assembly_data = form_data.get('assembly_instructions', [])
        for i, instruction_data in enumerate(assembly_data):
            if i < len(form.assembly_instructions):
                form.assembly_instructions[i].instruction.data = instruction_data.get('instruction', '')
        
        # Populate recommended websites
        websites_data = form_data.get('recommended_websites', [])
        for i, website_data in enumerate(websites_data):
            if i < len(form.recommended_websites):
                form.recommended_websites[i].title.data = website_data.get('title', '')
                form.recommended_websites[i].url.data = website_data.get('url', '')
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_moduleGuide.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading module guide draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_module_guide'))

@app.route('/delete-moduleguide-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_moduleguide_draft(draft_id):
    """Delete module guide draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='moduleguide').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

@app.route('/autosave-moduleanswerkey-draft', methods=['POST'])
@login_required
def autosave_moduleanswerkey_draft():
    """AJAX endpoint for autosaving module answer key draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'module_acronym': data.get('module_acronym', ''),
            'pretest_questions': [],
            'rca_sessions': [],
            'posttest_questions': [],
            'pba_sessions': [],
            'vocabulary': [],
            'portfolio_checklist': [],
            'enrichment_dynamic_content': data.get('enrichment_dynamic_content', []),
            'worksheet_answer_keys': data.get('worksheet_answer_keys', [])
        }
        
        # Process pre-test questions
        pretest_data = data.get('pretest_questions', [])
        for question in pretest_data:
            # Check if question is a valid dictionary (not None)
            if question and isinstance(question, dict):
                if question.get('question_text') or question.get('choice_a') or question.get('choice_b') or question.get('choice_c') or question.get('choice_d'):
                    form_data['pretest_questions'].append({
                        'question_text': question.get('question_text', ''),
                        'choice_a': question.get('choice_a', ''),
                        'choice_b': question.get('choice_b', ''),
                        'choice_c': question.get('choice_c', ''),
                        'choice_d': question.get('choice_d', ''),
                        'correct_answer': question.get('correct_answer', '')
                    })
        
        # Process RCA sessions (4 sessions with 3 questions each)
        rca_data = data.get('rca_sessions', [])
        for session in rca_data:
            # Check if session is a valid dictionary (not None)
            if session and isinstance(session, dict):
                questions = []
                for question in session.get('questions', []):
                    # Check if question is a valid dictionary (not None)
                    if question and isinstance(question, dict):
                        if question.get('question_text') or question.get('choice_a'):
                            questions.append({
                                'question_text': question.get('question_text', ''),
                                'choice_a': question.get('choice_a', ''),
                                'choice_b': question.get('choice_b', ''),
                                'choice_c': question.get('choice_c', ''),
                                'choice_d': question.get('choice_d', ''),
                                'correct_answer': question.get('correct_answer', '')
                            })
                
                if questions:
                    form_data['rca_sessions'].append({
                        'session_number': session.get('session_number', ''),
                        'questions': questions
                    })
        
        # Process post-test questions
        posttest_data = data.get('posttest_questions', [])
        for question in posttest_data:
            # Check if question is a valid dictionary (not None)
            if question and isinstance(question, dict):
                if question.get('question_text') or question.get('choice_a') or question.get('choice_b') or question.get('choice_c') or question.get('choice_d'):
                    form_data['posttest_questions'].append({
                        'question_text': question.get('question_text', ''),
                        'choice_a': question.get('choice_a', ''),
                        'choice_b': question.get('choice_b', ''),
                        'choice_c': question.get('choice_c', ''),
                        'choice_d': question.get('choice_d', ''),
                        'correct_answer': question.get('correct_answer', '')
                    })
        
        # Process PBA sessions
        pba_data = data.get('pba_sessions', [])
        for session in pba_data:
            # Check if session is a valid dictionary (not None)
            if session and isinstance(session, dict):
                questions = []
                for question in session.get('assessment_questions', []):
                    # Check if question is a valid dictionary (not None)
                    if question and isinstance(question, dict):
                        if question.get('question') or question.get('correct_answer'):
                            questions.append({
                                'question': question.get('question', ''),
                                'correct_answer': question.get('correct_answer', '')
                            })
                
                if session.get('session_number') or session.get('activity_name') or questions:
                    form_data['pba_sessions'].append({
                        'session_number': session.get('session_number', ''),
                        'activity_name': session.get('activity_name', ''),
                        'assessment_questions': questions
                    })
        
        # Process vocabulary terms
        vocab_data = data.get('vocabulary', [])
        for term in vocab_data:
            # Check if term is a valid dictionary (not None)
            if term and isinstance(term, dict):
                if term.get('term') or term.get('definition'):
                    form_data['vocabulary'].append({
                        'term': term.get('term', ''),
                        'definition': term.get('definition', '')
                    })
        
        # Process portfolio checklist
        portfolio_data = data.get('portfolio_checklist', [])
        for item in portfolio_data:
            # Check if item is a valid dictionary (not None)
            if item and isinstance(item, dict):
                if item.get('product') or item.get('session_number'):
                    form_data['portfolio_checklist'].append({
                        'product': item.get('product', ''),
                        'session_number': item.get('session_number', '')
                    })
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='moduleanswerkey').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                # Update title based on form data
                if form_data['module_acronym']:
                    draft.title = f"Module Answer Key - {form_data['module_acronym']}"
                    draft.module_acronym = form_data['module_acronym']
            else:
                return jsonify({'success': False, 'error': 'Draft not found'})
        else:
            # Create new draft
            title = f"Module Answer Key - {form_data['module_acronym']}" if form_data['module_acronym'] else "Module Answer Key - Untitled"
            draft = FormDraft(
                user_id=current_user.id,
                form_type='moduleanswerkey',
                title=title,
                module_acronym=form_data['module_acronym'],
                form_data=form_data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'draft_id': draft.id,
            'message': 'Draft saved automatically',
            'timestamp': datetime.utcnow().strftime('%I:%M:%S %p')
        })
        
    except Exception as e:
        print(f"Error in module answer key autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/load-moduleanswerkey-draft/<int:draft_id>')
@login_required
def load_moduleanswerkey_draft(draft_id):
    """Load module answer key draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='moduleanswerkey').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_module_answer_key'))
    
    try:
        # Create form and populate with draft data
        form = ModuleAnswerKeyForm()
        form_data = draft.form_data
        
        # Populate basic fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        
        # Populate pre-test questions
        pretest_data = form_data.get('pretest_questions', [])
        for i, question_data in enumerate(pretest_data):
            if i < len(form.pretest_questions):
                form.pretest_questions[i].question_text.data = question_data.get('question_text', '')
                form.pretest_questions[i].choice_a.data = question_data.get('choice_a', '')
                form.pretest_questions[i].choice_b.data = question_data.get('choice_b', '')
                form.pretest_questions[i].choice_c.data = question_data.get('choice_c', '')
                form.pretest_questions[i].choice_d.data = question_data.get('choice_d', '')
                # Clean correct_answer to ensure it's valid
                correct_answer = question_data.get('correct_answer', '').strip().upper()
                form.pretest_questions[i].correct_answer.data = correct_answer if correct_answer in ['A', 'B', 'C', 'D'] else ''
        
        # Populate RCA sessions
        rca_data = form_data.get('rca_sessions', [])
        print(f"🔍 RCA sessions data: {len(rca_data)} sessions found")
        for i, session_data in enumerate(rca_data):
            print(f"🔍 RCA session {i+1}: {len(session_data.get('questions', []))} questions")
            if i < len(form.rca_sessions):
                questions = session_data.get('questions', [])
                for j, question_data in enumerate(questions):
                    print(f"🔍 RCA session {i+1}, question {j+1}: {question_data.get('question_text', '')[:50]}...")
                    if j < len(form.rca_sessions[i].questions):
                        form.rca_sessions[i].questions[j].question_text.data = question_data.get('question_text', '')
                        form.rca_sessions[i].questions[j].choice_a.data = question_data.get('choice_a', '')
                        form.rca_sessions[i].questions[j].choice_b.data = question_data.get('choice_b', '')
                        form.rca_sessions[i].questions[j].choice_c.data = question_data.get('choice_c', '')
                        form.rca_sessions[i].questions[j].choice_d.data = question_data.get('choice_d', '')
                        # Clean correct_answer to ensure it's valid
                        correct_answer = question_data.get('correct_answer', '').strip().upper()
                        form.rca_sessions[i].questions[j].correct_answer.data = correct_answer if correct_answer in ['A', 'B', 'C', 'D'] else ''
        
        # Populate post-test questions
        posttest_data = form_data.get('posttest_questions', [])
        for i, question_data in enumerate(posttest_data):
            if i < len(form.posttest_questions):
                form.posttest_questions[i].question_text.data = question_data.get('question_text', '')
                form.posttest_questions[i].choice_a.data = question_data.get('choice_a', '')
                form.posttest_questions[i].choice_b.data = question_data.get('choice_b', '')
                form.posttest_questions[i].choice_c.data = question_data.get('choice_c', '')
                form.posttest_questions[i].choice_d.data = question_data.get('choice_d', '')
                # Clean correct_answer to ensure it's valid
                correct_answer = question_data.get('correct_answer', '').strip().upper()
                form.posttest_questions[i].correct_answer.data = correct_answer if correct_answer in ['A', 'B', 'C', 'D'] else ''
        
        # Populate PBA sessions
        pba_data = form_data.get('pba_sessions', [])
        print(f"🔍 PBA sessions data: {len(pba_data)} sessions found")
        for i, session_data in enumerate(pba_data):
            print(f"🔍 PBA session {i+1}: {len(session_data.get('assessment_questions', []))} questions")
            if i < len(form.pba_sessions):
                form.pba_sessions[i].session_number.data = session_data.get('session_number', '')
                form.pba_sessions[i].activity_name.data = session_data.get('activity_name', '')
                
                questions = session_data.get('assessment_questions', [])
                for j, question_data in enumerate(questions):
                    print(f"🔍 PBA session {i+1}, question {j+1}: {question_data.get('question', '')[:50]}...")
                    if j < len(form.pba_sessions[i].assessment_questions):
                        form.pba_sessions[i].assessment_questions[j].question.data = question_data.get('question', '')
                        form.pba_sessions[i].assessment_questions[j].correct_answer.data = question_data.get('correct_answer', '')
        
        # Populate vocabulary
        vocab_data = form_data.get('vocabulary', [])
        for i, term_data in enumerate(vocab_data):
            if i < len(form.vocabulary):
                form.vocabulary[i].term.data = term_data.get('term', '')
                form.vocabulary[i].definition.data = term_data.get('definition', '')
        
        # Populate portfolio checklist
        portfolio_data = form_data.get('portfolio_checklist', [])
        print(f"🔍 Portfolio checklist data: {len(portfolio_data)} items found")
        print(f"🔍 Raw portfolio data: {portfolio_data}")
        print(f"🔍 Form portfolio_checklist length: {len(form.portfolio_checklist)}")
        
        for i, item_data in enumerate(portfolio_data):
            print(f"🔍 Portfolio item {i+1}: {item_data}")
            if i < len(form.portfolio_checklist):
                print(f"🔍 Setting portfolio_checklist[{i}].product = '{item_data.get('product', '')}'")
                print(f"🔍 Setting portfolio_checklist[{i}].session_number = '{item_data.get('session_number', '')}'")
                form.portfolio_checklist[i].product.data = item_data.get('product', '')
                form.portfolio_checklist[i].session_number.data = item_data.get('session_number', '')
            else:
                print(f"🔍 Skipping item {i+1} - index {i} >= form length {len(form.portfolio_checklist)}")
        
        # Pass dynamic fields data to template for JavaScript reconstruction
        enrichment_data = form_data.get('enrichment_dynamic_content', [])
        worksheet_data = form_data.get('worksheet_answer_keys', [])
        
        # Debug logging
        print(f"🔍 Module Answer Key draft {draft.id} data keys: {list(form_data.keys())}")
        print(f"🔍 Enrichment data: {enrichment_data}")
        print(f"🔍 Worksheet data: {worksheet_data}")
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_moduleAnswerKey.html', form=form, draft_id=draft.id, 
                             draft_data=form_data, enrichment_dynamic_content=enrichment_data, 
                             worksheet_answer_keys=worksheet_data, portfolio_checklist_data=portfolio_data)
        
    except Exception as e:
        print(f"Error loading module answer key draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_module_answer_key'))

@app.route('/delete-moduleanswerkey-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_moduleanswerkey_draft(draft_id):
    """Delete module answer key draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='moduleanswerkey').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

# ===== MODULE ANSWER KEY 2.0 DRAFT MANAGEMENT ROUTES =====

@app.route('/autosave-module-answer-key2-draft', methods=['POST'])
@login_required
def autosave_module_answer_key2_draft():
    """AJAX endpoint for autosaving Module Answer Key 2.0 draft with enhanced reliability"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage with explicit validation
        form_data = {
            'module_acronym': data.get('module_acronym', ''),
            'pretest_questions': [],
            'rca_sessions': [],
            'posttest_questions': [],
            'pba_sessions': [],
            'vocabulary': [],
            'portfolio_checklist': []
        }
        
        # Process pre-test questions with index validation
        pretest_questions = data.get('pretest_questions', [])
        for i, question in enumerate(pretest_questions):
            if question and any(question.get(key, '') for key in ['question_text', 'choice_a', 'choice_b', 'choice_c', 'choice_d', 'correct_answer']):
                form_data['pretest_questions'].append({
                    'question_text': question.get('question_text', ''),
                    'choice_a': question.get('choice_a', ''),
                    'choice_b': question.get('choice_b', ''),
                    'choice_c': question.get('choice_c', ''),
                    'choice_d': question.get('choice_d', ''),
                    'correct_answer': question.get('correct_answer', '').upper() if question.get('correct_answer') else ''
                })
        
        # Process RCA sessions with nested question validation
        rca_sessions = data.get('rca_sessions', [])
        for session in rca_sessions:
            if session:
                questions = []
                for question in session.get('questions', []):
                    if question and any(question.get(key, '') for key in ['question_text', 'choice_a', 'choice_b', 'choice_c', 'choice_d', 'correct_answer']):
                        questions.append({
                            'question_text': question.get('question_text', ''),
                            'choice_a': question.get('choice_a', ''),
                            'choice_b': question.get('choice_b', ''),
                            'choice_c': question.get('choice_c', ''),
                            'choice_d': question.get('choice_d', ''),
                            'correct_answer': question.get('correct_answer', '').upper() if question.get('correct_answer') else ''
                        })
                if questions:  # Only add session if it has questions
                    form_data['rca_sessions'].append({'questions': questions})
        
        # Process post-test questions with index validation
        posttest_questions = data.get('posttest_questions', [])
        for i, question in enumerate(posttest_questions):
            if question and any(question.get(key, '') for key in ['question_text', 'choice_a', 'choice_b', 'choice_c', 'choice_d', 'correct_answer']):
                form_data['posttest_questions'].append({
                    'question_text': question.get('question_text', ''),
                    'choice_a': question.get('choice_a', ''),
                    'choice_b': question.get('choice_b', ''),
                    'choice_c': question.get('choice_c', ''),
                    'choice_d': question.get('choice_d', ''),
                    'correct_answer': question.get('correct_answer', '').upper() if question.get('correct_answer') else ''
                })
        
        # Process PBA sessions with validation
        pba_sessions = data.get('pba_sessions', [])
        for session in pba_sessions:
            if session and (session.get('session_number') or session.get('activity_name') or session.get('assessment_questions')):
                assessment_questions = []
                for question in session.get('assessment_questions', []):
                    if question and (question.get('question') or question.get('correct_answer')):
                        assessment_questions.append({
                            'question': question.get('question', ''),
                            'correct_answer': question.get('correct_answer', '')
                        })
                form_data['pba_sessions'].append({
                    'session_number': session.get('session_number', ''),
                    'activity_name': session.get('activity_name', ''),
                    'assessment_questions': assessment_questions
                })
        
        # Process vocabulary with validation
        vocabulary = data.get('vocabulary', [])
        for term in vocabulary:
            if term and (term.get('term') or term.get('definition')):
                form_data['vocabulary'].append({
                    'term': term.get('term', ''),
                    'definition': term.get('definition', '')
                })
        
        # Process portfolio checklist with validation
        portfolio_checklist = data.get('portfolio_checklist', [])
        for item in portfolio_checklist:
            if item and (item.get('product') or item.get('session_number')):
                form_data['portfolio_checklist'].append({
                    'product': item.get('product', ''),
                    'session_number': item.get('session_number', '')
                })
        
        # Create title from module acronym
        module_acronym = form_data.get('module_acronym', 'Untitled')
        title = f"Module Answer Key 2.0 - {module_acronym}" if module_acronym else "Module Answer Key 2.0 - Untitled"
        
        # Check if user already has a draft for this form type
        existing_draft = FormDraft.query.filter_by(
            user_id=current_user.id, 
            form_type='module_answer_key2',
            is_current=True
        ).first()
        
        if existing_draft:
            # Update existing draft
            existing_draft.form_data = form_data
            existing_draft.title = title
            existing_draft.module_acronym = module_acronym
            existing_draft.updated_at = datetime.utcnow()
            print(f"🔍 Updated existing Module Answer Key 2.0 draft {existing_draft.id}")
        else:
            # Create new draft
            new_draft = FormDraft(
                user_id=current_user.id,
                form_type='module_answer_key2',
                title=title,
                module_acronym=module_acronym,
                form_data=form_data,
                is_current=True
            )
            db.session.add(new_draft)
            print(f"🔍 Created new Module Answer Key 2.0 draft")
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Draft saved successfully'})
        
    except Exception as e:
        print(f"🔍 Error autosaving Module Answer Key 2.0 draft: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})

@app.route('/load-module-answer-key2-draft/<int:draft_id>')
@login_required
def load_module_answer_key2_draft(draft_id):
    """Load Module Answer Key 2.0 draft"""
    try:
        # Get the draft
        draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='module_answer_key2').first()
        
        if not draft:
            flash('Draft not found', 'error')
            return redirect(url_for('create_module_answer_key2'))
        
        # Create form and load data
        form = ModuleAnswerKey2Form()
        form_data = draft.form_data
        
        # Load data into form with enhanced error handling
        success = load_module_answer_key2_draft_into_form(form, current_user.id)
        
        if not success:
            flash('Error loading draft data', 'error')
            return redirect(url_for('create_module_answer_key2'))
        
        print(f"🔍 Module Answer Key 2.0 draft {draft.id} loaded successfully")
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_module_answer_key2.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading Module Answer Key 2.0 draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_module_answer_key2'))

@app.route('/delete-module-answer-key2-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_module_answer_key2_draft(draft_id):
    """Delete Module Answer Key 2.0 draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='module_answer_key2').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Module Answer Key 2.0 draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

@app.route('/autosave-familybriefing-draft', methods=['POST'])
@login_required
def autosave_familybriefing_draft():
    """AJAX endpoint for autosaving family briefing draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'module_name': data.get('module_name', ''),
            'introsentence': data.get('introsentence', ''),
            'learningobjective1': data.get('learningobjective1', ''),
            'learningobjective2': data.get('learningobjective2', ''),
            'learningobjective3': data.get('learningobjective3', ''),
            'learningobjective4': data.get('learningobjective4', ''),
            'learningobjective5': data.get('learningobjective5', ''),
            'learningobjective6': data.get('learningobjective6', ''),
            'activityname1': data.get('activityname1', ''),
            'activityname2': data.get('activityname2', ''),
            'activityname3': data.get('activityname3', ''),
            'activityname4': data.get('activityname4', ''),
            'activityname5': data.get('activityname5', ''),
            'activityname6': data.get('activityname6', ''),
            'activityname7': data.get('activityname7', ''),
            'term1': data.get('term1', ''),
            'term2': data.get('term2', ''),
            'term3': data.get('term3', ''),
            'term4': data.get('term4', ''),
            'term5': data.get('term5', ''),
            'term6': data.get('term6', ''),
            'term7': data.get('term7', ''),
            'term8': data.get('term8', ''),
            'term9': data.get('term9', ''),
            'term10': data.get('term10', ''),
            'term11': data.get('term11', ''),
            'term12': data.get('term12', ''),
            'term13': data.get('term13', ''),
            'term14': data.get('term14', ''),
            'term15': data.get('term15', ''),
            'term16': data.get('term16', ''),
            'term17': data.get('term17', ''),
            'term18': data.get('term18', ''),
            'term19': data.get('term19', ''),
            'term20': data.get('term20', ''),
            'term21': data.get('term21', ''),
            'keyconcept1_name': data.get('keyconcept1_name', ''),
            'keyconcept1_explanation': data.get('keyconcept1_explanation', ''),
            'keyconcept2_name': data.get('keyconcept2_name', ''),
            'keyconcept2_explanation': data.get('keyconcept2_explanation', ''),
            'keyconcept3_name': data.get('keyconcept3_name', ''),
            'keyconcept3_explanation': data.get('keyconcept3_explanation', ''),
            'keyconcept4_name': data.get('keyconcept4_name', ''),
            'keyconcept4_explanation': data.get('keyconcept4_explanation', ''),
            'keyconcept5_name': data.get('keyconcept5_name', ''),
            'keyconcept5_explanation': data.get('keyconcept5_explanation', '')
        }
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='familybriefing').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                # Update title based on form data
                if form_data['module_name']:
                    draft.title = f"Family Briefing - {form_data['module_name']}"
                    draft.module_acronym = form_data['module_name']
            else:
                return jsonify({'success': False, 'error': 'Draft not found'})
        else:
            # Create new draft
            title = f"Family Briefing - {form_data['module_name']}" if form_data['module_name'] else "Family Briefing - Untitled"
            draft = FormDraft(
                user_id=current_user.id,
                form_type='familybriefing',
                title=title,
                module_acronym=form_data['module_name'],
                form_data=form_data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'draft_id': draft.id,
            'message': 'Draft saved automatically',
            'timestamp': datetime.utcnow().strftime('%I:%M:%S %p')
        })
        
    except Exception as e:
        print(f"Error in family briefing autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/autosave-rca-draft', methods=['POST'])
@login_required
def autosave_rca_draft():
    """AJAX endpoint for autosaving RCA worksheet draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'module_acronym': data.get('module_acronym', ''),
            'session_number': data.get('session_number', ''),
            'questions': []
        }
        
        # Process RCA questions (Research, Challenge, Application)
        questions_data = data.get('questions', [])
        for question_data in questions_data:
            question_text = question_data.get('question_text', '').strip()
            choice_a = question_data.get('choice_a', '').strip()
            choice_b = question_data.get('choice_b', '').strip()
            choice_c = question_data.get('choice_c', '').strip()
            choice_d = question_data.get('choice_d', '').strip()
            
            # Only include questions with some content
            if question_text or choice_a or choice_b or choice_c or choice_d:
                form_data['questions'].append({
                    'question_text': question_text,
                    'choice_a': choice_a,
                    'choice_b': choice_b,
                    'choice_c': choice_c,
                    'choice_d': choice_d
                })
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='rca').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                # Update title based on form data
                if form_data['module_acronym'] or form_data['session_number']:
                    draft.title = f"RCA Worksheet - {form_data['module_acronym']} Session {form_data['session_number']}"
                    draft.module_acronym = form_data['module_acronym']
            else:
                return jsonify({'success': False, 'error': 'Draft not found'})
        else:
            # Create new draft
            title_parts = []
            if form_data['module_acronym']:
                title_parts.append(form_data['module_acronym'])
            if form_data['session_number']:
                title_parts.append(f"Session {form_data['session_number']}")
            
            title = f"RCA Worksheet - {' '.join(title_parts)}" if title_parts else "RCA Worksheet - Untitled"
            
            draft = FormDraft(
                user_id=current_user.id,
                form_type='rca',
                title=title,
                module_acronym=form_data['module_acronym'],
                form_data=form_data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'draft_id': draft.id,
            'message': 'Draft saved automatically',
            'timestamp': datetime.utcnow().strftime('%I:%M:%S %p')
        })
        
    except Exception as e:
        print(f"Error in RCA autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

# Missing Load/Delete Draft Routes - CRITICAL FIX
@app.route('/load-familybriefing-draft/<int:draft_id>')
@login_required
def load_familybriefing_draft(draft_id):
    """Load family briefing draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='familybriefing').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_familybriefing'))
    
    try:
        # Create form and populate with draft data
        form = FamilyBriefingForm()
        form_data = draft.form_data
        
        # Populate all form fields
        form.module_name.data = form_data.get('module_name', '')
        form.introsentence.data = form_data.get('introsentence', '')
        form.learningobjective1.data = form_data.get('learningobjective1', '')
        form.learningobjective2.data = form_data.get('learningobjective2', '')
        form.learningobjective3.data = form_data.get('learningobjective3', '')
        form.learningobjective4.data = form_data.get('learningobjective4', '')
        form.learningobjective5.data = form_data.get('learningobjective5', '')
        form.learningobjective6.data = form_data.get('learningobjective6', '')
        form.activityname1.data = form_data.get('activityname1', '')
        form.activityname2.data = form_data.get('activityname2', '')
        form.activityname3.data = form_data.get('activityname3', '')
        form.activityname4.data = form_data.get('activityname4', '')
        form.activityname5.data = form_data.get('activityname5', '')
        form.activityname6.data = form_data.get('activityname6', '')
        form.activityname7.data = form_data.get('activityname7', '')
        form.term1.data = form_data.get('term1', '')
        form.term2.data = form_data.get('term2', '')
        form.term3.data = form_data.get('term3', '')
        form.term4.data = form_data.get('term4', '')
        form.term5.data = form_data.get('term5', '')
        form.term6.data = form_data.get('term6', '')
        form.term7.data = form_data.get('term7', '')
        form.term8.data = form_data.get('term8', '')
        form.term9.data = form_data.get('term9', '')
        form.term10.data = form_data.get('term10', '')
        form.term11.data = form_data.get('term11', '')
        form.term12.data = form_data.get('term12', '')
        form.term13.data = form_data.get('term13', '')
        form.term14.data = form_data.get('term14', '')
        form.term15.data = form_data.get('term15', '')
        form.term16.data = form_data.get('term16', '')
        form.term17.data = form_data.get('term17', '')
        form.term18.data = form_data.get('term18', '')
        form.term19.data = form_data.get('term19', '')
        form.term20.data = form_data.get('term20', '')
        form.term21.data = form_data.get('term21', '')
        form.keyconcept1_name.data = form_data.get('keyconcept1_name', '')
        form.keyconcept1_explanation.data = form_data.get('keyconcept1_explanation', '')
        form.keyconcept2_name.data = form_data.get('keyconcept2_name', '')
        form.keyconcept2_explanation.data = form_data.get('keyconcept2_explanation', '')
        form.keyconcept3_name.data = form_data.get('keyconcept3_name', '')
        form.keyconcept3_explanation.data = form_data.get('keyconcept3_explanation', '')
        form.keyconcept4_name.data = form_data.get('keyconcept4_name', '')
        form.keyconcept4_explanation.data = form_data.get('keyconcept4_explanation', '')
        form.keyconcept5_name.data = form_data.get('keyconcept5_name', '')
        form.keyconcept5_explanation.data = form_data.get('keyconcept5_explanation', '')
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_familybriefing.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading family briefing draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_familybriefing'))

@app.route('/delete-familybriefing-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_familybriefing_draft(draft_id):
    """Delete family briefing draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='familybriefing').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

# ===== STUDENT LOGBOOK ROUTES =====

@app.route('/create-studentmoduleworkbook', methods=['GET', 'POST'])
@login_required
def create_studentmoduleworkbook():
    form = StudentModuleWorkbookForm()
    
    if request.method == 'POST':
        print("Student Logbook form submitted!")
        print(f"Form data: {request.form}")
        print(f"Form valid: {form.validate_on_submit()}")
        if form.errors:
            print(f"Form errors: {form.errors}")
    
    if form.validate_on_submit():
        print("Student Logbook form validation passed!")
        # Generate the document
        try:
            print("Attempting to generate student logbook...")
            doc_path = generate_student_module_workbook(form)
            filename = os.path.basename(doc_path)
            
            print(f"Student logbook generated at: {doc_path}")
            
            # Save document info to database
            doc_record = GeneratedDocument(
                user_id=current_user.id,
                document_type='studentmoduleworkbook',
                filename=filename,
                file_path=doc_path,
                module_acronym=form.module_name.data or 'Student Workbook',  # Use the module name
                file_size=os.path.getsize(doc_path)
            )
            db.session.add(doc_record)
            db.session.commit()
            
            flash('Student Logbook generated successfully!', 'success')
            return redirect(url_for('my_documents'))
        except Exception as e:
            print(f"Error generating student logbook: {e}")
            flash(f'Error generating document: {str(e)}', 'error')
    
    return render_template('create_studentmoduleworkbook.html', form=form)

@app.route('/autosave-studentmoduleworkbook-draft', methods=['POST'])
@login_required
def autosave_studentmoduleworkbook_draft():
    """AJAX endpoint for autosaving student logbook draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'module_name': data.get('module_name', ''),
        }
        
        # Session focus
        for i in range(1, 8):
            form_data[f'focus_s{i}'] = data.get(f'focus_s{i}', '')
        
        # Session goals, vocabulary, and assessments
        for s in range(1, 8):
            for g in range(1, 4):
                form_data[f's{s}_goal{g}'] = data.get(f's{s}_goal{g}', '')
            for v in range(1, 6):
                form_data[f's{s}_vocab{v}'] = data.get(f's{s}_vocab{v}', '')
            for a in range(1, 5):
                form_data[f's{s}_assessment{a}'] = data.get(f's{s}_assessment{a}', '')
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='studentmoduleworkbook').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                db.session.commit()
                
                timestamp = draft.updated_at.strftime('%I:%M %p')
                return jsonify({
                    'success': True,
                    'draft_id': draft.id,
                    'timestamp': timestamp
                })
        
        # Create new draft
        # Generate title from first non-empty focus
        title = 'Student Logbook Draft'
        for i in range(1, 8):
            focus_text = data.get(f'focus_s{i}', '').strip()
            if focus_text:
                title = f'Student Logbook - Session {i}: {focus_text[:30]}...' if len(focus_text) > 30 else f'Student Logbook - Session {i}: {focus_text}'
                break
        
        draft = FormDraft(
            user_id=current_user.id,
            form_type='studentmoduleworkbook',
            title=title,
            form_data=form_data
        )
        db.session.add(draft)
        db.session.commit()
        
        timestamp = draft.created_at.strftime('%I:%M %p')
        return jsonify({
            'success': True,
            'draft_id': draft.id,
            'timestamp': timestamp
        })
        
    except Exception as e:
        print(f"Error in autosave_studentmoduleworkbook_draft: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/load-studentmoduleworkbook-draft/<int:draft_id>')
@login_required
def load_studentmoduleworkbook_draft(draft_id):
    """Load student logbook draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='studentmoduleworkbook').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_studentmoduleworkbook'))
    
    try:
        # Create form and populate with draft data
        form = StudentModuleWorkbookForm()
        form_data = draft.form_data
        
        # Populate module information
        form.module_name.data = form_data.get('module_name', '')
        
        # Populate session focus fields
        for i in range(1, 8):
            getattr(form, f'focus_s{i}').data = form_data.get(f'focus_s{i}', '')
        
        # Populate session goals, vocabulary, and assessments
        for s in range(1, 8):
            for g in range(1, 4):
                getattr(form, f's{s}_goal{g}').data = form_data.get(f's{s}_goal{g}', '')
            for v in range(1, 6):
                getattr(form, f's{s}_vocab{v}').data = form_data.get(f's{s}_vocab{v}', '')
            for a in range(1, 5):
                getattr(form, f's{s}_assessment{a}').data = form_data.get(f's{s}_assessment{a}', '')
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_studentmoduleworkbook.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading student logbook draft: {e}")
        flash('Error loading draft', 'error')
        return redirect(url_for('create_studentmoduleworkbook'))

@app.route('/delete-studentmoduleworkbook-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_studentmoduleworkbook_draft(draft_id):
    """Delete student logbook draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='studentmoduleworkbook').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

@app.route('/load-rca-draft/<int:draft_id>')
@login_required
def load_rca_draft(draft_id):
    """Load RCA worksheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='rca').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_rca'))
    
    try:
        # Create form and populate with draft data
        form = RCAWorksheetForm()
        form_data = draft.form_data
        
        # Populate form fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        form.session_number.data = form_data.get('session_number', '')
        
        # Populate questions
        questions_data = form_data.get('questions', [])
        for i, question_data in enumerate(questions_data):
            if i < len(form.questions):
                form.questions[i].question_text.data = question_data.get('question_text', '')
                form.questions[i].choice_a.data = question_data.get('choice_a', '')
                form.questions[i].choice_b.data = question_data.get('choice_b', '')
                form.questions[i].choice_c.data = question_data.get('choice_c', '')
                form.questions[i].choice_d.data = question_data.get('choice_d', '')
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_rca.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading RCA draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_rca'))

@app.route('/delete-rca-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_rca_draft(draft_id):
    """Delete RCA worksheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='rca').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

@app.route('/load-generic-draft/<int:draft_id>')
@login_required
def load_generic_draft(draft_id):
    """Load generic worksheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='generic').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_generic'))
    
    try:
        # Create form and populate with draft data
        form = GenericWorksheetForm()
        form_data = draft.form_data
        
        # Populate basic fields
        form.module_acronym.data = form_data.get('module_acronym', '')
        form.worksheet_title.data = form_data.get('worksheet_title', '')
        
        # Pass dynamic fields data to template for JavaScript reconstruction
        dynamic_fields_data = form_data.get('dynamic_fields', [])
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_generic.html', form=form, draft_id=draft.id, 
                             draft_data=form_data, dynamic_fields_data=dynamic_fields_data)
        
    except Exception as e:
        print(f"Error loading generic draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_generic'))

@app.route('/delete-generic-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_generic_draft(draft_id):
    """Delete generic worksheet draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='generic').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

# ===== HORIZONTAL LESSON PLAN DRAFT ROUTES =====

@app.route('/autosave-horizontal-lesson-plan-draft', methods=['POST'])
@login_required
def autosave_horizontal_lesson_plan_draft():
    """AJAX endpoint for autosaving horizontal lesson plan draft"""
    try:
        # Get JSON data from the request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Prepare form data for JSON storage
        form_data = {
            'school_name': data.get('school_name', ''),
            'teacher_name': data.get('teacher_name', ''),
            'term': data.get('term', ''),
            'modules': data.get('modules', [])
        }
        
        # Check if this is updating an existing draft
        draft_id = data.get('draft_id')
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='horizontal_lesson_plan').first()
            if draft:
                draft.form_data = form_data
                draft.updated_at = datetime.utcnow()
                # Update title if school name changed
                if form_data['school_name']:
                    draft.title = f"Horizontal Lesson Plan - {form_data['school_name']}"
            else:
                return jsonify({'success': False, 'error': 'Draft not found'})
        else:
            # Create new draft
            title = f"Horizontal Lesson Plan - {form_data['school_name']}" if form_data['school_name'] else "Horizontal Lesson Plan - Untitled"
            draft = FormDraft(
                user_id=current_user.id,
                form_type='horizontal_lesson_plan',
                title=title,
                form_data=form_data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'draft_id': draft.id,
            'message': 'Draft saved automatically',
            'timestamp': datetime.utcnow().strftime('%I:%M:%S %p')
        })
        
    except Exception as e:
        print(f"Error in autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/load-horizontal-lesson-plan-draft/<int:draft_id>')
@login_required
def load_horizontal_lesson_plan_draft(draft_id):
    """Load horizontal lesson plan draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='horizontal_lesson_plan').first()
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('create_horizontal_lesson_plan'))
    
    try:
        # Create form and populate with draft data
        form = HorizontalLessonPlanForm()
        form_data = draft.form_data
        
        # Populate basic fields
        form.school_name.data = form_data.get('school_name', '')
        form.teacher_name.data = form_data.get('teacher_name', '')
        form.term.data = form_data.get('term', '')
        
        # Populate modules
        modules_data = form_data.get('modules', [])
        for i, module_data in enumerate(modules_data):
            if i < len(form.modules):
                form.modules[i].module_name.data = module_data.get('module_name', '')
                form.modules[i].enrichment_activities.data = module_data.get('enrichment_activities', '')
                
                # Populate sessions
                sessions_data = module_data.get('sessions', [])
                for j, session_data in enumerate(sessions_data):
                    if j < len(form.modules[i].sessions):
                        form.modules[i].sessions[j].focus.data = session_data.get('focus', '')
                        form.modules[i].sessions[j].objectives.data = session_data.get('objectives', '')
                        form.modules[i].sessions[j].materials.data = session_data.get('materials', '')
                        form.modules[i].sessions[j].teacher_prep.data = session_data.get('teacher_prep', '')
                        form.modules[i].sessions[j].assessments.data = session_data.get('assessments', '')
        
        flash(f'Draft "{draft.title}" loaded successfully!', 'success')
        return render_template('create_horizontal_lesson_plan.html', form=form, draft_id=draft.id)
        
    except Exception as e:
        print(f"Error loading horizontal lesson plan draft: {e}")
        flash(f'Error loading draft: {str(e)}', 'error')
        return redirect(url_for('create_horizontal_lesson_plan'))

@app.route('/delete-horizontal-lesson-plan-draft/<int:draft_id>', methods=['POST'])
@login_required
def delete_horizontal_lesson_plan_draft(draft_id):
    """Delete horizontal lesson plan draft"""
    draft = FormDraft.query.filter_by(id=draft_id, user_id=current_user.id, form_type='horizontal_lesson_plan').first()
    
    if not draft:
        flash('Draft not found', 'error')
    else:
        db.session.delete(draft)
        db.session.commit()
        flash('Draft deleted successfully!', 'success')
    
    return redirect(url_for('drafts'))

@app.route('/create-curriculum-design-build', methods=['GET', 'POST'])
@login_required
def create_curriculum_design_build():
    form = CurriculumDesignBuildForm()
    
    if request.method == 'POST' and form.validate_on_submit():
        try:
            # Generate the document
            doc_path = generate_curriculum_design_build_document(form)
            
            # Store in database
            doc_record = GeneratedDocument(
                user_id=current_user.id,
                document_type='curriculum_design_build',
                filename=os.path.basename(doc_path),
                file_path=doc_path,
                file_size=os.path.getsize(doc_path)
            )
            db.session.add(doc_record)
            db.session.commit()
            
            flash('Curriculum Design Build document generated successfully!', 'success')
            return redirect(url_for('my_documents'))
            
        except Exception as e:
            print(f"Error generating Curriculum Design Build document: {e}")
            import traceback
            traceback.print_exc()
            flash(f'Error generating document: {str(e)}', 'error')
    
    # Load existing draft if editing
    draft = None
    draft_id = request.args.get('draft_id')
    if draft_id:
        try:
            draft = FormDraft.query.get(draft_id)
            if draft and draft.user_id == current_user.id:
                load_curriculum_design_build_draft_into_form(form, draft)
        except Exception as e:
            print(f"Error loading draft: {e}")
    
    # Pass draft data to template if loading a draft
    draft_data = None
    if draft:
        draft_data = draft.form_data
    
    return render_template('create_curriculum_design_build.html', form=form, draft_id=draft_id, draft_data=draft_data)

@app.route('/autosave-curriculum-design-build-draft', methods=['POST'])
@login_required
def autosave_curriculum_design_build_draft():
    try:
        data = request.get_json()
        draft_id = data.get('draft_id')
        
        if draft_id:
            # Update existing draft
            draft = FormDraft.query.get(draft_id)
            if draft and draft.user_id == current_user.id:
                draft.form_data = data
                draft.updated_at = datetime.utcnow()
        else:
            # Create new draft
            title = f"{data.get('site_name', 'Untitled')} - Curriculum Design Build"
            draft = FormDraft(
                user_id=current_user.id,
                form_type='curriculum_design_build',
                title=title,
                form_data=data
            )
            db.session.add(draft)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'draft_id': draft.id,
            'timestamp': datetime.now().strftime('%I:%M %p')
        })
    
    except Exception as e:
        print(f"Error in autosave: {e}")
        return jsonify({'success': False, 'error': str(e)})

def load_curriculum_design_build_draft_into_form(form, draft):
    """Load draft data into the curriculum design build form"""
    try:
        form_data = draft.form_data
        
        # Basic fields
        form.star_academy_model.data = form_data.get('star_academy_model', '')
        form.grade_level.data = form_data.get('grade_level', '')
        form.soft_start.data = form_data.get('soft_start', False)
        form.site_name.data = form_data.get('site_name', '')
        form.school_district.data = form_data.get('school_district', '')
        form.science_course_grades.data = form_data.get('science_course_grades', '')
        form.math_course_grades.data = form_data.get('math_course_grades', '')
        form.ss_course_grades.data = form_data.get('ss_course_grades', '')
        form.ela_course_grades.data = form_data.get('ela_course_grades', '')
        form.science_rotations.data = form_data.get('science_rotations', '')
        form.math_rotations.data = form_data.get('math_rotations', '')
        form.tier_one_component.data = form_data.get('tier_one_component', '')
        form.science_design_domain.data = form_data.get('science_design_domain', '')
        form.state_math_domains.data = form_data.get('state_math_domains', '')
        form.ipls_additional_coverage.data = form_data.get('ipls_additional_coverage', '')
        form.ipls_critical_standards.data = form_data.get('ipls_critical_standards', '')
        form.tci_program_title.data = form_data.get('tci_program_title', '')
        form.ss_course_title.data = form_data.get('ss_course_title', '')
        
        
        
        
        print(f"Successfully loaded curriculum design build draft data")
        
    except Exception as e:
        print(f"Error loading draft data: {e}")

def generate_curriculum_design_build_document(form):
    """Generate curriculum design build document with dynamic tables"""
    master_template_path = 'templates/docx_templates/curriculum_design_build_master.docx'
    working_template_path = 'templates/docx_templates/curriculum_design_build.docx'
    
    # Copy master to working template
    shutil.copy2(master_template_path, working_template_path)
    
    # Create temporary file for processing
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(working_template_path, temp_file.name)
        
        doc = DocxTemplate(temp_file.name)
        
        # Prepare context data - build step by step to avoid conflicts
        context = {}
        
        # Star Academy Details
        context['star'] = {'academy': {'model': escape_xml(form.star_academy_model.data)}}
        context['grade'] = {'level': escape_xml(form.grade_level.data)}
        context['soft'] = {'start': form.soft_start.data}
        context['site'] = {'name': escape_xml(form.site_name.data)}
        context['school'] = {'district': escape_xml(form.school_district.data)}
        
        # Essential Program Elements - separate nested structure
        context['grade_level_courses'] = {
            'science': escape_xml(form.science_course_grades.data),
            'math': escape_xml(form.math_course_grades.data),
            'ss': escape_xml(form.ss_course_grades.data),
            'ela': escape_xml(form.ela_course_grades.data)
        }
        
        # Generate dynamic table subdocuments
        print("Generating dynamic table subdocuments...")
        science_table_subdoc = generate_dynamic_table_subdocument(doc, request.form, 'science')
        math_table_subdoc = generate_dynamic_table_subdocument(doc, request.form, 'math')
        social_studies_table_subdoc = generate_dynamic_table_subdocument(doc, request.form, 'social_studies')
        print(f"Dynamic subdocs created - Science: {type(science_table_subdoc)}, Math: {type(math_table_subdoc)}, SS: {type(social_studies_table_subdoc)}")
        
        # Curriculum Elements - build science dict carefully
        context['science'] = {
            'rotations': escape_xml(form.science_rotations.data),
            'design': {'domain': escape_xml(form.science_design_domain.data)},
            'table': {'content': science_table_subdoc}
        }
        
        # Math elements - build math dict carefully  
        context['math'] = {
            'rotations': escape_xml(form.math_rotations.data),
            'table': {'content': math_table_subdoc}
        }
        
        # Other elements
        context['tier'] = {'one': {'component': escape_xml(form.tier_one_component.data)}}
        context['state'] = {'math': {'domains': escape_xml(form.state_math_domains.data)}}
        context['ipls'] = {
            'additional': {'coverage': escape_xml(form.ipls_additional_coverage.data or '')},
            'critical': {'standards': escape_xml(form.ipls_critical_standards.data or '')}
        }
        
        # Social Studies
        context['tci'] = {'program': {'title': escape_xml(form.tci_program_title.data)}}
        context['ss'] = {'course': {'title': escape_xml(form.ss_course_title.data)}}
        
        # Social studies table content
        context['social'] = {'studies': {'table': {'content': social_studies_table_subdoc}}}
        
        print(f"Final context keys: {list(context.keys())}")
        print(f"Science context: {type(context['science'])}, Math context: {type(context['math'])}")
        print(f"Context structure check - Science keys: {list(context['science'].keys()) if isinstance(context['science'], dict) else 'Not a dict'}")
        
        doc.render(context)
        
        # Generate filename with site name
        site_name_clean = re.sub(r'[^a-zA-Z0-9\s]', '', form.site_name.data).strip()[:30]
        filename = f"Curriculum_Design_Build_{site_name_clean}.docx"
        
        # Save to generated_docs directory
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)
        
        return output_path

def generate_dynamic_table_subdocument(doc, form_data, subject_type):
    """Generate dynamic table subdocument for any subject"""
    subdoc = doc.new_subdoc()
    
    # Get table configuration from form data
    table_rows = int(form_data.get(f'{subject_type}_table_rows', 3))
    table_cols = int(form_data.get(f'{subject_type}_table_cols', 3))
    
    print(f"DEBUG: {subject_type} table - rows={table_rows}, cols={table_cols}")
    
    # Create the dynamic table
    table = subdoc.add_table(rows=table_rows, cols=table_cols)
    table.style = None  # No borders by default, but can be customized
    
    # Populate table cells with user data
    for row_idx in range(table_rows):
        for col_idx in range(table_cols):
            cell_name = f'{subject_type}_table_cell_{row_idx}_{col_idx}'
            cell_value = form_data.get(cell_name, '').strip()
            
            if cell_value:
                cell = table.rows[row_idx].cells[col_idx]
                
                # Handle multi-line content - split by lines and create paragraphs for each
                lines = [line.strip() for line in cell_value.split('\n') if line.strip()]
                
                if lines:
                    # Clear the default paragraph
                    cell._element.clear_content()
                    
                    # Add each line as a separate paragraph in the cell
                    for i, line in enumerate(lines):
                        # Format first row as headers (bold and italicized for grade levels)
                        if row_idx == 0:
                            p = cell.add_paragraph()
                            run = p.add_run(line)
                            run.bold = True
                            # Italicize grade level headers (columns 1 and beyond)
                            if col_idx >= 1:
                                run.italic = True
                        
                        # Format first column as row labels (bold, colon, right-aligned)
                        elif col_idx == 0:
                            # Add colon if not already present
                            text_with_colon = f"{line}:" if not line.endswith(':') else line
                            p = cell.add_paragraph()
                            run = p.add_run(text_with_colon)
                            run.bold = True
                            # Right-align the paragraph
                            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        
                        # Regular content cells
                        else:
                            p = cell.add_paragraph(line)
                        
                        # Adjust spacing between lines
                        if i < len(lines) - 1:  # Not the last line
                            p.paragraph_format.space_after = Pt(2)
    
    # Add extra spacing between table rows for better readability
    for row_idx, row in enumerate(table.rows):
        if row_idx > 0:  # Skip the header row
            # Add space after the last paragraph in each cell of content rows
            for cell in row.cells:
                # Ensure cell has at least one paragraph, then add space after the last one
                if not cell.paragraphs:
                    cell.add_paragraph()  # Add empty paragraph if none exists
                
                # Add spacing after the last paragraph in each cell
                last_para = cell.paragraphs[-1]
                last_para.paragraph_format.space_after = Pt(18)  # Increased to 18pt for visibility
    
    return subdoc

@app.route('/create-correlation-report', methods=['GET', 'POST'])
@login_required
def create_correlation_report():
    form = CorrelationReportForm()
    
    # Populate state choices
    states = State.query.order_by(State.name).all()
    form.state.choices = [('', 'Select State...')] + [(s.code, s.name) for s in states]
    
    if request.method == 'POST':
        # Get form data manually to avoid validation issues with dynamic choices
        state = request.form.get('state')
        grade = request.form.get('grade') 
        subject = request.form.get('subject')
        selected_modules = request.form.getlist('selected_modules')
        
        # Basic validation
        if not state or not grade or not subject or not selected_modules:
            flash('Please fill in all fields and select at least one module.', 'error')
        else:
            try:
                # Generate the document
                doc_path = generate_correlation_report_document(
                    state,
                    grade,
                    subject,
                    selected_modules
                )
                
                # Store in database
                doc_record = GeneratedDocument(
                    user_id=current_user.id,
                    document_type='correlation_report',
                    filename=os.path.basename(doc_path),
                    file_path=doc_path,
                    file_size=os.path.getsize(doc_path)
                )
                db.session.add(doc_record)
                db.session.commit()
                
                flash('Correlation Report generated successfully!', 'success')
                return redirect(url_for('my_documents'))
                
            except Exception as e:
                print(f"Error generating Correlation Report: {e}")
                import traceback
                traceback.print_exc()
                flash(f'Error generating document: {str(e)}', 'error')
    
    return render_template('create_correlation_report.html', form=form)

@app.route('/api/modules')
@login_required
def get_modules_api():
    """API endpoint to get all modules for correlation report form"""
    try:
        modules = Module.query.order_by(Module.title).all()
        module_data = []
        
        for module in modules:
            # Skip "Unnamed: 1" modules which are Excel artifacts
            if module.title.startswith('Unnamed:'):
                continue
                
            module_data.append({
                'id': module.id,
                'title': module.title,
                'subject': module.subject,
                'grade_level': module.grade_level
            })
        
        return {'modules': module_data}
        
    except Exception as e:
        return {'error': str(e)}, 500

@app.route('/api/modules/<subject>', methods=['GET'])
@login_required
def get_modules_by_subject(subject):
    """API endpoint to get modules filtered by subject"""
    modules = Module.query.filter_by(subject=subject).order_by(Module.title).all()
    return jsonify([{'id': m.id, 'title': m.title, 'acronym': m.acronym} for m in modules])

# Streamlined Horizontal Lesson Plan Routes
@app.route('/api/lesson-plan-modules')
@login_required
def get_lesson_plan_modules_api():
    """API endpoint to get all lesson plan modules"""
    from models import LessonPlanModule
    try:
        modules = LessonPlanModule.query.filter_by(active=True).order_by(LessonPlanModule.name).all()
        module_data = []

        for module in modules:
            module_data.append({
                'id': module.id,
                'name': module.name,
                'subject': module.subject,
                'grade_level': module.grade_level,
                'session_count': module.sessions.count()
            })

        return jsonify({'modules': module_data})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def generate_streamlined_horizontal_lesson_plan(school_name, teacher_name, school_year, subject, module_ids):
    """Generate horizontal lesson plan document from database-driven modules"""
    import tempfile
    import shutil
    from models import LessonPlanModule, GeneratedDocument
    
    # Load selected modules with sessions
    modules = LessonPlanModule.query.filter(LessonPlanModule.id.in_(module_ids)).all()

    if not modules:
        raise ValueError("No valid modules found for selected IDs")

    # Template handling - follow correlation report pattern exactly
    template_path = 'templates/docx_templates/hlp_master_template.docx'
    
    # Create temporary copy (like correlation report does)
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(template_path, temp_file.name)
        temp_template_path = temp_file.name
    
    try:
        # Create DocxTemplate instance from temp file
        doc = DocxTemplate(temp_template_path)
        
        # Generate table subdoc using the template instance
        print(f"DEBUG HLP MAIN: About to create subdoc for module IDs: {module_ids}")
        hlp_table_subdoc = create_hlp_table_subdoc(doc, module_ids)
        print(f"DEBUG HLP MAIN: Received subdoc: {type(hlp_table_subdoc)}")
        

        # Create context for template - TEST MULTIPLE APPROACHES
        context = {
            'school': {
                'name': school_name,
                'year': school_year
            },
            'teacher': {
                'name': teacher_name
            },
            'subject': subject,
            'modules': modules,
            # Template expects {{ hlp.table }} so provide exactly that structure:
            'hlp': {'table': hlp_table_subdoc}
        }
        
        print(f"DEBUG HLP MAIN: Context keys: {list(context.keys())}")

        # Render and save
        doc.render(context)
        
        # Create final output path
        filename = f"{school_name.replace(' ', '_')}_{school_year}_HLP.docx"
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
        
        doc.save(output_path)
        
        return output_path
        
    finally:
        # Clean up temp file
        if os.path.exists(temp_template_path):
            os.unlink(temp_template_path)

@app.route('/create-horizontal-lesson-plan-streamlined', methods=['GET', 'POST'])
@login_required
def create_streamlined_horizontal_lesson_plan():
    """Streamlined horizontal lesson plan creation - database driven"""
    from models import LessonPlanModule, GeneratedDocument
    
    form = StreamlinedHorizontalLessonPlanForm()

    if request.method == 'POST':
        # Manual form handling (like correlation report pattern)
        school_name = request.form.get('school_name')
        teacher_name = request.form.get('teacher_name')
        school_year = request.form.get('school_year', '2025-2026')
        subject = request.form.get('subject')
        selected_modules = request.form.getlist('selected_modules')

        # Basic validation
        if not school_name or not teacher_name or not subject or not selected_modules:
            flash('Please fill in all fields and select at least one module.', 'error')
        elif len(selected_modules) > 10:
            flash('Please select no more than 10 modules.', 'error')
        else:
            try:
                # Generate the document using database-driven approach
                doc_path = generate_streamlined_horizontal_lesson_plan(
                    school_name=school_name,
                    teacher_name=teacher_name,
                    school_year=school_year,
                    subject=subject,
                    module_ids=[int(mid) for mid in selected_modules]
                )

                # Save document record
                doc_record = GeneratedDocument(
                    user_id=current_user.id,
                    document_type='horizontal_lesson_plan_streamlined',
                    filename=os.path.basename(doc_path),
                    file_path=doc_path,
                    file_size=os.path.getsize(doc_path)
                )
                db.session.add(doc_record)
                db.session.commit()

                flash('Streamlined Horizontal Lesson Plan generated successfully!', 'success')
                return redirect(url_for('my_documents'))

            except Exception as e:
                print(f"Error generating Streamlined HLP: {e}")
                import traceback
                traceback.print_exc()
                flash(f'Error generating document: {str(e)}', 'error')

    return render_template('create_horizontal_lesson_plan_streamlined.html', form=form)

# Correlation Report Helper Functions
def safe_clear_cell(cell):
    """Safely clear a cell without nuking XML relationships."""
    cell.text = ""  # resets to a single empty paragraph

def set_cell_text(cell, text, font_name, size_pt, bold=False, align='center', vcenter=True):
    """Write formatted text into a cell safely."""
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import html
    
    safe_clear_cell(cell)
    p = cell.paragraphs[0]
    
    # Properly handle ampersand symbols and other HTML entities
    # First unescape any existing HTML entities, then let python-docx handle XML escaping
    if text:
        clean_text = html.unescape(str(text))
    else:
        clean_text = ""
    
    run = p.add_run(clean_text)
    
    # Set font; python-docx sometimes needs rFonts for reliability
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold

    # Force font via rPr.rFonts (helps Rockwell/Arial stick)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    if align == 'center':
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align == 'left':
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif align == 'right':
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    if vcenter:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def shade_cell(cell, hex_fill):
    """Apply background fill to a cell (e.g., '8DC593' or 'EFEFEF')."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill.upper().lstrip('#'))
    tcPr.append(shd)

def set_cell_margins(cell, top_inches=0, bottom_inches=0, left_inches=0, right_inches=0):
    """Set cell margins in inches."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Get cell properties
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create margin element
    tc_mar = tcPr.find(qn('w:tcMar'))
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tcPr.append(tc_mar)
    
    # Convert inches to twips (1 inch = 1440 twips)
    def inches_to_twips(inches):
        return str(int(inches * 1440))
    
    # Set margins
    if top_inches > 0:
        top = tc_mar.find(qn('w:top'))
        if top is None:
            top = OxmlElement('w:top')
            tc_mar.append(top)
        top.set(qn('w:w'), inches_to_twips(top_inches))
        top.set(qn('w:type'), 'dxa')
    
    if bottom_inches > 0:
        bottom = tc_mar.find(qn('w:bottom'))
        if bottom is None:
            bottom = OxmlElement('w:bottom')
            tc_mar.append(bottom)
        bottom.set(qn('w:w'), inches_to_twips(bottom_inches))
        bottom.set(qn('w:type'), 'dxa')
    
    if left_inches > 0:
        left = tc_mar.find(qn('w:left'))
        if left is None:
            left = OxmlElement('w:left')
            tc_mar.append(left)
        left.set(qn('w:w'), inches_to_twips(left_inches))
        left.set(qn('w:type'), 'dxa')
    
    if right_inches > 0:
        right = tc_mar.find(qn('w:right'))
        if right is None:
            right = OxmlElement('w:right')
            tc_mar.append(right)
        right.set(qn('w:w'), inches_to_twips(right_inches))
        right.set(qn('w:type'), 'dxa')

def merge_cells_vertically(table, col, start_row, end_row):
    """Merge cells vertically in a table column from start_row to end_row (inclusive)"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    if start_row >= end_row:
        return  # Nothing to merge
    
    # Set vertical merge for the first cell (vMerge restart)
    first_cell = table.cell(start_row, col)
    first_tc = first_cell._tc
    first_tcPr = first_tc.get_or_add_tcPr()
    first_vmerge = first_tcPr.find(qn('w:vMerge'))
    if first_vmerge is None:
        first_vmerge = OxmlElement('w:vMerge')
        first_tcPr.append(first_vmerge)
    first_vmerge.set(qn('w:val'), 'restart')
    
    # Set vertical merge for subsequent cells (vMerge continue)
    for row_idx in range(start_row + 1, end_row + 1):
        cell = table.cell(row_idx, col)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vmerge = tcPr.find(qn('w:vMerge'))
        if vmerge is None:
            vmerge = OxmlElement('w:vMerge')
            tcPr.append(vmerge)
        # For continuation cells, vMerge has no val attribute or val="continue"
        # We'll leave it empty which defaults to continue

def set_header_row_repeat(row):
    """Set table row to repeat as header on each page"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = trPr.find(qn('w:tblHeader'))
    if tblHeader is None:
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)

def set_row_height(row, inches, exact=True):
    """Set row height; exact prevents Word from auto-expanding."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    # 1 inch = 1440 twips
    h_twips = int(inches * 1440)
    trHeight.set(qn('w:val'), str(h_twips))
    if exact:
        trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def set_table_borders(table):
    """Add all borders to the table (top, bottom, left, right, inside)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Create table borders element
    tblBorders = OxmlElement('w:tblBorders')
    
    # Define all border sides
    border_sides = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    
    for side in border_sides:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')  # Single line border
        border.set(qn('w:sz'), '4')        # Border width (4 = 0.5pt)
        border.set(qn('w:space'), '0')     # No spacing
        border.set(qn('w:color'), '000000') # Black color
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def set_table_left_indent(table, inches=0.13):
    """Set left indent via tblInd."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    tblPr = table._tbl.tblPr
    tblInd = tblPr.xpath('./w:tblInd')
    if tblInd:
        tblInd = tblInd[0]
    else:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    # 1 inch = 1440 twips
    tblInd.set(qn('w:w'), str(int(inches * 1440)))
    tblInd.set(qn('w:type'), 'dxa')


def build_correlation_subdoc(doc, title_set, all_standards, module_to_standards):
    """Build correlation table as subdocument with exact formatting"""
    from docx.shared import Inches
    from docxtpl import Subdoc
    
    print(f"DEBUG: Building subdoc with {len(title_set)} modules, {len(all_standards)} standards")
    print(f"DEBUG: Column widths - Standards: {Inches(0.67)}, Modules: {Inches(1.62)} each")
    
    # Create subdocument
    subdoc = doc.new_subdoc()
    
    # Add table (standards + modules)
    num_cols = len(title_set) + 1  # +1 for standards column
    num_rows = len(all_standards) + 1  # +1 for header
    
    table = subdoc.add_table(rows=num_rows, cols=num_cols)
    
    # Set table-level properties
    set_table_left_indent(table, 0.13)
    set_table_borders(table)
    print(f"DEBUG: Set table indent to 0.13 inches (187 twips) and added all borders")
    
    # Set column widths
    table.columns[0].width = Inches(0.67)  # Standards column
    for i in range(1, num_cols):
        table.columns[i].width = Inches(1.62)  # Module columns
    
    print(f"DEBUG: Set column widths - Col 0: {Inches(0.67)}, Cols 1-{num_cols-1}: {Inches(1.62)}")
    
    # Header row
    header_row = table.rows[0]
    set_row_height(header_row, 0.38, exact=True)
    
    # Header: Standards column
    header_cell = header_row.cells[0]
    set_cell_text(header_cell, "Standards", "Rockwell", 10, bold=True, align='center')
    
    # Header: Module columns
    for i, module_title in enumerate(title_set):
        header_cell = header_row.cells[i + 1]
        set_cell_text(header_cell, module_title, "Rockwell", 10, bold=True, align='center')
    
    print(f"DEBUG: Header row height set to 0.38 inches EXACT")
    
    # Set header row to repeat on each page
    set_header_row_repeat(header_row)
    
    # Data rows
    for row_idx, standard_code in enumerate(all_standards):
        data_row = table.rows[row_idx + 1]
        set_row_height(data_row, 0.31, exact=True)
        
        # Determine row shading (alternating: #EFEFEF, white)
        is_gray_row = row_idx % 2 == 0  # First data row (row_idx=0) gets gray
        row_shade = "EFEFEF" if is_gray_row else "FFFFFF"
        
        # Standards column
        std_cell = data_row.cells[0]
        set_cell_text(std_cell, standard_code, "Arial", 8, bold=False, align='center')
        shade_cell(std_cell, row_shade)
        
        # Module columns
        for col_idx, module_title in enumerate(title_set):
            module_cell = data_row.cells[col_idx + 1]
            
            # Check if this module covers this standard
            module_standards = module_to_standards.get(module_title, set())
            has_coverage = standard_code in module_standards
            
            if has_coverage:
                set_cell_text(module_cell, "X", "Arial", 8, bold=True, align='center')
                shade_cell(module_cell, "8DC593")  # Green override
            else:
                set_cell_text(module_cell, "", "Arial", 8, bold=False, align='center')
                shade_cell(module_cell, row_shade)  # Row alternating color
    
    print(f"DEBUG: Data rows height set to 0.31 inches EXACT, alternating shades applied")
    print(f"DEBUG: Green cells (#8DC593) applied for coverage, alternating row shades: #EFEFEF/white")
    
    return subdoc

def build_coverage_report_subdoc(doc, title_set, all_standards, module_to_standards, standard_descriptions):
    """Build coverage report table showing standards and which modules cover them"""
    from docx.shared import Inches
    from docxtpl import Subdoc
    import html
    
    print(f"DEBUG: Building coverage report subdoc with {len(all_standards)} standards, {len(title_set)} modules")
    
    # Create subdocument
    subdoc = doc.new_subdoc()
    
    # Create table with 3 columns: Standard, Statement, Selected Module
    num_cols = 3
    num_rows = len(all_standards) + 1  # +1 for header
    
    table = subdoc.add_table(rows=num_rows, cols=num_cols)
    
    # Set table-level properties
    set_table_left_indent(table, 0.13)
    set_table_borders(table)
    
    # Enable autofit for coverage report table (handles long descriptions)
    table.autofit = True
    
    # Set initial column widths as proportional guidelines for autofit
    table.columns[0].width = Inches(0.8)   # Standard code column
    table.columns[1].width = Inches(3.5)   # Standard description column (wider for descriptions)
    table.columns[2].width = Inches(2.0)   # Selected modules column
    
    print(f"DEBUG: Set coverage table column widths - Standard: {Inches(0.8)}, Description: {Inches(3.5)}, Modules: {Inches(2.0)}, Autofit: ENABLED")
    
    # Header row
    header_row = table.rows[0]
    set_row_height(header_row, 0.38, exact=True)
    
    # Header cells - Rockwell 11pt bold
    set_cell_text(header_row.cells[0], "Standard", "Rockwell", 11, bold=True, align='center')
    set_cell_text(header_row.cells[1], "Statement", "Rockwell", 11, bold=True, align='center')
    set_cell_text(header_row.cells[2], "Selected Module", "Rockwell", 11, bold=True, align='center')
    
    # Apply cell margins to header row
    for cell in header_row.cells:
        set_cell_margins(cell, top_inches=0.05, bottom_inches=0.05)
    
    # Set header row to repeat on each page
    set_header_row_repeat(header_row)
    
    # Build reverse mapping: standard -> list of modules that cover it
    standard_to_modules = {}
    for standard_code in all_standards:
        covering_modules = []
        for module_title in title_set:
            module_standards = module_to_standards.get(module_title, set())
            if standard_code in module_standards:
                covering_modules.append(module_title)
        standard_to_modules[standard_code] = covering_modules
    
    # Data rows
    for row_idx, standard_code in enumerate(all_standards):
        data_row = table.rows[row_idx + 1]
        # Allow rows to auto-size based on content (don't set exact height for long descriptions)
        # set_row_height(data_row, 0.31, exact=True)  # Commented out to allow auto-sizing
        
        # Get covering modules for this standard
        covering_modules = standard_to_modules[standard_code]
        is_uncovered = len(covering_modules) == 0
        
        # Determine row shading
        if is_uncovered:
            # Gray shading for uncovered standards
            row_shade = "D3D3D3"  # Light gray for uncovered standards
        else:
            # Alternating shading for covered standards
            is_gray_row = row_idx % 2 == 0
            row_shade = "EFEFEF" if is_gray_row else "FFFFFF"
        
        # Standard code column - Times New Roman 10pt
        std_cell = data_row.cells[0]
        set_cell_text(std_cell, standard_code, "Times New Roman", 10, bold=False, align='center')
        shade_cell(std_cell, row_shade)
        set_cell_margins(std_cell, top_inches=0.05, bottom_inches=0.05)
        
        # Standard description column - Times New Roman 10pt
        desc_cell = data_row.cells[1]
        # Get description from standard_descriptions dict
        description = standard_descriptions.get(standard_code, "")
        # Handle math symbols and HTML entities
        clean_description = html.unescape(description) if description else ""
        set_cell_text(desc_cell, clean_description, "Times New Roman", 10, bold=False, align='left')
        shade_cell(desc_cell, row_shade)
        set_cell_margins(desc_cell, top_inches=0.05, bottom_inches=0.05)
        
        # Selected modules column - Rockwell 10pt bold centered
        modules_cell = data_row.cells[2]
        if covering_modules:
            # Join module names with line breaks
            modules_text = "\n".join(covering_modules)
            set_cell_text(modules_cell, modules_text, "Rockwell", 10, bold=True, align='center')
        else:
            set_cell_text(modules_cell, "", "Rockwell", 10, bold=True, align='center')
        shade_cell(modules_cell, row_shade)
        set_cell_margins(modules_cell, top_inches=0.05, bottom_inches=0.05)
    
    print(f"DEBUG: Coverage report table completed - {len([s for s, m in standard_to_modules.items() if not m])} uncovered standards (gray)")
    print(f"DEBUG: Coverage report - uncovered standards are shaded #D3D3D3, others alternate #EFEFEF/white")
    print(f"DEBUG: Coverage report - AUTOFIT ENABLED, rows will auto-size for long descriptions")
    print(f"DEBUG: Coverage report formatting - Headers: Rockwell 11pt bold, Standards: Times New Roman 10pt, Modules: Rockwell 10pt bold centered")
    print(f"DEBUG: Coverage report - Cell margins: 0.05\" top/bottom applied to all cells")
    
    return subdoc

def build_coverage_report_by_product_subdoc(doc, title_set, all_standards, module_to_standards, standard_descriptions):
    """Build coverage report table organized by product/module with individual rows for each standard"""
    from docx.shared import Inches
    from docxtpl import Subdoc
    import html
    
    print(f"DEBUG: Building coverage report by product subdoc with {len(title_set)} modules")
    
    # Create subdocument
    subdoc = doc.new_subdoc()
    
    # Calculate total rows needed - one row per standard across all modules
    total_rows = 0
    sorted_modules = sorted(title_set)
    for module_title in sorted_modules:
        module_standards = module_to_standards.get(module_title, set())
        total_rows += max(1, len(module_standards))  # At least 1 row per module
    
    # Create table with 3 columns: Title, Standards, Statement
    num_cols = 3
    num_rows = total_rows + 1  # +1 for header
    
    table = subdoc.add_table(rows=num_rows, cols=num_cols)
    
    # Set table-level properties
    set_table_left_indent(table, 0.13)
    set_table_borders(table)
    
    # Enable autofit for by-product table
    table.autofit = True
    
    # Set column widths - Title, Standards, Statement
    table.columns[0].width = Inches(1.8)   # Title column
    table.columns[1].width = Inches(0.8)   # Standards column (narrow for codes)
    table.columns[2].width = Inches(3.7)   # Statement column (wide for descriptions)
    
    print(f"DEBUG: Set by-product table column widths - Title: {Inches(1.8)}, Standards: {Inches(0.8)}, Statement: {Inches(3.7)}")
    
    # Header row
    header_row = table.rows[0]
    set_row_height(header_row, 0.38, exact=True)
    
    # Header cells - Rockwell 11pt bold
    set_cell_text(header_row.cells[0], "Title", "Rockwell", 11, bold=True, align='center')
    set_cell_text(header_row.cells[1], "Standards", "Rockwell", 11, bold=True, align='center')
    set_cell_text(header_row.cells[2], "Statement", "Rockwell", 11, bold=True, align='center')
    
    # Apply cell margins to header row
    for cell in header_row.cells:
        set_cell_margins(cell, top_inches=0.05, bottom_inches=0.05)
    
    # Set header row to repeat on each page
    set_header_row_repeat(header_row)
    
    # Data rows - create individual rows for each standard with merged title cells
    current_row = 1
    
    for module_idx, module_title in enumerate(sorted_modules):
        # Get standards covered by this module
        module_standards = module_to_standards.get(module_title, set())
        sorted_standards = sorted(list(module_standards))
        
        if not sorted_standards:
            # Handle modules with no standards - create one row
            sorted_standards = [""]
        
        # Alternating shading per module (not per row)
        is_gray_module = module_idx % 2 == 0
        row_shade = "EFEFEF" if is_gray_module else "FFFFFF"
        
        # Remember the start row for this module (for merging)
        module_start_row = current_row
        
        # Create a row for each standard in this module
        for standard_idx, standard_code in enumerate(sorted_standards):
            data_row = table.rows[current_row]
            
            # Title column - only set text in first row, then we'll merge all rows for this module
            title_cell = data_row.cells[0]
            if standard_idx == 0:  # Only set text in the first row
                set_cell_text(title_cell, module_title, "Rockwell", 10, bold=True, align='center')
            shade_cell(title_cell, row_shade)
            set_cell_margins(title_cell, top_inches=0.05, bottom_inches=0.05)
            
            # Standards column - Times New Roman 10pt centered
            standards_cell = data_row.cells[1]
            set_cell_text(standards_cell, standard_code, "Times New Roman", 10, bold=False, align='center')
            shade_cell(standards_cell, row_shade)
            set_cell_margins(standards_cell, top_inches=0.05, bottom_inches=0.05)
            
            # Statement column - Times New Roman 10pt left-aligned
            statement_cell = data_row.cells[2]
            if standard_code and standard_code in standard_descriptions:
                description = standard_descriptions[standard_code]
                clean_description = html.unescape(description) if description else ""
                set_cell_text(statement_cell, clean_description, "Times New Roman", 10, bold=False, align='left')
            else:
                placeholder_text = "No standards mapped" if not standard_code else ""
                set_cell_text(statement_cell, placeholder_text, "Times New Roman", 10, bold=False, align='left')
            shade_cell(statement_cell, row_shade)
            set_cell_margins(statement_cell, top_inches=0.05, bottom_inches=0.05)
            
            current_row += 1
        
        # Merge title cells vertically for this module (if more than one row)
        module_end_row = current_row - 1
        if module_end_row > module_start_row:
            merge_cells_vertically(table, 0, module_start_row, module_end_row)  # Column 0 is the Title column
            print(f"DEBUG: Merged title cells for '{module_title}' from row {module_start_row} to {module_end_row}")
    
    print(f"DEBUG: Coverage by product table completed - {len(sorted_modules)} modules, {total_rows} total standard rows")
    print(f"DEBUG: By-product formatting - Headers: Rockwell 11pt bold, Titles: Rockwell 10pt bold (merged), Standards: Times New Roman 10pt centered, Statements: Times New Roman 10pt left")
    print(f"DEBUG: By-product table - alternating module shades: #EFEFEF/white, WITH vertical cell merging for module titles")
    
    return subdoc

def build_uncorrelated_standards_subdoc(doc, all_standards, module_to_standards, standard_descriptions, title_set):
    """Build table showing standards NOT covered by the selected modules"""
    from docx.shared import Inches
    from docxtpl import Subdoc
    import html
    
    # Find uncorrelated standards (standards not covered by any selected module)
    covered_standards = set()
    for module_standards in module_to_standards.values():
        covered_standards.update(module_standards)
    
    uncorrelated_standards = [std for std in all_standards if std not in covered_standards]
    
    print(f"DEBUG: Building uncorrelated standards subdoc with {len(uncorrelated_standards)} uncovered standards out of {len(all_standards)} total")
    
    if not uncorrelated_standards:
        # If no uncorrelated standards, create empty subdoc with message
        subdoc = doc.new_subdoc()
        para = subdoc.add_paragraph("All standards are covered by the selected modules.")
        return subdoc
    
    # Create subdocument
    subdoc = doc.new_subdoc()
    
    # Create table with 3 columns: Standard, Statement, District Resources
    num_cols = 3
    num_rows = len(uncorrelated_standards) + 1  # +1 for header
    
    table = subdoc.add_table(rows=num_rows, cols=num_cols)
    
    # Set table-level properties
    set_table_left_indent(table, 0.13)
    set_table_borders(table)
    
    # Enable autofit for uncorrelated standards table (handles long descriptions)
    table.autofit = True
    
    # Set initial column widths
    table.columns[0].width = Inches(0.8)   # Standard code column
    table.columns[1].width = Inches(3.5)   # Standard description column (wider for descriptions)
    table.columns[2].width = Inches(2.0)   # District resources column
    
    print(f"DEBUG: Set uncorrelated standards table column widths - Standard: {Inches(0.8)}, Description: {Inches(3.5)}, Resources: {Inches(2.0)}, Autofit: ENABLED")
    
    # Header row
    header_row = table.rows[0]
    set_row_height(header_row, 0.38, exact=True)
    
    # Header cells - Rockwell 11pt bold
    set_cell_text(header_row.cells[0], "Standard", "Rockwell", 11, bold=True, align='center')
    set_cell_text(header_row.cells[1], "Statement", "Rockwell", 11, bold=True, align='center')
    set_cell_text(header_row.cells[2], "District Resources", "Rockwell", 11, bold=True, align='center')
    
    # Apply cell margins to header row
    for cell in header_row.cells:
        set_cell_margins(cell, top_inches=0.05, bottom_inches=0.05)
    
    # Set header row to repeat on each page
    set_header_row_repeat(header_row)
    
    # Data rows
    for row_idx, standard_code in enumerate(uncorrelated_standards):
        data_row = table.rows[row_idx + 1]
        
        # Alternating shading for rows
        is_gray_row = row_idx % 2 == 0
        row_shade = "EFEFEF" if is_gray_row else "FFFFFF"
        
        # Standard code column - Times New Roman 10pt centered
        std_cell = data_row.cells[0]
        set_cell_text(std_cell, standard_code, "Times New Roman", 10, bold=False, align='center')
        shade_cell(std_cell, row_shade)
        set_cell_margins(std_cell, top_inches=0.05, bottom_inches=0.05)
        
        # Standard description column - Times New Roman 10pt left-aligned
        desc_cell = data_row.cells[1]
        # Get description from standard_descriptions dict
        description = standard_descriptions.get(standard_code, "")
        # Handle math symbols and HTML entities
        clean_description = html.unescape(description) if description else ""
        set_cell_text(desc_cell, clean_description, "Times New Roman", 10, bold=False, align='left')
        shade_cell(desc_cell, row_shade)
        set_cell_margins(desc_cell, top_inches=0.05, bottom_inches=0.05)
        
        # District resources column - empty for user to fill in manually
        resources_cell = data_row.cells[2]
        set_cell_text(resources_cell, "", "Times New Roman", 10, bold=False, align='left')
        shade_cell(resources_cell, row_shade)
        set_cell_margins(resources_cell, top_inches=0.05, bottom_inches=0.05)
    
    print(f"DEBUG: Uncorrelated standards table completed - {len(uncorrelated_standards)} standards not covered by selected modules")
    
    return subdoc

def generate_correlation_report_document(state, grade_level, subject, selected_module_ids):
    """Generate correlation report document using subdoc approach"""
    print(f"DEBUG: Generating report for {state}, {grade_level}, {subject}")
    print(f"DEBUG: Selected module IDs: {selected_module_ids}")
    
    # Get basic data
    state_obj = State.query.filter_by(code=state).first()
    state_name = state_obj.name if state_obj else state
    
    # Get modules (ordered by selection)
    modules = Module.query.filter(Module.id.in_([int(mid) for mid in selected_module_ids])).all()
    # Maintain order from selection
    module_lookup = {m.id: m for m in modules}
    ordered_modules = [module_lookup[int(mid)] for mid in selected_module_ids if int(mid) in module_lookup]
    title_set = [m.title for m in ordered_modules]
    
    # Get standards using clean helper function
    # Extract grade number from formats like "8th Grade" -> 8
    grade_num = None
    if 'Grade' in grade_level:
        grade_part = grade_level.split()[0]  # "8th" from "8th Grade"
        grade_num = int(grade_part.rstrip('thrdns'))  # Remove ordinal suffixes
    all_standards = get_all_standards(state, grade_num, subject)
    
    print(f"DEBUG: Found {len(ordered_modules)} modules, {len(all_standards)} standards")
    
    # Get module-to-standards mapping for ALL selected modules (any subject/grade)
    module_to_standards = {}
    if ordered_modules:
        module_ids = [m.id for m in ordered_modules]
        mappings = (db.session.query(Module.title, Standard.code)
                   .join(ModuleStandardMapping, Module.id==ModuleStandardMapping.module_id)
                   .join(Standard, Standard.id==ModuleStandardMapping.standard_id)
                   .filter(Module.id.in_(module_ids))
                   .all())
        
        for title, code in mappings:
            if title not in module_to_standards:
                module_to_standards[title] = set()
            module_to_standards[title].add(code)
    
    # Subset for selected modules only (but we already filtered above)
    subset = {t: module_to_standards.get(t, set()) for t in title_set}
    
    print(f"DEBUG: Prepared module-to-standards mapping for {len(subset)} modules")
    for title, codes in subset.items():
        print(f"DEBUG: {title} covers {len(codes)} standards")
    
    # Template handling
    template_path = 'templates/docx_templates/correlation_report_master.docx'
    print(f"DEBUG: Using template at: {template_path}")
    
    # Load template
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(template_path, temp_file.name)
        temp_template_path = temp_file.name
    
    try:
        doc = DocxTemplate(temp_template_path)
        print(f"DEBUG: TPL instance for subdoc+render: {id(doc)}")
        
        # Generate subdocument FROM THIS EXACT INSTANCE
        subdoc = build_correlation_subdoc(doc, title_set, all_standards, subset)
        
        # Get standard descriptions for coverage report
        standard_descriptions = {}
        if all_standards:
            standards_with_desc = (db.session.query(Standard.code, Standard.description)
                                  .filter(Standard.code.in_(all_standards))
                                  .all())
            standard_descriptions = {code: desc for code, desc in standards_with_desc}
            print(f"DEBUG: Fetched descriptions for {len(standard_descriptions)} standards")
        
        # Generate coverage report subdocument
        coverage_subdoc = build_coverage_report_subdoc(doc, title_set, all_standards, subset, standard_descriptions)
        
        # Generate coverage report by product subdocument
        coverage_by_product_subdoc = build_coverage_report_by_product_subdoc(doc, title_set, all_standards, subset, standard_descriptions)
        
        # Generate uncorrelated standards subdocument
        district_covered_standards = build_uncorrelated_standards_subdoc(doc, all_standards, subset, standard_descriptions, title_set)
        
        # Create modules list for template (maintains compatibility with existing template)
        import html
        
        # Clean module titles and prepare for template
        modules_list = []
        for title in title_set:
            # Unescape any XML/HTML entities to ensure symbols like & display correctly
            clean_title = html.unescape(title) if title else title
            modules_list.append(clean_title)
            print(f"DEBUG: Module title '{title}' -> cleaned: '{clean_title}'")
        
        # Create a formatted string that maintains line breaks for the existing template placeholder
        modules_newline_list = '\n'.join(modules_list)
        print(f"DEBUG: modules_newline_list content: '{modules_newline_list}'")
        
        # Context for template
        context = {
            'state': state_name,
            'grade': grade_level,
            'subject': subject,
            'module_list': modules_list,  # For loop-based template (matches your template variable)
            'modules_list': modules_list,  # For future use
            'modules_newline_list': modules_newline_list,  # For current template compatibility
            'correlation_table': subdoc,
            'coverage_report': coverage_subdoc,  # Standards-first coverage report table
            'coverage_report_by_product': coverage_by_product_subdoc,  # Product-first coverage report table
            'district_covered_standards': district_covered_standards  # Uncorrelated standards table
        }
        
        print("DEBUG: Rendering template with subdoc...")
        doc.render(context)
        
        # Save
        filename = f"Correlation_Report_{subject}_{grade_level.replace(' ', '_')}.docx"
        
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
        
        print(f"DEBUG: Saving to {output_path}")
        print(f"DEBUG: Rendered with tpl id: {id(doc)} -> {output_path}")
        doc.save(output_path)
        
        return output_path
        
    except Exception as e:
        print(f"Error generating correlation report: {e}")
        import traceback
        traceback.print_exc()
        raise
    finally:
        # Clean up temp file
        if os.path.exists(temp_template_path):
            os.unlink(temp_template_path)

@app.route('/test-correlation-table/<grade>/<subject>/<int:num_modules>')
@login_required
def test_correlation_table(grade, subject, num_modules):
    """Test route for correlation table generation"""
    try:
        # Get sample modules for the subject (limit by num_modules)
        modules = Module.query.filter_by(subject=subject).limit(num_modules).all()
        module_ids = [str(m.id) for m in modules]
        
        if not modules:
            return f"No modules found for subject: {subject}", 404
        
        # Generate document
        doc_path = generate_correlation_report_document(
            state='LA',  # Louisiana
            grade_level=grade,
            subject=subject,
            selected_module_ids=module_ids
        )
        
        # Return file
        return send_file(doc_path, as_attachment=True, download_name=f"TEST_{subject}_{grade}_{num_modules}modules.docx")
        
    except Exception as e:
        return f"Error: {str(e)}", 500

# IPL (Individual Pacing List) Routes
@app.route('/create-ipl-report', methods=['GET', 'POST'])
@login_required
def create_ipl_report():
    form = IplReportForm()
    
    if request.method == 'POST':
        # Get form data manually to handle dynamic module selection
        module_acronym = request.form.get('module_acronym')
        selected_modules = request.form.getlist('selected_modules')
        
        # Basic validation
        if not module_acronym or not selected_modules:
            flash('Please fill in the module acronym and select at least one module.', 'error')
        else:
            try:
                # Generate the document
                doc_path = generate_ipl_report_document(
                    module_acronym,
                    selected_modules
                )
                
                # Store in database
                doc_record = GeneratedDocument(
                    user_id=current_user.id,
                    document_type='ipl_report',
                    filename=os.path.basename(doc_path),
                    file_path=doc_path,
                    file_size=os.path.getsize(doc_path)
                )
                db.session.add(doc_record)
                db.session.commit()
                
                flash('Module IPL List generated successfully!', 'success')
                return redirect(url_for('my_documents'))
                
            except Exception as e:
                print(f"Error generating Module IPL List: {e}")
                import traceback
                traceback.print_exc()
                flash(f'Error generating document: {str(e)}', 'error')
    
    return render_template('create_ipl_report.html', form=form)

@app.route('/api/ipl-modules')
@login_required
def get_ipl_modules_api():
    """API endpoint to get all IPL modules for form"""
    try:
        modules = IplModule.query.filter_by(active=True).order_by(IplModule.name).all()
        module_data = []
        
        for module in modules:
            module_data.append({
                'id': module.id,
                'name': module.name
            })
        
        return {'modules': module_data}
        
    except Exception as e:
        print(f"Error loading IPL modules: {e}")
        return {'error': 'Failed to load modules'}, 500

def create_ipl_table_subdoc(doc, selected_modules):
    """Generate a formatted IPL table subdoc matching the expected format"""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docxtpl import Subdoc
    
    print(f"DEBUG IPL TABLE: Starting subdoc creation for {len(selected_modules)} modules")
    
    # Create subdoc from the DocxTemplate instance
    subdoc = doc.new_subdoc()
    print(f"DEBUG IPL TABLE: Created subdoc: {type(subdoc)}")
    
    # Get all entries for selected modules, ordered properly
    all_entries = []
    for module in selected_modules:
        entries = IplEntry.query.filter_by(module_id=module.id).order_by(IplEntry.order_index).all()
        
        # Add module header entry first
        module_header_entry = type('obj', (object,), {
            'module_name': module.name,
            'unit_name': None,
            'ipl_title': None,
            'goal_text': None,
            'is_module_header': True
        })()
        all_entries.append(module_header_entry)
        
        # Add all entries for this module
        all_entries.extend(entries)
        print(f"DEBUG IPL TABLE: Module {module.name} has {len(entries)} entries")
    
    if not all_entries:
        print("DEBUG IPL TABLE: No entries found, adding error paragraph")
        p = subdoc.add_paragraph("No IPL entries found.")
        return subdoc
    
    # Helper function to set cell text with advanced formatting
    def set_cell_text(cell, text, font_name='Arial', font_size=10, bold=False, italic=False, center=False, left=False, font_color=None):
        if text is None:
            text = ""
        cell.text = str(text)
        
        # Set vertical alignment to center
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run.font.italic = italic
                if font_color:
                    run.font.color.rgb = font_color
                    
            if center:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif left:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Helper function to shade a cell with specific color
    def shade_cell(cell, color_hex):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Helper function to merge cells
    def merge_cells(cell1, cell2):
        """Merge cell2 into cell1"""
        cell1.merge(cell2)
    
    # Process data and create structured rows (skip module headers since they're in document header)
    structured_rows = []
    
    # Add column header once at the beginning
    structured_rows.append({
        'type': 'column_header'
    })
    
    for entry in all_entries:
        if hasattr(entry, 'is_module_header') and entry.is_module_header:
            # Skip module header since it's already in document header
            continue
        else:
            structured_rows.append({
                'type': 'data',
                'unit': entry.unit_name,
                'ipl': entry.ipl_title, 
                'goal': entry.goal_text
            })
    
    # Group data by Unit, then by IPL within each unit
    grouped_data = []
    i = 0
    while i < len(structured_rows):
        row = structured_rows[i]
        
        if row['type'] != 'data':
            grouped_data.append(row)
            i += 1
        else:
            # Group consecutive rows with same unit
            unit_group = []
            current_unit = row['unit']
            
            while i < len(structured_rows) and structured_rows[i]['type'] == 'data':
                current_row = structured_rows[i]
                if current_row['unit'] == current_unit or not current_row['unit']:
                    unit_group.append(current_row)
                    i += 1
                else:
                    # Different unit, start new group
                    break
            
            # Now group within the unit by IPL
            ipl_groups = []
            j = 0
            while j < len(unit_group):
                current_ipl = unit_group[j]['ipl']
                ipl_rows = []
                goals_for_ipl = []
                
                # Collect all rows for this IPL (including those with empty IPL that belong to previous IPL)
                while j < len(unit_group):
                    row_data = unit_group[j]
                    if row_data['ipl'] == current_ipl or (not row_data['ipl'] and current_ipl):
                        ipl_rows.append(row_data)
                        if row_data['goal']:
                            goals_for_ipl.append(row_data['goal'])
                        j += 1
                    elif row_data['ipl'] and row_data['ipl'] != current_ipl:
                        # Different IPL, start new group
                        break
                    else:
                        j += 1
                
                if current_ipl or goals_for_ipl:  # Only add if we have an IPL or goals
                    ipl_groups.append({
                        'ipl': current_ipl,
                        'goals': goals_for_ipl
                    })
            
            grouped_data.append({
                'type': 'unit_group', 
                'unit': current_unit,
                'ipl_groups': ipl_groups
            })
    
    # Calculate the actual number of rows needed after processing the data structure
    total_rows_needed = 0
    
    # Count rows from grouped_data structure
    for group in grouped_data:
        if group['type'] == 'column_header':
            total_rows_needed += 1  # Column header
        elif group['type'] == 'unit_group':
            total_rows_needed += len(group['ipl_groups'])  # One row per IPL
    
    # Create table with 3 columns: IPL Units, IPLs, Goals
    num_cols = 3
    num_rows = total_rows_needed
    
    print(f"DEBUG IPL TABLE: Creating table with {num_rows} rows and {num_cols} cols")
    table = subdoc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set cell margins: 0.1" top/bottom, 0.08" left/right
    for row in table.rows:
        for cell in row.cells:
            # Set cell margins using table cell properties
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            # Create margins element
            tcMar = OxmlElement('w:tcMar')
            
            # Top margin - 0.08" = 115 twips (1" = 1440 twips)
            top = OxmlElement('w:top')
            top.set(qn('w:w'), '115')
            top.set(qn('w:type'), 'dxa')
            tcMar.append(top)
            
            # Bottom margin - 0.08" = 115 twips  
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:w'), '115')
            bottom.set(qn('w:type'), 'dxa')
            tcMar.append(bottom)
            
            # Left margin - 0.08" = 115 twips
            left = OxmlElement('w:left')
            left.set(qn('w:w'), '115')
            left.set(qn('w:type'), 'dxa')
            tcMar.append(left)
            
            # Right margin - 0.08" = 115 twips
            right = OxmlElement('w:right')
            right.set(qn('w:w'), '115')
            right.set(qn('w:type'), 'dxa')
            tcMar.append(right)
            
            # Add margins to cell properties
            tcPr.append(tcMar)
    
    # Now populate the table with proper merging
    row_idx = 0
    
    for group in grouped_data:
        if group['type'] == 'column_header':
            # Column header row
            header_row = table.rows[row_idx]
            set_cell_text(header_row.cells[0], 'IPL Unit', 'Rockwell', 12, bold=True, center=True)
            set_cell_text(header_row.cells[1], 'IPLs', 'Rockwell', 12, bold=True, center=True)
            set_cell_text(header_row.cells[2], 'Goals', 'Rockwell', 12, bold=True, center=True)
            
            # Shade column headers
            for cell in header_row.cells:
                shade_cell(cell, 'F2F2F2')
                
            row_idx += 1
            
        elif group['type'] == 'unit_group':
            # Data rows with merging for IPL Units - one row per IPL with combined goals
            ipl_groups = group['ipl_groups']
            unit_start_row = row_idx
            total_ipls_for_unit = len(ipl_groups)
            
            for ipl_idx, ipl_group in enumerate(ipl_groups):
                data_row = table.rows[row_idx]
                
                # Unit name: only in first row of entire unit group
                if ipl_idx == 0 and group['unit']:
                    set_cell_text(data_row.cells[0], group['unit'], 'Rockwell', 11, bold=True, center=True)
                else:
                    set_cell_text(data_row.cells[0], '', 'Rockwell', 11)
                
                # IPL name
                set_cell_text(data_row.cells[1], ipl_group['ipl'] or '', 'Times New Roman', 10.5, italic=True, center=True)
                
                # Combined goals: join all goals with line breaks
                combined_goals = '\n'.join(ipl_group['goals']) if ipl_group['goals'] else ''
                set_cell_text(data_row.cells[2], combined_goals, 'Times New Roman', 10, center=True)
                
                row_idx += 1
            
            # Merge Unit cells if multiple IPLs for this unit
            if total_ipls_for_unit > 1 and group['unit']:
                first_unit_cell = table.rows[unit_start_row].cells[0]
                for merge_row_idx in range(unit_start_row + 1, row_idx):
                    merge_cells(first_unit_cell, table.rows[merge_row_idx].cells[0])
    
    # Set column widths to match the image proportions
    table.columns[0].width = Inches(1.8)  # IPL Units
    table.columns[1].width = Inches(2.2)  # IPLs  
    table.columns[2].width = Inches(3.5)  # Goals (widest)
    
    print(f"DEBUG IPL TABLE: Table creation completed successfully, returning subdoc")
    return subdoc

def generate_ipl_report_document(module_acronym, selected_module_ids):
    """Generate Module IPL List document using subdoc approach"""
    print(f"DEBUG IPL: Generating report for modules: {selected_module_ids}")
    
    # Get selected modules
    modules = IplModule.query.filter(IplModule.id.in_([int(mid) for mid in selected_module_ids])).all()
    
    if not modules:
        raise ValueError("No valid modules found for selected IDs")
    
    print(f"DEBUG IPL: Found {len(modules)} modules")
    
    # Template handling - follow correlation report pattern
    template_path = 'templates/docx_templates/ipl_template_master.docx'
    
    # Create temporary copy
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        shutil.copy2(template_path, temp_file.name)
        temp_template_path = temp_file.name
    
    try:
        # Create DocxTemplate instance
        doc = DocxTemplate(temp_template_path)
        
        # Generate IPL table subdoc
        print(f"DEBUG IPL: About to create subdoc for modules: {[m.name for m in modules]}")
        ipl_table_subdoc = create_ipl_table_subdoc(doc, modules)
        print(f"DEBUG IPL: Received subdoc: {type(ipl_table_subdoc)}")
        
        # Create context for template
        context = {
            'module_acronym': module_acronym,
            'modules': modules,
            'ipl': {'table': ipl_table_subdoc}
        }
        
        print(f"DEBUG IPL: Context keys: {list(context.keys())}")
        
        # Render and save
        doc.render(context)
        
        # Create final output path
        filename = f"{module_acronym}_IPL_List.docx"
        output_dir = 'generated_docs'
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
        
        doc.save(output_path)
        
        return output_path
        
    finally:
        # Clean up temp file
        if os.path.exists(temp_template_path):
            os.unlink(temp_template_path)

if __name__ == '__main__':
    # Create default admin if none exists (for development/initial setup)
    with app.app_context():
        try:
            # Check if any admin exists
            if not User.query.filter_by(is_admin=True).first():
                # Create default admin
                default_admin = User(
                    email='admin@nola.docs',
                    username='admin',
                    first_name='Admin',
                    last_name='User',
                    is_admin=True,
                    is_active=True
                )
                default_admin.set_password('admin123')  # Change this password!
                
                db.session.add(default_admin)
                db.session.commit()
                print("✅ Default admin created: admin@nola.docs / admin123")
        except Exception as e:
            print(f"Note: {e}")
    
    app.run(debug=True, port=5002) 