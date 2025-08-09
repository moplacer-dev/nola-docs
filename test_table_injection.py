#!/usr/bin/env python3
"""
Test that the table is actually injected into the document
"""
import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from app import app, build_correlation_subdoc, Module, Standard
from docxtpl import DocxTemplate
from docx import Document
import tempfile

def create_minimal_template_and_test():
    """Create a minimal template with known placeholder and test injection"""
    
    with app.app_context():
        print("=== Creating minimal test template ===")
        
        # Create a simple test document template
        doc = Document()
        doc.add_heading('Test Correlation Report', 0)
        doc.add_paragraph('State: {{ state }}')
        doc.add_paragraph('Grade: {{ grade }}')
        doc.add_paragraph('Subject: {{ subject }}')
        doc.add_paragraph('Table below:')
        
        # Add a paragraph with the placeholder
        p = doc.add_paragraph()
        p.add_run('{{ correlation_table }}')
        
        # Save as template
        template_path = 'generated_docs/TEST_TEMPLATE.docx'
        doc.save(template_path)
        print(f"✓ Minimal template created: {template_path}")
        
        # Now test with this template
        print("\n=== Testing with minimal template ===")
        
        # Get test data
        modules = Module.query.filter_by(subject='Math').limit(3).all()
        title_set = [m.title for m in modules]
        
        standards = Standard.query.filter_by(subject='Math', grade_level='8th Grade').limit(5).all()
        all_standards = [s.code for s in standards]
        
        # Create test mapping
        module_to_standards = {
            title_set[0]: {all_standards[0], all_standards[1]} if len(all_standards) > 1 else set(),
            title_set[1]: {all_standards[2]} if len(all_standards) > 2 else set(),
            title_set[2]: {all_standards[0]} if len(all_standards) > 0 else set(),
        }
        
        print(f"Test data - Modules: {title_set}")
        print(f"Test data - Standards: {all_standards}")
        print(f"Test data - Mappings: {module_to_standards}")
        
        # Load template
        doc_template = DocxTemplate(template_path)
        
        # Create subdoc
        subdoc = build_correlation_subdoc(doc_template, title_set, all_standards, module_to_standards)
        print(f"✓ Subdoc created: {type(subdoc)}")
        
        # Render with context
        context = {
            'state': 'Louisiana', 
            'grade': '8th Grade',
            'subject': 'Math',
            'correlation_table': subdoc
        }
        
        doc_template.render(context)
        
        # Save result
        result_path = 'generated_docs/TEST_TABLE_INJECTION_RESULT.docx'
        doc_template.save(result_path)
        
        print(f"✓ Test document generated: {result_path}")
        print("✓ Open this document to visually verify the table appears!")
        
        return result_path

if __name__ == '__main__':
    create_minimal_template_and_test()