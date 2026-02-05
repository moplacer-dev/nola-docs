#!/usr/bin/env python3
"""
Script to generate the Module Guide V2 DOCX template.
Run this once to create the template file that docxtpl will use.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def add_heading_style(doc, name, font_size, bold=True, color=None):
    """Add a custom heading style"""
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(font_size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = color
    return style


def create_module_guide_v2_template():
    """Create the Module Guide V2 DOCX template with Jinja2 placeholders"""
    doc = Document()

    # Set up document margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title = doc.add_heading('MODULE GUIDE V2', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Module info
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Module: ').bold = True
    p.add_run('{{ module.name }}')

    p = doc.add_paragraph()
    p.add_run('Acronym: ').bold = True
    p.add_run('{{ module.acronym }}')

    p = doc.add_paragraph()
    p.add_run('Grade Level: ').bold = True
    p.add_run('{{ module.grade_level }}')

    p = doc.add_paragraph()
    p.add_run('Module Type: ').bold = True
    p.add_run('{{ module.type }}')

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # Sessions loop
    doc.add_paragraph('{%- for session in sessions %}')

    # Session Header
    session_header = doc.add_heading('Session {{ session.number }}: {{ session.title }}', 1)

    # Introduction
    doc.add_heading('INTRODUCTION', 2)
    doc.add_paragraph('{{ session.introduction }}')
    doc.add_paragraph()

    # Assembly & Maintenance
    doc.add_heading('ASSEMBLY & MAINTENANCE', 2)

    p = doc.add_paragraph()
    p.add_run('Advance Prep:').bold = True
    doc.add_paragraph('{%- for item in session.advance_prep %}')
    doc.add_paragraph('  - {{ item }}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Station Setup:').bold = True
    doc.add_paragraph('{%- for item in session.station_setup %}')
    doc.add_paragraph('  - {{ item }}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Equipment Notes:').bold = True
    doc.add_paragraph('{%- for item in session.equipment_notes %}')
    doc.add_paragraph('  - {{ item }}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Cleanup:').bold = True
    doc.add_paragraph('{%- for item in session.cleanup %}')
    doc.add_paragraph('  - {{ item }}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()

    # Goals & Standards
    doc.add_heading('GOALS & STANDARDS', 2)

    p = doc.add_paragraph()
    p.add_run('Learning Goals:').bold = True
    doc.add_paragraph('{%- for goal in session.learning_goals %}')
    doc.add_paragraph('  {{ loop.index }}. {{ goal }}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Standards Addressed:').bold = True
    doc.add_paragraph('{%- for std in session.standards %}')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('{{ std.code }} ({{ std.type }})').bold = True
    doc.add_paragraph('{{ std.description }}')
    p = doc.add_paragraph()
    p.add_run('How It Shows Up: ').italic = True
    p.add_run('{{ std.how_it_shows_up }}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()

    # Materials
    doc.add_heading('MATERIALS', 2)

    # Create materials table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Item'
    header_cells[1].text = 'Quantity'

    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Jinja loop for materials
    doc.add_paragraph('{%- for mat in session.materials %}')
    mat_row = doc.add_paragraph('{{ mat.item }} | {{ mat.quantity }}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()

    # Safety
    doc.add_heading('SAFETY', 2)

    p = doc.add_paragraph()
    p.add_run('General Safety Rules:').bold = True
    doc.add_paragraph('{%- for rule in session.general_safety %}')
    doc.add_paragraph('  - {{ rule }}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Item-Specific Safety:').bold = True
    doc.add_paragraph('{%- for item in session.safety_items %}')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('{{ item.item }}:').bold = True
    doc.add_paragraph('{%- for warning in item.warnings %}')
    doc.add_paragraph('    - {{ warning }}')
    doc.add_paragraph('{%- endfor %}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()

    # Vocabulary
    doc.add_heading('VOCABULARY', 2)

    doc.add_paragraph('{%- for v in session.vocabulary %}')
    p = doc.add_paragraph()
    p.add_run('{{ v.term }}').bold = True
    p.add_run(' - {{ v.definition }}')
    p.add_run(' (Slides: {{ v.slides }})').italic = True
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()

    # Career Spotlight
    doc.add_heading('CAREER SPOTLIGHT', 2)

    p = doc.add_paragraph()
    p.add_run('{{ session.career.name }}, {{ session.career.title }}').bold = True
    doc.add_paragraph('{{ session.career.connection }}')

    doc.add_paragraph()

    # Teacher Tips & Assessment
    doc.add_heading('TEACHER TIPS & ASSESSMENT', 2)

    p = doc.add_paragraph()
    p.add_run('Tips:').bold = True
    doc.add_paragraph('{%- for tip in session.tips %}')
    doc.add_paragraph('  - {{ tip }}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Look For:').bold = True
    doc.add_paragraph('{%- for item in session.look_fors %}')
    doc.add_paragraph('  - {{ item }}')
    doc.add_paragraph('{%- endfor %}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Discussion Questions:').bold = True
    doc.add_paragraph('{%- for q in session.questions %}')
    doc.add_paragraph('  - {{ q }}')
    doc.add_paragraph('{%- endfor %}')

    # Session divider
    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # End session loop
    doc.add_paragraph('{%- endfor %}')

    # Save the template
    output_path = 'templates/docx_templates/module_guide_v2_master.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Template created at: {output_path}")
    return output_path


if __name__ == '__main__':
    create_module_guide_v2_template()
    print("Module Guide V2 template created successfully!")
